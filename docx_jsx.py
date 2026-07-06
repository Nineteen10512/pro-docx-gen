"""docx_jsx — 对外统一 API。

PaperJSX 语义编译架构：LLM 只写语义，不写数值。

v1.2 新增：
- load(path) 读取已有 docx 语义摘要
- update_document(original_path, edits, output_path, theme, lang) 编辑已有文档
- to_pdf(docx_path, output_dir) LibreOffice 转 PDF
- to_images(pdf_path, output_dir, dpi) pdftoppm 转图片
- extract_text(docx_path, fmt, track_changes) 提取文本（pandoc）
- list_themes() 返回可用主题列表

v1.3.0 新增（PaperJSX v1.3）：
- 原生 OMML 公式（equation 节点生成 Word 可编辑公式）
- 接受/拒绝修订：accept_all_revisions / reject_all_revisions / list_revisions / accept_revision_by_id / reject_revision_by_id
- .doc 旧格式自动转 .docx（load/update_document/to_pdf/to_images/extract_text 均支持）
- svg_shape 节点注入 SVG 装饰图形
- quality_check 质量自检（结构/排版/字数/引用/表格/图片/一致性七维度）
"""

import os
import re
import json
import shutil
import subprocess
import tempfile
from typing import Union, Optional

from . import __version__ as _PKG_VERSION
from .tokens import get_theme, merge_theme, BASE_TOKENS, list_themes as _list_themes
from .compiler.parser import parse_document, extract_outline, count_words
from .compiler.markdown_parser import markdown_to_document
from .engine.renderer import DocxRenderer


# ─── 支持的语义节点类型（v1.3 扩展） ──────────────────────────────

NODE_TYPES = [
    # v1.1
    "heading", "paragraph", "list", "table", "figure", "chart",
    "kpi_card", "callout", "page_break", "toc", "references", "appendix",
    # v1.2
    "revision", "comment", "footnote", "endnote",
    "watermark", "page_border", "equation",
    "signature_block", "signature_line",
    # v1.3
    "svg_shape",
]


# ─── 核心生成 API ────────────────────────────────────────────────

_BAD_FN_CHARS = re.compile(r'[\\/:*?"<>|\s]+')


def _default_docx_output_path(content, ext: str = "docx") -> str:
    """v1.4 P2-5: 默认输出文件名 {title}_v{version}.{ext}。"""
    title = "untitled"
    if isinstance(content, dict):
        meta = content.get("meta") or {}
        title = meta.get("title") or content.get("title") or "untitled"
    elif isinstance(content, str):
        # markdown 字符串：从首行 # 标题提取
        for line in content.splitlines():
            m = re.match(r"^#\s+(.+)$", line.strip())
            if m:
                title = m.group(1).strip()
                break
    safe = _BAD_FN_CHARS.sub("_", title).strip("_") or "untitled"
    if len(safe) > 40:
        safe = safe[:40].rstrip("_")
    parts = _PKG_VERSION.split(".")
    short_ver = ".".join(parts[:2]) if len(parts) >= 2 else _PKG_VERSION
    return f"{safe}_v{short_ver}.{ext}"


# ─── v1.5.3: 模板库 (D 能力域) ────────────────────────────────────

def _apply_template(doc: dict, template_name: Optional[str]) -> dict:
    """如果 template_name 命中注册表，则用 DOCXTemplate.default_structure
    填充 doc 中缺失的 sections/abstract/references/toc 等字段。

    约束：
    - 仅在对应字段缺失时注入，不覆盖已有内容（用户/LLM 优先）
    - template.theme_overrides 由 generate() 在 tokens 层合并
    - 模板未命中或 template_name=None 时返回原 doc 零修改（保持字节级向后兼容）
    """
    if not template_name:
        return doc
    try:
        from .templates import get_template
    except ImportError:
        return doc
    tpl = get_template(template_name)
    if tpl is None or not tpl.default_structure:
        return doc
    struct = tpl.default_structure
    # v1.5.3: 确保 doc 至少有最小 meta（title），否则 parser 校验失败
    if "meta" not in doc or not isinstance(doc.get("meta"), dict) or not doc["meta"].get("title"):
        meta = dict(doc.get("meta") or {})
        meta.setdefault("title", template_name.replace("_", " ").title())
        doc["meta"] = meta
    # 仅在缺失时注入（不覆盖）
    for k, v in struct.items():
        if k not in doc or doc.get(k) in (None, [], {}, "", False):
            doc[k] = v
    return doc


def _merge_template_overrides(theme: Union[str, dict], template_name: Optional[str]) -> Union[str, dict]:
    """把 template.theme_overrides 浅合并到 effective_theme。
    theme 是 str 时保持 str（不破坏原路径）；是 dict 时浅合并。
    """
    if not template_name:
        return theme
    try:
        from .templates import get_template as _gt
    except ImportError:
        return theme
    tpl = _gt(template_name)
    if tpl is None or not tpl.theme_overrides:
        return theme
    overrides = tpl.theme_overrides
    if isinstance(theme, dict):
        merged = dict(overrides)
        for k, v in theme.items():
            if isinstance(v, dict) and isinstance(merged.get(k), dict):
                merged[k] = {**merged[k], **v}
            else:
                merged[k] = v
        return merged
    return theme  # str 主题名不参与 override（保持路径稳定）


def _hex_to_rgb(hex_str):
    """hex str → docx.shared.RGBColor；已是 RGBColor 则原样返回；None → None。"""
    if hex_str is None:
        return None
    if isinstance(hex_str, str):
        from docx.shared import RGBColor
        try:
            return RGBColor.from_string(hex_str.lstrip("#"))
        except Exception:
            return hex_str
    return hex_str


def _normalize_theme_to_rgb(theme: dict) -> dict:
    """把 theme["color"] 里的 hex str 全部转 RGBColor。"""
    if not isinstance(theme, dict):
        return theme
    out = dict(theme)
    color = out.get("color")
    if isinstance(color, dict):
        out["color"] = {k: _hex_to_rgb(v) for k, v in color.items()}
    return out


# ─── v1.5.3: A 能力域 DOCX 端集成 ────────────────────────────────

def _extract_and_apply_template_path(
    tokens: dict,
    template_path: Optional[str],
) -> tuple[dict, Optional[dict]]:
    """从已有 .docx/.dotx/.wpt 文件提取主题并合并到 tokens。

    Returns:
        (merged_tokens, extracted_dict_or_None)
    """
    if not template_path:
        return tokens, None
    from .theme_extractor import extract_docx_theme, apply_extracted_theme, ThemeExtractionError
    try:
        extracted = extract_docx_theme(template_path)
    except (FileNotFoundError, ThemeExtractionError) as e:
        # 不静默降级，抛清晰错误
        raise ThemeExtractionError(
            f"template_path 提取失败 ({template_path}): {e}"
        ) from e
    merged = apply_extracted_theme(tokens, extracted)
    return merged, extracted


def _check_docx_template_match(
    tokens: dict,
    template_name: Optional[str],
    extracted: Optional[dict] = None,
) -> dict:
    """DOCX 端轻量模板匹配评分（0~100）。

    评分维度：
    - 颜色对比度：text vs bg（WCAG AA ≥4.5）、primary vs bg（≥3.0）
    - 字体完整性：family.heading / body / default 都有
    - 模板一致性：若传 template_name，验证 tokens.primary 与模板主题色协调

    Returns:
        {
            "match_score": int (0-100),
            "passed": bool (≥60),
            "warnings": list[str],
            "details": dict,
        }
    """
    warnings: list[str] = []
    score = 100

    color = tokens.get("color", {}) or {}
    bg = color.get("bg", "#FFFFFF")
    text = color.get("text", "#333333")
    primary = color.get("primary", "#1F3864")

    # 1) 颜色对比度
    try:
        from shared.quality import contrast_ratio
        from shared.color_palette import hex_to_rgb
        try:
            body_ratio = contrast_ratio(hex_to_rgb(text), hex_to_rgb(bg))
            if body_ratio < 4.5:
                score -= 20
                warnings.append(f"正文对比度 {body_ratio:.2f}:1 低于 WCAG AA 4.5:1")
        except Exception:
            score -= 5
            warnings.append("正文对比度计算失败")
        try:
            primary_ratio = contrast_ratio(hex_to_rgb(primary), hex_to_rgb(bg))
            if primary_ratio < 3.0:
                score -= 8
                warnings.append(f"主色对比度 {primary_ratio:.2f}:1 偏低（<3.0）")
        except Exception:
            pass
    except Exception as e:
        warnings.append(f"对比度计算异常: {e}")

    # 2) 字体完整性
    family = (tokens.get("font", {}) or {}).get("family", {}) or {}
    for key in ("heading", "body", "default"):
        if not family.get(key):
            score -= 5
            warnings.append(f"字体缺少 {key}")

    # 3) 模板一致性（若提供 template_name 或 extracted）
    if template_name and extracted and extracted.get("tokens", {}).get("color"):
        ext_primary = extracted["tokens"]["color"].get("primary", "").upper()
        cur_primary = primary.upper()
        if ext_primary and cur_primary and ext_primary != cur_primary:
            # 仅 warning，不扣太多分（用户可主动覆盖）
            warnings.append(
                f"当前 primary {cur_primary} 与模板 {template_name} 的 {ext_primary} 不一致"
            )
            score -= 3

    score = max(0, min(100, score))
    return {
        "match_score": score,
        "passed": score >= 60,
        "warnings": warnings,
        "details": {
            "text_contrast": body_ratio if 'body_ratio' in dir() else None,
        },
    }


def generate(
    content: Union[dict, str] = None,
    output_path: Optional[str] = None,
    theme: Union[str, dict] = "academic",
    lang: str = "en",
    template_name: Optional[str] = None,
    template_path: Optional[str] = None,
    auto_taste_match: bool = False,
) -> str:
    """编译语义内容并渲染为 .docx 文件。

    Compile semantic content (JSON dict or Markdown string) and render a .docx.

    Args:
        content: 语义 JSON 字典，或 Markdown 字符串。v1.5.3 起允许 None（搭配 template_name）。
        output_path: 输出 docx 文件路径。v1.4 起默认为 None，按 ``{title}_v{version}.docx`` 自动命名。
        theme: 主题名（含 v1.2 新增 10+ 主题与中文别名），或自定义 tokens dict。
        lang: 语言 "en" | "cn"，影响默认字体、首行缩进和关键词标签。
        template_name: v1.5.3 新增。可选 DOCX 模板名（academic/business/thesis_full/...）；
            命中注册表时，用模板的 default_structure 填充 doc 缺失字段，浅合并 theme_overrides。
            不传或未命中时与 v1.5 行为一致（字节级兼容）。
        template_path: v1.5.3 新增。可选 .docx/.dotx/.wpt 文件路径；
            从该文件提取主题（色板/字体）合并到 tokens。优先于 template_name 的 theme_overrides。
        auto_taste_match: v1.5.3 新增。True 时执行模板匹配评分（对比度+字体完整性+一致性），
            低分（<60）时给 warning（不阻断生成）。

    Returns:
        输出文件绝对路径。
    """
    if content is None:
        if not template_name and not template_path:
            raise TypeError(
                "content 不能为 None，除非同时传入 template_name 或 template_path。修复建议：传入语义 JSON/Markdown 字符串，或指定 template_name/template_path"
            )
        # 给 content=None 一个最小 doc 骨架，避免 parser 校验失败
        # template_name 会在后续 _apply_template 注入；template_path 也需要 meta
        from os.path import basename
        default_title = (template_name or basename(template_path or "untitled")).replace("_", " ").title()
        content = {"meta": {"title": default_title}, "sections": []}

    if isinstance(content, str):
        doc = markdown_to_document(content)
        doc["theme"] = theme
    elif isinstance(content, dict):
        doc = content
    else:
        # English: content must be a dict (semantic JSON) or str (Markdown)
        raise TypeError(
            "content 必须是 dict（语义 JSON）或 str（Markdown 文本）。"
            "修复建议：传入 Markdown 字符串，或按 SKILL.md 构造语义 JSON 对象"
        )

    # v1.5.3: 模板注入（在 output_path 和 effective_theme 计算之前）
    if template_name:
        doc = _apply_template(doc, template_name)

    if output_path is None:
        output_path = _default_docx_output_path(doc, "docx")

    effective_theme = doc.get("theme", theme)
    # v1.5.3: 模板 theme_overrides 合并
    effective_theme = _merge_template_overrides(effective_theme, template_name)
    plan = parse_document(doc)

    if isinstance(effective_theme, dict):
        # v1.5.3: 用户传 hex str 时转 RGBColor（merge_theme 不做类型转换）
        normalized = _normalize_theme_to_rgb(effective_theme)
        tokens = merge_theme(BASE_TOKENS, normalized)
    else:
        tokens = get_theme(effective_theme)

    # v1.5.3: A 能力域 — template_path 提取并合并到 tokens
    _extracted_for_match: Optional[dict] = None
    if template_path:
        tokens, _extracted_for_match = _extract_and_apply_template_path(tokens, template_path)

    # v1.5.3: A 能力域 — auto_taste_match 检查
    if auto_taste_match:
        result = _check_docx_template_match(tokens, template_name, _extracted_for_match)
        if not result["passed"]:
            # 不阻断生成，但把 warning 写到 metadata 让用户能看到
            # 通过 core_properties 的 keywords 字段承载（避免破坏 docx 结构）
            # 实际项目里调用方应读取 result；这里只把 score 写进 doc 的 meta
            doc.setdefault("meta", {})
            doc["meta"]["taste_match_score"] = result["match_score"]
            doc["meta"]["taste_match_passed"] = result["passed"]
            # 重新解析（meta 变化需反映到 plan）
            plan = parse_document(doc)

    # 应用 meta.page_setup 到 tokens.page
    page_setup = plan.get("page_setup")
    if page_setup:
        pg = dict(tokens.get("page", {}))
        for k in ("size", "orientation", "different_first_page", "different_odd_even"):
            if k in page_setup:
                pg[k] = page_setup[k]
        for k in ("margin_top", "margin_bottom", "margin_left", "margin_right",
                  "gutter", "header_distance", "footer_distance"):
            if k in page_setup:
                from docx.shared import Inches
                pg[k] = Inches(float(page_setup[k]))
        tokens["page"] = pg
        # 同步到 spacing
        from .tokens.design_tokens import PAGE_SIZES
        pw, ph = PAGE_SIZES.get(pg.get("size", "A4"), PAGE_SIZES["A4"])
        if pg.get("orientation", "portrait") == "landscape":
            pw, ph = ph, pw
        tokens["spacing"]["page_width"] = pw
        tokens["spacing"]["page_height"] = ph

    renderer = DocxRenderer(tokens, lang=lang)
    out = renderer.render(plan, output_path)
    return out


def generate_with_collaboration(
    content=None,
    output_path: Optional[str] = None,
    theme: Union[str, dict] = "academic",
    lang: str = "en",
    template_name: Optional[str] = None,
    template_path: Optional[str] = None,
    auto_taste_match: bool = False,
    enable_track_changes: bool = True,
) -> "CollaborationSession":  # noqa: F821
    """生成 docx 并立即开启协作会话（v1.5.3 C 能力域快捷入口）。

    便捷封装：调用 generate() 后用 start_collaboration() 创建协作会话。
    Returns:
        CollaborationSession（已开启 trackChanges）
    """
    from .collaboration.session import start_collaboration
    out = generate(
        content=content,
        output_path=output_path,
        theme=theme,
        lang=lang,
        template_name=template_name,
        template_path=template_path,
        auto_taste_match=auto_taste_match,
    )
    return start_collaboration(out, track_changes=enable_track_changes)


def generate_from_markdown(
    md_text: str,
    output_path: Optional[str] = None,
    theme: Union[str, dict] = "academic",
    lang: str = "zh",
    **kwargs,
) -> str:
    """从 Markdown 文本直接生成 .docx 文档（P1-2, v1.4）。

    Generate a .docx directly from a Markdown string.

    Args:
        md_text: Markdown 源文本（支持 # 标题、段落、列表、表格、引用、代码块、图片、--- 分页等）。
        output_path: 输出 .docx 路径，None 时按 ``{title}_v{version}.docx`` 自动命名。
        theme: 主题名或自定义 tokens dict。
        lang: 语言 ``"zh"`` | ``"en"``。
        **kwargs: 作为 meta 透传（如 author/subtitle/institution/date）。

    Returns:
        输出文件绝对路径。
    """
    # markdown_to_document 接受 title 与 meta_kwargs；theme 不直接传入，后续由 generate() 处理
    title = kwargs.pop("title", None)
    md_kwargs = {}
    if title:
        md_kwargs["title"] = title
    md_kwargs.update(kwargs)
    doc = markdown_to_document(md_text, **md_kwargs)
    doc["theme"] = theme
    if output_path is None:
        output_path = _default_docx_output_path(doc, "docx")
    # 复用 generate 的渲染流程：把 doc 作为 content 传入（dict 分支）
    return generate(doc, output_path=output_path, theme=theme, lang=lang)


def outline(content: Union[dict, str]) -> dict:
    """返回文档大纲（章节树）。"""
    if isinstance(content, str):
        content = markdown_to_document(content)
    return extract_outline(content)


def word_count(content: Union[dict, str], lang: str = "en") -> int:
    """估算正文字数。"""
    if isinstance(content, str):
        content = markdown_to_document(content)
    return count_words(content, lang=lang)


def list_themes() -> list[dict]:
    """返回所有可用主题的简要信息。"""
    return _list_themes()


def taste_check(
    content,
    theme: str = "academic",
    lang: str = "cn",
    strict: bool = False,
) -> dict:
    """Run taste/craft preflight on DOCX semantic content before rendering."""
    from .taste import taste_check as _taste_check

    return _taste_check(content, theme=theme, lang=lang, strict=strict)


# ─── v1.2: 文档加载与编辑 ───────────────────────────────────────

def load(docx_path: str) -> dict:
    """加载已有 .docx（或自动转换 .doc→.docx）并解析为语义结构摘要。

    Returns:
        {
            "path": str,
            "core_props": {title, author, subject, ...},
            "sections": [{"heading": str, "level": int, "paragraph_count": int, "tables": int, "start_offset": int}],
            "paragraphs": [{"style": str, "text": str, "is_heading": bool, "heading_level": int|None}],
            "tables_count": int,
            "images_count": int,
        }

    @since v1.3.0: auto-converts .doc → .docx via LibreOffice.
    """
    from .engine.doc_converter import ensure_docx
    from docx import Document as _D
    docx_path = ensure_docx(docx_path)
    if not os.path.exists(docx_path):
        # English: FileNotFoundError(docx_path) — load()
        raise FileNotFoundError(
            f"找不到文件：{docx_path}。修复建议：请检查传入的 docx 路径是否正确"
        )
    d = _D(docx_path)
    cp = d.core_properties
    paragraphs = []
    sections = []
    tables_count = len(d.tables)
    images_count = 0
    # 统计图片（inline shapes）
    try:
        images_count = len(d.inline_shapes)
    except Exception:
        pass
    current_sec = None
    for i, p in enumerate(d.paragraphs):
        style_name = p.style.name if p.style else ""
        is_heading = style_name.startswith("Heading")
        lvl = None
        if is_heading:
            try:
                lvl = int(style_name.split()[-1])
            except Exception:
                lvl = 1
            if current_sec is not None:
                sections.append(current_sec)
            current_sec = {
                "heading": p.text,
                "level": lvl or 1,
                "paragraph_count": 0,
                "tables": 0,
                "start_paragraph_index": i,
            }
        else:
            if current_sec is None:
                current_sec = {
                    "heading": "(preamble)",
                    "level": 0,
                    "paragraph_count": 0,
                    "tables": 0,
                    "start_paragraph_index": 0,
                }
            current_sec["paragraph_count"] += 1
        paragraphs.append({
            "style": style_name,
            "text": p.text,
            "is_heading": is_heading,
            "heading_level": lvl,
        })
    if current_sec is not None:
        sections.append(current_sec)

    return {
        "path": os.path.abspath(docx_path),
        "core_props": {
            "title": cp.title,
            "author": cp.author,
            "subject": cp.subject,
            "keywords": cp.keywords,
            "category": cp.category,
            "created": str(cp.created) if cp.created else None,
            "modified": str(cp.modified) if cp.modified else None,
        },
        "sections": sections,
        "paragraphs": paragraphs,
        "tables_count": tables_count,
        "images_count": images_count,
    }


def update_document(
    original_path: str,
    edits: list[dict],
    output_path: str,
    theme: str = "academic",
    lang: str = "en",
) -> str:
    """在已有 docx 上追加/替换/删除内容。

    Args:
        original_path: 原 docx 路径（v1.3: 支持 .doc 自动转换）。
        edits: 指令列表，每个指令是一个 dict，支持 action:
            - {"action": "append_section", "title": str, "content": [nodes...]}
              在文档末尾追加一个章节（Heading 1 + 内容）
            - {"action": "append_node", "node": {...}}  在文档末尾追加一个语义节点
            - {"action": "append_paragraphs", "after_heading": str|None, "content": [nodes...]}
              在指定标题后（或末尾）追加多个段落/节点
            - {"action": "replace_section", "title_match": str, "content": [nodes...], "level": int|None}
              找到匹配标题的章节，删除其内容直到下一个同级或更高级标题，再插入新内容
            - {"action": "replace_text", "find": str, "replace": str}
              全文简单字符串替换（保留首个 run 的样式）
            - {"action": "append_paragraph", "text": str, "style": str}
              追加一个简单段落
            - {"action": "delete_section", "title_match": str}
              删除指定章节
            - {"action": "accept_all_revisions"} (v1.3) 接受所有修订
            - {"action": "reject_all_revisions"} (v1.3) 拒绝所有修订
        output_path: 输出 docx 路径。
        theme: （仅在插入新内容时作为参考样式，不会修改原文档其他样式）
        lang: 语言。

    Returns:
        输出文件路径。
    """
    import copy
    from docx import Document as _D
    from .engine.doc_converter import ensure_docx
    original_path = ensure_docx(original_path)
    if not os.path.exists(original_path):
        # English: FileNotFoundError(original_path)
        raise FileNotFoundError(
            f"找不到待编辑文件：{original_path}。修复建议：请检查 original_path 是否正确"
        )

    shutil.copy2(original_path, output_path)
    d = _D(output_path)

    # 为新内容准备 tokens 和一个 helper renderer（只用来构造 OOXML，不保存）
    tokens = get_theme(theme)
    helper = DocxRenderer(tokens, lang=lang, output_dir=os.path.dirname(output_path) or ".")
    helper.doc = d  # 直接把渲染挂到现有 doc
    helper._output_dir = os.path.dirname(os.path.abspath(output_path)) or "."

    # 预处理：收集段落 → (索引, 段落对象, 标题级别)
    body = d.element.body
    # 段落索引映射
    para_map = []  # list of (paragraph_object, heading_level_or_None)
    for p in d.paragraphs:
        sn = p.style.name if p.style else ""
        lvl = None
        if sn.startswith("Heading"):
            try:
                lvl = int(sn.split()[-1])
            except Exception:
                lvl = 1
        para_map.append((p, lvl))

    def find_section_range(title_match: str, level: int | None = None):
        """返回 (start_para_index, end_para_index_exclusive)。
        end 指向下一个同级或更高级标题，或 len(para_map)。
        """
        start = None
        for i, (p, lvl) in enumerate(para_map):
            if lvl is not None and title_match in p.text:
                if level is None or lvl == level:
                    start = i
                    break
        if start is None:
            return None
        # 找到标题段落所在 body element
        start_elem = para_map[start][0]._p
        # 向后找到下一个同级或更高级标题
        end_idx = len(para_map)
        for j in range(start + 1, len(para_map)):
            pj, lvlj = para_map[j]
            if lvlj is not None and lvlj <= (level or para_map[start][1] or 1):
                end_idx = j
                break
        return start, end_idx, para_map[start][1] or 1

    def remove_paragraph_range(start_idx: int, end_idx: int, keep_title: bool = True):
        """删除 [start_idx, end_idx) 的段落。keep_title=True 时保留标题段落本身（只删内容到下一节前）。"""
        begin = start_idx + 1 if keep_title else start_idx
        for j in range(end_idx - 1, begin - 1, -1):
            p = para_map[j][0]
            p._element.getparent().remove(p._element)

    def insert_nodes_after(paragraph, nodes: list, heading_level: int = 1):
        """在指定 paragraph 之后插入若干语义节点。

        实现策略：在 paragraph 后追加一个临时占位段落，移动到其前面逐个 add 节点，
        然后删除占位段落。由于 renderer._render_node 默认通过 doc.add_* 追加到末尾，
        我们用一个 anchor 机制：先在末尾插入一个锚点段落，把新节点依次 add，然后
        把这些新产生的 XML 元素移动到 paragraph 之后，最后去掉锚点。
        """
        anchor = d.add_paragraph()
        # 记录插入前 body 中最后一个元素的位置
        body = d.element.body
        before_children = list(body)
        # 逐个渲染节点（追加到末尾）
        new_node_count = 0
        for n in nodes:
            if isinstance(n, dict) and "type" in n:
                from .compiler.parser import _expand_node
                try:
                    flat = _expand_node(n, heading_level)
                except Exception:
                    flat = [{"node_type": "paragraph", "text": str(n)}]
                for fn in flat:
                    helper._render_node(fn)
                    new_node_count += 1
            elif isinstance(n, str):
                p = d.add_paragraph(n)
                from .engine.renderer import _set_run_font
                for r in p.runs:
                    _set_run_font(r, tokens["font"]["family"]["default"], tokens["font"]["family"]["cn"])
                new_node_count += 1
        # body 末尾新增的 XML 元素 = 从 anchor 之后开始的全部（anchor 是末尾 add 的，然后又 add 了其他）
        after_children = list(body)
        anchor_idx = after_children.index(anchor._p)
        new_elements = after_children[anchor_idx + 1:]
        # 将 new_elements 移动到 paragraph 之后
        ref_elem = paragraph._p
        for el in new_elements:
            body.remove(el)
            ref_elem.addnext(el)
            ref_elem = el
        # 删除 anchor
        body.remove(anchor._p)

    # 处理编辑指令（按顺序）
    for edit in edits:
        action = edit.get("action")
        if action == "append_section":
            title = edit["title"]
            content = edit.get("content", [])
            nodes = [{"type": "heading", "level": edit.get("level", 1), "text": title}] + list(content)
            # 追加到末尾
            for n in nodes:
                if isinstance(n, dict) and "type" in n:
                    from .compiler.parser import _expand_node
                    for fn in _expand_node(n, 1):
                        helper._render_node(fn)
                elif isinstance(n, str):
                    d.add_paragraph(n)

        elif action == "append_paragraph":
            p = d.add_paragraph(edit.get("text", ""))
            from .engine.renderer import _set_run_font
            for r in p.runs:
                _set_run_font(r, tokens["font"]["family"]["default"], tokens["font"]["family"]["cn"])
                r.font.size = tokens["font"]["size"]["body"]

        elif action == "append_node":
            node = edit["node"]
            if isinstance(node, str):
                node = json.loads(node)
            if isinstance(node, dict) and "type" in node:
                from .compiler.parser import _expand_node
                for fn in _expand_node(node, 1):
                    helper._render_node(fn)

        elif action == "append_paragraphs":
            after_match = edit.get("after_heading")
            content = edit.get("content", [])
            if after_match:
                rng = find_section_range(after_match)
                if rng is None:
                    # 找不到则追加到末尾
                    target = d.paragraphs[-1] if d.paragraphs else None
                else:
                    target = para_map[rng[1] - 1][0] if rng[1] > 0 else None
                if target is not None:
                    insert_nodes_after(target, content, heading_level=1)
                else:
                    for n in content:
                        if isinstance(n, dict) and "type" in n:
                            from .compiler.parser import _expand_node
                            for fn in _expand_node(n, 1):
                                helper._render_node(fn)
            else:
                for n in content:
                    if isinstance(n, dict) and "type" in n:
                        from .compiler.parser import _expand_node
                        for fn in _expand_node(n, 1):
                            helper._render_node(fn)

        elif action == "replace_section":
            match = edit["title_match"]
            content = edit.get("content", [])
            level = edit.get("level")
            rng = find_section_range(match, level)
            if rng is None:
                raise ValueError(f"找不到章节 {match!r}。修复建议：请检查章节标题是否正确，或使用 replace_text 进行全文替换")
            start_idx, end_idx, hlevel = rng
            title_para = para_map[start_idx][0]
            # 删除标题之后到下一节之前的所有内容
            remove_paragraph_range(start_idx, end_idx, keep_title=True)
            # 注意：表格等非段落 block 也可能在该区间，我们用 XML 直接清理：
            # 删除 title_para 之后、下一个标题之前的所有 body 子元素
            self_body = title_para._p.getparent()
            siblings = list(self_body)
            try:
                ti = siblings.index(title_para._p)
            except ValueError:
                ti = 0
            # 找下一个 heading 段落（同级或更高）
            end_pos = len(siblings)
            for k in range(ti + 1, len(siblings)):
                el = siblings[k]
                if el.tag.endswith('}p'):
                    # 检查是不是 heading
                    pPr = el.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                    if pPr is not None:
                        pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                        if pStyle is not None:
                            val = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            if val and val.startswith('Heading'):
                                try:
                                    hl = int(val.replace('Heading', '').strip())
                                    if hl <= hlevel:
                                        end_pos = k
                                        break
                                except Exception:
                                    pass
            for k in range(end_pos - 1, ti, -1):
                self_body.remove(siblings[k])
            # 在标题后插入新内容
            insert_nodes_after(title_para, content, heading_level=hlevel)

        elif action == "delete_section":
            match = edit["title_match"]
            rng = find_section_range(match, edit.get("level"))
            if rng is None:
                continue
            start_idx, end_idx, hlevel = rng
            # 删除从标题开始到下一节前所有元素
            start_para = para_map[start_idx][0]
            self_body = start_para._p.getparent()
            siblings = list(self_body)
            ti = siblings.index(start_para._p)
            end_pos = len(siblings)
            for k in range(ti + 1, len(siblings)):
                el = siblings[k]
                if el.tag.endswith('}p'):
                    pPr = el.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                    if pPr is not None:
                        pStyle = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle')
                        if pStyle is not None:
                            val = pStyle.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                            if val and val.startswith('Heading'):
                                try:
                                    hl = int(val.replace('Heading', '').strip())
                                    if hl <= hlevel:
                                        end_pos = k
                                        break
                                except Exception:
                                    pass
            for k in range(end_pos - 1, ti - 1, -1):
                self_body.remove(siblings[k])

        elif action == "replace_text":
            find = edit["find"]
            repl = edit["replace"]
            # 简单替换：遍历段落，在 runs 中做文本替换（跨 run 可能遗漏，但基本可用）
            for p in d.paragraphs:
                if find in p.text:
                    # 保留第一个 run 的样式，合并所有 run 文本替换后写回
                    full = p.text
                    new_text = full.replace(find, repl)
                    if p.runs:
                        first = p.runs[0]
                        first.text = new_text
                        for r in p.runs[1:]:
                            r.text = ""
                    else:
                        p.add_run(new_text)
            # 表格单元格内文本
            for tbl in d.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            if find in p.text:
                                full = p.text
                                new_text = full.replace(find, repl)
                                if p.runs:
                                    p.runs[0].text = new_text
                                    for r in p.runs[1:]:
                                        r.text = ""
                                else:
                                    p.add_run(new_text)

        elif action == "accept_all_revisions":
            # v1.3: apply accept to current working file
            from .engine.revisions import accept_all_revisions as _accept
            d.save(output_path)
            _accept(output_path, output_path)
            d = _D(output_path)
            continue

        elif action == "reject_all_revisions":
            from .engine.revisions import reject_all_revisions as _reject
            d.save(output_path)
            _reject(output_path, output_path)
            d = _D(output_path)
            continue

        else:
            print(f"[update_document] unknown action: {action}")

    d.save(output_path)
    return output_path


# ─── v1.2: PDF/图片预览 ─────────────────────────────────────────

def to_pdf(docx_path: str, output_dir: Optional[str] = None) -> str:
    """使用 LibreOffice 将 docx 转为 PDF（v1.3: 自动转换 .doc→.docx）。

    Returns:
        生成的 PDF 路径。失败时抛出 RuntimeError。
    """
    from .engine.doc_converter import ensure_docx
    docx_path = ensure_docx(docx_path)
    if not os.path.exists(docx_path):
        # English: FileNotFoundError(docx_path) — to_pdf
        raise FileNotFoundError(
            f"找不到待转换文件：{docx_path}。修复建议：请检查 docx_path 是否存在"
        )
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError("未找到 LibreOffice (soffice)，无法转换 PDF。修复建议：请安装 LibreOffice 或将其加入 PATH")
    out_dir = output_dir or os.path.dirname(os.path.abspath(docx_path)) or "."
    os.makedirs(out_dir, exist_ok=True)
    cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(f"soffice 转换失败：{result.stderr}。修复建议：检查输入文档是否损坏或被其他程序占用")
    base = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(out_dir, base + ".pdf")
    if not os.path.exists(pdf_path):
        # 尝试从 stdout 中找到输出路径
        raise RuntimeError(f"PDF 未生成，soffice 输出：{result.stdout} {result.stderr}。修复建议：检查输出目录权限与磁盘空间")
    return pdf_path


def to_images(
    pdf_path: str,
    output_dir: Optional[str] = None,
    dpi: int = 150,
    fmt: str = "png",
) -> list[str]:
    """使用 pdftoppm 将 PDF 转为图片（v1.3: 若传入 .doc/.docx 则先转 PDF）。

    Returns:
        图片路径列表（按页序）。
    """
    from .engine.doc_converter import is_doc_path
    # If input is .doc/.docx, convert to PDF first
    lower = pdf_path.lower()
    if lower.endswith(".docx") or is_doc_path(pdf_path):
        pdf_dir = output_dir or tempfile.mkdtemp(prefix="docx_pdf_")
        pdf_path = to_pdf(pdf_path, output_dir=pdf_dir)
    if not os.path.exists(pdf_path):
        # English: FileNotFoundError(pdf_path) — to_images
        raise FileNotFoundError(
            f"找不到 PDF 文件：{pdf_path}。修复建议：请先确认 to_pdf 成功生成了 PDF"
        )
    pdftoppm = shutil.which("pdftoppm")
    if not pdftoppm:
        raise RuntimeError("未找到 pdftoppm。修复建议：请安装 poppler-utils（apt install poppler-utils 或 brew install poppler）")
    out_dir = output_dir or tempfile.mkdtemp(prefix="docx_preview_")
    os.makedirs(out_dir, exist_ok=True)
    prefix = os.path.join(out_dir, os.path.splitext(os.path.basename(pdf_path))[0])
    cmd = [pdftoppm, f"-{fmt}", "-r", str(dpi), pdf_path, prefix]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
    if result.returncode != 0:
        raise RuntimeError(f"pdftoppm 转换失败：{result.stderr}。修复建议：检查 PDF 文件是否有效")
    # 找到 prefix-*.png
    import glob
    files = sorted(glob.glob(prefix + "-*." + fmt))
    return files


# ─── v1.2: 文本提取 ─────────────────────────────────────────────

def extract_text(docx_path: str, fmt: str = "markdown", track_changes: str = "accept") -> str:
    """提取 docx 文本内容。

    Args:
        docx_path: docx 文件路径。
        fmt: "markdown"（pandoc）| "text"（纯文本，用 python-docx）。
        track_changes: "accept" | "reject" | "all"，仅 pandoc 支持。

    Returns:
        文本字符串。
    """
    if not os.path.exists(docx_path):
        # English: FileNotFoundError(docx_path) — extract_text
        raise FileNotFoundError(
            f"找不到文件：{docx_path}。修复建议：请检查 docx_path 路径是否正确"
        )
    from .engine.doc_converter import ensure_docx
    docx_path = ensure_docx(docx_path)
    pandoc = shutil.which("pandoc")
    if fmt == "markdown" and pandoc:
        # pandoc 提取 markdown
        track_flag = {
            "accept": "--track-changes=accept",
            "reject": "--track-changes=reject",
            "all": "--track-changes=all",
        }.get(track_changes, "--track-changes=accept")
        cmd = [pandoc, "-f", "docx", "-t", "markdown", track_flag, docx_path]
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            if result.returncode == 0:
                return result.stdout
        except Exception:
            pass
    # fallback：python-docx 纯文本
    from docx import Document as _D
    d = _D(docx_path)
    parts = []
    for p in d.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for tbl in d.tables:
        for row in tbl.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)


# ─── v1.3.0: 修订 API（包装 engine.revisions） ─────────────────

def _require_docx_path(p: str) -> str:
    if not isinstance(p, str) or not os.path.exists(p):
        raise FileNotFoundError(f"找不到 docx 文件：{p!r}。修复建议：请确认文件路径与扩展名是否正确")
    from .engine.doc_converter import ensure_docx
    return ensure_docx(p)


def list_revisions(docx_path: str) -> list:
    """List all tracked changes in a docx file.

    Returns list of dicts with keys: id, type (ins/del), author, date, text.
    """
    p = _require_docx_path(docx_path)
    from .engine.revisions import list_revisions as _lr
    return _lr(p)


def accept_all_revisions(docx_path: str, output_path: Optional[str] = None) -> str:
    """Accept all tracked changes and save to output_path (or in-place if None)."""
    p = _require_docx_path(docx_path)
    from .engine.revisions import accept_all_revisions as _aar
    return _aar(p, output_path=output_path)


def reject_all_revisions(docx_path: str, output_path: Optional[str] = None) -> str:
    """Reject all tracked changes and save to output_path (or in-place if None)."""
    p = _require_docx_path(docx_path)
    from .engine.revisions import reject_all_revisions as _rar
    return _rar(p, output_path=output_path)


def accept_revision_by_id(docx_path: str, rev_id, output_path: Optional[str] = None) -> str:
    """Accept a single revision by its id (str or int)."""
    p = _require_docx_path(docx_path)
    from .engine.revisions import accept_revision_by_id as _a
    return _a(p, rev_id=rev_id, output_path=output_path)


def reject_revision_by_id(docx_path: str, rev_id, output_path: Optional[str] = None) -> str:
    """Reject a single revision by its id."""
    p = _require_docx_path(docx_path)
    from .engine.revisions import reject_revision_by_id as _r
    return _r(p, rev_id=rev_id, output_path=output_path)


# ─── v1.3.0: quality_check ──────────────────────────────────────

def quality_check(doc, theme: str = "academic", lang: str = "cn") -> dict:
    """Run seven-dimension QA self-check on a document spec.

    七维质量自检：结构完整性、排版规范、字数达标、色板对比度、引用格式、图表存在性、
    前后一致性。返回 score + 各维度 pass/fail + suggestions 列表。"""
    try:
        from ..shared.quality import (
            contrast_ratio, hex_to_rgb, count_words, count_chars_cjk,
            weighted_score,
        )
    except ImportError:
        try:
            from shared.quality import (
                contrast_ratio, hex_to_rgb, count_words, count_chars_cjk,
                weighted_score,
            )
        except ImportError:
            try:
                from skills.shared.quality import (
                    contrast_ratio, hex_to_rgb, count_words, count_chars_cjk,
                    weighted_score,
                )
            except ImportError:
                return {"total_score": 0, "passed": False,
                        "error": "shared.quality module not available",
                        "dimensions": {}}

    # If doc is a path, load into semantic dict via python-docx heuristics
    if isinstance(doc, str):
        if not os.path.exists(doc):
            raise FileNotFoundError(f"找不到文档文件：{doc}。修复建议：请传入语义 JSON dict 或存在的文件路径")
        from .engine.doc_converter import ensure_docx
        p = ensure_docx(doc)
        from docx import Document as _D
        d = _D(p)
        nodes = []
        for para in d.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            if "heading" in style or "title" in style:
                try:
                    lvl = int(''.join(c for c in style if c.isdigit()) or "1")
                except Exception:
                    lvl = 1
                nodes.append({"type":"heading","level":lvl,"text":text})
            else:
                nodes.append({"type":"paragraph","text":text})
        for tbl in d.tables:
            rows = []
            for row in tbl.rows:
                rows.append([c.text for c in row.cells])
            if rows:
                nodes.append({"type":"table","headers":rows[0],"rows":rows[1:]})
        doc = {"meta":{"title":os.path.basename(p)},"content":nodes}

    if not isinstance(doc, dict):
        raise TypeError("quality_check 入参必须是语义 JSON dict 或 docx/doc 文件路径。修复建议：传入 generate() 返回的 dict，或有效的 .docx/.doc 文件路径")

    # Flatten sections[].content → a single nodes list for QA
    nodes = doc.get("content") or doc.get("nodes") or []
    if not nodes:
        sects = doc.get("sections") or []
        flat = []
        for s in sects:
            if not isinstance(s, dict):
                continue
            t = s.get("title")
            if t:
                flat.append({"type": "heading", "level": 1, "text": t})
            flat.extend(s.get("content") or [])
        nodes = flat

    # Resolve theme colors
    resolved = get_theme(theme)
    color = resolved.get("color", {})
    bg_hex = color.get("bg", "#FFFFFF")
    text_hex = color.get("text", "#333333")
    heading_hex = color.get("heading", "#1F3864")
    def _to_hex(c):
        if isinstance(c, str) and c.startswith("#") and len(c) == 7:
            return c
        try:
            return "#{:02X}{:02X}{:02X}".format(c[0], c[1], c[2])
        except Exception:
            return None
    bg_hex = _to_hex(bg_hex) or "#FFFFFF"
    text_hex = _to_hex(text_hex) or "#333333"
    heading_hex = _to_hex(heading_hex) or "#1F3864"

    dimensions = {
        "structure": {"score": 100, "issues": []},
        "typography": {"score": 100, "issues": []},
        "word_count": {"score": 100, "issues": []},
        "color_contrast": {"score": 100, "issues": []},
        "references": {"score": 100, "issues": []},
        "tables_figures": {"score": 100, "issues": []},
        "consistency": {"score": 100, "issues": []},
    }

    n_headings = 0; n_paragraphs = 0; n_tables = 0; n_figures = 0; n_charts = 0
    n_citations = 0; n_eq = 0; n_long_para = 0
    total_words = 0
    heading_levels = []
    has_h1 = False

    for i, node in enumerate(nodes):
        if not isinstance(node, dict):
            continue
        t = node.get("type","")
        if t == "heading":
            n_headings += 1
            lvl = int(node.get("level",1))
            heading_levels.append(lvl)
            if lvl == 1:
                has_h1 = True
            txt = node.get("text","")
            if txt and count_words(txt, lang) > 35:
                dimensions["typography"]["score"] -= 4
                dimensions["typography"]["issues"].append({"index":i,"type":"heading","problem":"标题过长 (>35字)","suggestion":"精简标题至20字以内"})
        elif t in ("paragraph",):
            n_paragraphs += 1
            txt = node.get("text","")
            w = count_words(txt, lang)
            total_words += w
            if w > 250:
                n_long_para += 1
                dimensions["typography"]["score"] -= 5
                dimensions["typography"]["issues"].append({"index":i,"problem":f"段落过长({w}字)","suggestion":"拆分为多个短段落"})
            if "@" in txt or "et al" in txt or "（" in txt and "）" in txt:
                n_citations += 1
        elif t == "list":
            items = node.get("items") or []
            for it in items:
                tt = it.get("text","") if isinstance(it,dict) else str(it)
                total_words += count_words(tt, lang)
        elif t == "table":
            n_tables += 1
            rows = node.get("rows") or []
            headers = node.get("headers") or []
            if headers and rows:
                col_count = len(headers)
                for r in rows:
                    if len(r) != col_count:
                        dimensions["tables_figures"]["score"] -= 6
                        dimensions["tables_figures"]["issues"].append({"index":i,"problem":"表格行列数不齐","suggestion":"统一行列数"})
                        break
        elif t == "figure":
            n_figures += 1
            cap = node.get("caption","")
            if not cap:
                dimensions["tables_figures"]["score"] -= 4
                dimensions["tables_figures"]["issues"].append({"index":i,"problem":"图片缺少图题","suggestion":"添加caption字段"})
        elif t == "chart":
            n_charts += 1
            cap = node.get("caption","")
            if not cap:
                dimensions["tables_figures"]["score"] -= 3
        elif t == "references":
            n_citations += 1
        elif t == "equation":
            n_eq += 1

    # structure: must have at least one heading
    if not has_h1:
        dimensions["structure"]["score"] -= 15
        dimensions["structure"]["issues"].append({"problem":"缺少一级标题(heading level 1)","suggestion":"文档开头添加文档标题"})
    if n_headings < 2 and total_words > 200:
        dimensions["structure"]["score"] -= 8
        dimensions["structure"]["issues"].append({"problem":"章节数过少","suggestion":"使用heading切分为多个章节"})
    # Heading hierarchy monotonicity
    prev = 0
    for lv in heading_levels:
        if prev and lv - prev > 1:
            dimensions["structure"]["score"] -= 4
            dimensions["structure"]["issues"].append({"problem":f"标题层级跳级(从 H{prev} 跳到 H{lv})","suggestion":"按 H1→H2→H3 顺序组织"})
        prev = lv

    # word_count: doc target ~ 1000-5000 for formal doc
    if total_words and total_words < 300:
        dimensions["word_count"]["score"] -= 10
        dimensions["word_count"]["issues"].append({"problem":f"正文偏短({total_words}字)","suggestion":"补充内容至1000字以上"})

    # color contrast
    try:
        bg_rgb = hex_to_rgb(bg_hex); txt_rgb = hex_to_rgb(text_hex); h_rgb = hex_to_rgb(heading_hex)
        body_ratio = contrast_ratio(txt_rgb, bg_rgb)
        h_ratio = contrast_ratio(h_rgb, bg_rgb)
        if body_ratio < 4.5:
            dimensions["color_contrast"]["score"] -= 15
            dimensions["color_contrast"]["issues"].append({"problem":"正文对比度不足","ratio":round(body_ratio,2),"suggestion":"WCAG AA 要求 ≥4.5:1"})
        if h_ratio < 4.5:
            dimensions["color_contrast"]["score"] -= 8
    except Exception:
        pass

    # references: if doc has citations but no references section
    if n_citations >= 2 and not any(n.get("type")=="references" for n in nodes):
        dimensions["references"]["score"] -= 12
        dimensions["references"]["issues"].append({"problem":"正文含引用但缺少 references 节点","suggestion":"在文末添加参考文献节点"})

    # consistency: table/figure numbering (simple check)
    fig_cap_nums = []; tbl_cap_nums = []
    for node in nodes:
        if not isinstance(node, dict): continue
        cap = node.get("caption","")
        import re
        m = re.match(r"(图|Fig\.?|Figure)\s*(\d+)", cap)
        if m: fig_cap_nums.append(int(m.group(2)))
        m = re.match(r"(表|Tab\.?|Table)\s*(\d+)", cap)
        if m: tbl_cap_nums.append(int(m.group(2)))
    for nums,name in ((fig_cap_nums,"图"),(tbl_cap_nums,"表")):
        if nums and nums != list(range(1,len(nums)+1)):
            dimensions["consistency"]["score"] -= 5
            dimensions["consistency"]["issues"].append({"problem":f"{name}编号不连续","current":nums,"suggestion":f"按 {name}1,{name}2...顺序编号"})

    # clamp
    for k in dimensions:
        dimensions[k]["score"] = max(0, min(100, int(dimensions[k]["score"])))

    weights = {
        "structure": 20, "typography": 20, "word_count": 10,
        "color_contrast": 15, "references": 10, "tables_figures": 15,
        "consistency": 10,
    }
    total, _ = weighted_score(dimensions, weights)
    return {
        "total_score": total,
        "passed": total >= 75,
        "dimensions": dimensions,
        "weights": weights,
        "stats": {
            "headings": n_headings, "paragraphs": n_paragraphs,
            "tables": n_tables, "figures": n_figures, "charts": n_charts,
            "equations": n_eq, "citations": n_citations,
            "total_words": total_words,
        },
    }
