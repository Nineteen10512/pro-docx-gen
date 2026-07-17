import base64
import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from pro_docx_gen.docx_qa import audit_docx_layout, expand_review_pages


class LayoutQATests(unittest.TestCase):
    def build_defective_document(self, path: Path) -> None:
        image_path = path.with_suffix(".png")
        image_path.write_bytes(
            base64.b64decode(
                "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
            )
        )
        document = Document()
        for text, size in (("Title", 26), ("Rule", 9), ("Subtitle", 13.5), ("Author", 10.5)):
            paragraph = document.add_paragraph(text)
            paragraph.runs[0].font.size = Pt(size)
        first = document.add_paragraph("正文小号")
        first.runs[0].font.size = Pt(8)
        second = document.add_paragraph("正文大号")
        second.runs[0].font.size = Pt(14)
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        document.add_heading("1. 章节", level=1)
        document.add_picture(str(image_path), width=Inches(2.4))

        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "边框"
        tbl_borders = OxmlElement("w:tblBorders")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:color"), "184D43")
        tbl_borders.append(top)
        table._tbl.tblPr.append(tbl_borders)
        cell_borders = OxmlElement("w:tcBorders")
        cell_top = OxmlElement("w:top")
        cell_top.set(qn("w:val"), "single")
        cell_top.set(qn("w:color"), "FF0000")
        cell_borders.append(cell_top)
        table.cell(0, 0)._tc.get_or_add_tcPr().append(cell_borders)

        field_paragraph = document.add_paragraph()
        run = OxmlElement("w:r")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGEREF _Toc99999 \\h "
        run.append(instruction)
        field_paragraph._p.append(run)
        document.save(path)

    def test_audit_detects_known_docx_layout_failures(self):
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "defective.docx"
            self.build_defective_document(path)
            report = audit_docx_layout(path)
            codes = {issue["code"] for issue in report["issues"]}
            self.assertIn("cell_border_color_conflict", codes)
            self.assertIn("figure_too_small", codes)
            self.assertIn("dangling_toc_bookmark", codes)
            self.assertIn("redundant_page_break", codes)
            self.assertIn("role_font_size_drift", codes)
            self.assertTrue(report["requires_full_figure_review"])

    def test_figure_review_targets_include_neighbors(self):
        self.assertEqual(expand_review_pages([1, 5, 10], 10), [1, 2, 4, 5, 6, 9, 10])

    def test_cover_typography_is_not_reported_as_body_font_drift(self):
        with tempfile.TemporaryDirectory() as temp:
            path = Path(temp) / "cover.docx"
            document = Document()
            for text, size in (("Title", 26), ("Rule", 9), ("Subtitle", 13.5), ("Author", 10.5)):
                paragraph = document.add_paragraph(text)
                paragraph.runs[0].font.size = Pt(size)
            for text in ("Body one", "Body two"):
                paragraph = document.add_paragraph(text)
                paragraph.runs[0].font.size = Pt(10.5)
            document.save(path)
            report = audit_docx_layout(path)
            codes = {issue["code"] for issue in report["issues"]}
            self.assertNotIn("role_font_size_drift", codes)


if __name__ == "__main__":
    unittest.main()
