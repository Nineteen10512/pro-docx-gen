"""Smoke tests for PRO-DOCX v1.6.6 release packages."""

from __future__ import annotations

import argparse
import inspect
import json
import os
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]


def _run(cmd: list[str], cwd: Path | None = None) -> None:
    proc = subprocess.run(cmd, cwd=cwd or ROOT, text=True, encoding="utf-8", errors="replace", capture_output=True)
    if proc.returncode:
        print(proc.stdout)
        print(proc.stderr, file=sys.stderr)
        raise SystemExit(proc.returncode)


def _assert(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def test_imports() -> None:
    sys.path.insert(0, str(ROOT))
    import PIL
    import pro_docx_gen

    _assert(bool(PIL.__version__), "Pillow is not importable")
    _assert(pro_docx_gen.__version__ == "1.6.6", f"wrong version: {pro_docx_gen.__version__}")
    for name in ("generate", "generate_from_markdown", "quality_check", "taste_check"):
        _assert(hasattr(pro_docx_gen, name), f"missing public API: {name}")


def test_package_only_import(output_dir: Path) -> None:
    isolated = output_dir / "package_only_import"
    if isolated.exists():
        shutil.rmtree(isolated)
    isolated.mkdir(parents=True)
    shutil.copytree(ROOT / "pro_docx_gen", isolated / "pro_docx_gen")

    env = os.environ.copy()
    env["PYTHONPATH"] = str(isolated)
    env["PYTHONIOENCODING"] = "utf-8"
    code = (
        "import pro_docx_gen; "
        "from pro_docx_gen import generate; "
        "assert pro_docx_gen.__version__ == '1.6.6'; "
        "assert callable(generate); "
        "assert len(pro_docx_gen.list_templates()) >= 10; "
        "print('package-only import ok')"
    )
    proc = subprocess.run(
        [sys.executable, "-c", code],
        cwd=output_dir,
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
        env=env,
    )
    if proc.returncode:
        print(proc.stdout)
        print(proc.stderr, file=sys.stderr)
        raise SystemExit(proc.returncode)


def test_docx_quality_gate(output_dir: Path) -> None:
    docx_path = output_dir / "smoke_visual.docx"
    doc = Document()
    doc.add_heading("PRO-DOCX v1.6.6 Smoke", level=1)
    doc.add_paragraph("This document verifies import, document packaging, table geometry, and render quality gate.")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Quality gate"
    table.cell(1, 1).text = "Strict pass"
    doc.save(docx_path)

    _run(
        [
            sys.executable,
            str(ROOT / "quality_gates" / "run_quality_gate.py"),
            str(docx_path),
            "--json-report",
            str(output_dir / "quality_report.json"),
            "--output-dir",
            str(output_dir / "quality"),
        ]
    )


def test_agent_entrypoint_workflow(output_dir: Path) -> None:
    _run([sys.executable, str(ROOT / "agent_docx.py"), "doctor"])
    _run(
        [
            sys.executable,
            str(ROOT / "agent_docx.py"),
            "portability-check",
            "--json-report",
            str(output_dir / "portability_report.json"),
        ]
    )
    _run([sys.executable, str(ROOT / "agent_docx.py"), "start", "smoke-agent-entrypoint"])

    docx_path = output_dir / "agent_entrypoint.docx"
    doc = Document()
    doc.add_heading("Agent Entrypoint Smoke", level=1)
    doc.add_paragraph("This document verifies the one-command agent workflow.")
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Status"
    table.cell(1, 1).text = "Pass"
    doc.save(docx_path)

    _run(
        [
            sys.executable,
            str(ROOT / "agent_docx.py"),
            "deliver",
            str(docx_path),
            "--job",
            "smoke-agent-entrypoint",
        ]
    )
    manifest = ROOT / "_work" / "jobs" / "smoke-agent-entrypoint" / "submission_manifest.json"
    _assert(manifest.exists(), "agent entrypoint did not write submission manifest")
    data = json.loads(manifest.read_text(encoding="utf-8"))
    pngs = data.get("preview_pngs") or []
    _assert(len(pngs) == 1, f"expected one sampled preview, got {len(pngs)}")
    _assert(Path(pngs[0]).exists(), "sampled submission preview is missing")
    sampling = data.get("visual_review_sampling") or {}
    _assert(sampling.get("group_size") == 3, "visual review group size is not three pages")
    _assert(sampling.get("minimum_total") == 1, "visual review minimum is not one page")
    _assert(sampling.get("sampled_page_numbers") == [1], "single-page sampling did not select page 1")
    _assert(data.get("review_pdf") is None, "single sample should not create a review PDF")
    _run([sys.executable, str(ROOT / "agent_docx.py"), "read-text", str(ROOT / "pro_docx_gen" / "SKILL.md"), "--lines", "3"])


def test_sampled_four_up_review_bundle(output_dir: Path) -> None:
    from PIL import Image

    sys.path.insert(0, str(ROOT))
    import agent_docx

    review_dir = output_dir / "sampled_review_bundle"
    review_dir.mkdir(parents=True, exist_ok=True)
    previews = []
    for page_number in range(1, 14):
        preview = review_dir / f"page-{page_number:02d}.png"
        Image.new("RGB", (420, 594), (245, 245, 245)).save(preview)
        previews.append((page_number, str(preview)))

    sampled = agent_docx._select_sample_page_previews(previews)
    _assert([page for page, _ in sampled] == [1, 4, 7, 10, 13], "one-in-three sampling is incorrect")
    bundle = agent_docx._build_four_up_review_bundle(sampled, review_dir, "smoke")
    _assert(Path(bundle["pdf"]).exists(), "four-up review PDF is missing")
    _assert(len(bundle["sheet_pngs"]) == 2, "five sampled pages should produce two four-up sheets")
    _assert(all(Path(path).exists() for path in bundle["sheet_pngs"]), "four-up sheet PNG is missing")


def test_agent_deliver_rejects_no_render(output_dir: Path) -> None:
    docx_path = output_dir / "no_render_rejected.docx"
    doc = Document()
    doc.add_heading("No Render Rejection", level=1)
    doc.add_paragraph("Final delivery must not bypass PNG preview generation.")
    doc.save(docx_path)

    proc = subprocess.run(
        [sys.executable, str(ROOT / "agent_docx.py"), "deliver", str(docx_path), "--no-render"],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "agent deliver --no-render bypassed final visual review")
    _assert("final delivery requires render verification" in proc.stdout, "no-render rejection did not explain required render gate")


def test_portability_backend_guidance() -> None:
    text = (ROOT / "agent_docx.py").read_text(encoding="utf-8")
    _assert("Install WPS or Microsoft Word" in text, "missing selective page export install guidance")
    _assert("selective page export backend" in text, "missing selective rendering requirement")
    _assert("Install LibreOffice" not in text, "strict delivery still recommends full-document fallback")
    _assert("install PyMuPDF" in text, "missing PDF-to-PNG backend install guidance")
    _assert("Do not skip render verification" in text, "missing no-skip render warning")


def test_shared_compat_wrappers() -> None:
    sys.path.insert(0, str(ROOT))
    import shared.citation as citation
    import shared.color_palette as color_palette
    import shared.svg_engine as svg_engine

    _assert(hasattr(citation, "format_reference"), "shared.citation wrapper missing format_reference")
    _assert(hasattr(color_palette, "hex_to_rgb"), "shared.color_palette wrapper missing hex_to_rgb")
    _assert(hasattr(svg_engine, "svg_to_docx_drawing"), "shared.svg_engine wrapper missing svg_to_docx_drawing")


def test_chart_assets_do_not_leak(output_dir: Path) -> None:
    sys.path.insert(0, str(ROOT))
    from pro_docx_gen import generate

    leaked_assets = output_dir / "chart_assets"
    if leaked_assets.exists():
        raise AssertionError(f"pre-existing leaked chart assets block test: {leaked_assets}")

    docx_path = output_dir / "chart_no_leak.docx"
    generate(
        {
            "meta": {"title": "Chart Asset Cleanup Smoke", "author": "PRO-DOCX"},
            "sections": [
                {
                    "title": "Chart",
                    "level": 1,
                    "content": [
                        {
                            "type": "chart",
                            "chart_type": "column",
                            "title": "Revenue",
                            "categories": ["A", "B", "C"],
                            "series": [{"name": "Actual", "values": [1, 3, 2]}],
                        }
                    ],
                }
            ],
        },
        str(docx_path),
        theme="premium",
        lang="en",
        auto_style=False,
    )
    _assert(docx_path.exists(), "chart smoke docx was not generated")
    _assert(not leaked_assets.exists(), "chart renderer leaked chart_assets beside output docx")
    with zipfile.ZipFile(docx_path) as zf:
        media = [name for name in zf.namelist() if name.startswith("word/media/")]
    _assert(media, "chart was not embedded in docx media")


def test_chart_readability_helpers() -> None:
    sys.path.insert(0, str(ROOT))
    from pro_docx_gen.engine import chart_renderer

    chart_renderer._require_matplotlib()
    labels = [
        "EP-8 1-8x28 FFP LPVO",
        "LH-6 1-6x24 SFP LPVO",
        "SLx 1-8x24 FFP",
        "Strike Eagle 1-8x24 FFP",
        "TANGO-MSR 1-8x24",
        "Arrowhead 1-8x24",
        "RT-6 1-6x24",
    ]
    _assert(
        chart_renderer._should_auto_bar("column", labels, [{"name": "Max", "values": [8, 6, 8, 8, 8, 8, 6]}], 9, 6.27),
        "dense single-series column chart should auto-switch to horizontal bar",
    )
    bases = chart_renderer._baseline_array(0, 7)
    _assert(len(bases) == 7 and all(v == 0 for v in bases), "scalar bar baseline should expand for every data label")


def test_preinsertion_figure_gate(output_dir: Path) -> None:
    sys.path.insert(0, str(ROOT))
    from PIL import Image, ImageDraw
    from pro_docx_gen.figure_preflight import FigureAssetGateError, assert_figure_asset_ready

    bad = output_dir / "preflight_bad_chart.png"
    image = Image.new("RGB", (2000, 1000), "white")
    ImageDraw.Draw(image).rectangle((850, 430, 1150, 570), fill="#184D43")
    image.save(bad)
    try:
        assert_figure_asset_ready(
            bad,
            display_width_inches=6.0,
            source_width_inches=12.0,
            declared_min_font_pt=8.0,
            contains_text=True,
        )
    except FigureAssetGateError as exc:
        codes = {issue["code"] for issue in exc.report["issues"]}
        _assert("excessive_canvas_whitespace" in codes, "blank-canvas defect was not blocked")
        _assert("effective_text_too_small" in codes, "tiny effective font was not blocked")
        _assert(exc.report["remediation_steps"], "gate did not tell the agent how to rerender")
    else:
        raise AssertionError("pre-insertion gate allowed an illegible chart")


def test_quality_gate_blocks_warning(output_dir: Path) -> None:
    bad_path = output_dir / "warning_should_fail.docx"
    doc = Document()
    doc.add_paragraph("TODO")
    doc.save(bad_path)

    proc = subprocess.run(
        [sys.executable, str(ROOT / "quality_gates" / "run_quality_gate.py"), str(bad_path), "--no-render"],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "quality gate allowed warning-only document to pass")
    _assert("placeholder_text" in proc.stdout, "warning code not reported")


def test_quality_gate_auto_fixes_narrow_table_header(output_dir: Path) -> None:
    bad_path = output_dir / "narrow_header_should_fail.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    headers = ["Strategy", "Move", "Why it follows from SWOT"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        cell.text = header
        cell.width = Inches(0.35 if idx == 0 else 2.6)
    for idx, value in enumerate(["SO", "Use value proof.", "Avoid narrow header wrapping."]):
        cell = table.cell(1, idx)
        cell.text = value
        cell.width = Inches(0.35 if idx == 0 else 2.6)
    doc.save(bad_path)

    proc = subprocess.run(
        [sys.executable, str(ROOT / "quality_gates" / "run_quality_gate.py"), str(bad_path), "--no-render"],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "quality gate allowed narrow table header to pass")
    _assert("table_header_narrow" in proc.stdout, "narrow table header error not reported")

    _run(
        [
            sys.executable,
            str(ROOT / "quality_gates" / "run_quality_gate.py"),
            str(bad_path),
            "--no-render",
            "--auto-fix",
        ]
    )


def test_quality_gate_auto_fixes_splittable_table_rows(output_dir: Path) -> None:
    bad_path = output_dir / "splittable_rows_should_fail.docx"
    doc = Document()
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    for idx, header in enumerate(["Brand", "Signal", "Impact"]):
        table.cell(0, idx).text = header
    long_text = "Long row content should not split across pages because it damages visual readability and table meaning."
    for row_idx in range(1, 5):
        table.cell(row_idx, 0).text = f"Brand {row_idx}"
        table.cell(row_idx, 1).text = long_text
        table.cell(row_idx, 2).text = long_text
    doc.save(bad_path)

    proc = subprocess.run(
        [sys.executable, str(ROOT / "quality_gates" / "run_quality_gate.py"), str(bad_path), "--no-render"],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "quality gate allowed splittable long table rows to pass")
    _assert("table_row_can_split" in proc.stdout, "splittable row warning not reported")

    _run(
        [
            sys.executable,
            str(ROOT / "quality_gates" / "run_quality_gate.py"),
            str(bad_path),
            "--no-render",
            "--auto-fix",
        ]
    )


def test_quality_gate_auto_fixes_callout_split_risk(output_dir: Path) -> None:
    bad_path = output_dir / "callout_split_should_fail.docx"
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = (
        "Source confidence rule: medium confidence rows need review before purchase decisions, "
        "and this note must stay together as one visual callout block."
    )
    doc.save(bad_path)

    proc = subprocess.run(
        [sys.executable, str(ROOT / "quality_gates" / "run_quality_gate.py"), str(bad_path), "--no-render"],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "quality gate allowed a splittable callout block to pass")
    _assert("callout_row_can_split" in proc.stdout, "callout split risk was not reported")

    _run(
        [
            sys.executable,
            str(ROOT / "quality_gates" / "run_quality_gate.py"),
            str(bad_path),
            "--no-render",
            "--auto-fix",
        ]
    )


def test_quality_gate_auto_fixes_prose_table_orphan_risk(output_dir: Path) -> None:
    bad_path = output_dir / "prose_table_orphan_should_fail.docx"
    doc = Document()
    table = doc.add_table(rows=5, cols=2)
    table.cell(0, 0).text = "Quadrant"
    table.cell(0, 1).text = "Findings"
    long_findings = (
        "This SWOT finding contains multiple clauses and should not create an orphan row on the next page. "
        "The table needs keep-with-next protection so all compact quadrants read as one block."
    )
    for idx, label in enumerate(["Strength", "Weakness", "Opportunity", "Threat"], 1):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = long_findings
    doc.save(bad_path)

    proc = subprocess.run(
        [sys.executable, str(ROOT / "quality_gates" / "run_quality_gate.py"), str(bad_path), "--no-render"],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "quality gate allowed prose table orphan risk to pass")
    _assert("table_orphan_split_risk" in proc.stdout, "prose table orphan risk was not reported")

    _run(
        [
            sys.executable,
            str(ROOT / "quality_gates" / "run_quality_gate.py"),
            str(bad_path),
            "--no-render",
            "--auto-fix",
        ]
    )


def test_quality_gate_auto_fixes_narrow_table_body_word(output_dir: Path) -> None:
    bad_path = output_dir / "narrow_body_word_should_fail.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    for idx, header in enumerate(["Brand", "Signal", "Impact"]):
        cell = table.cell(0, idx)
        cell.text = header
        cell.width = Inches(0.3 if idx == 0 else 2.4)
    for idx, value in enumerate(["Nightforce", "Premium durability benchmark.", "Threat"]):
        cell = table.cell(1, idx)
        cell.text = value
        cell.width = Inches(0.3 if idx == 0 else 2.4)
    doc.save(bad_path)

    proc = subprocess.run(
        [sys.executable, str(ROOT / "quality_gates" / "run_quality_gate.py"), str(bad_path), "--no-render"],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "quality gate allowed narrow body word to pass")
    _assert("table_body_word_narrow" in proc.stdout, "narrow body word error not reported")

    _run(
        [
            sys.executable,
            str(ROOT / "quality_gates" / "run_quality_gate.py"),
            str(bad_path),
            "--no-render",
            "--auto-fix",
        ]
    )


def test_quality_gate_blocks_short_word_and_price_breaks(output_dir: Path) -> None:
    bad_path = output_dir / "short_word_price_breaks_should_fail.docx"
    doc = Document()
    table = doc.add_table(rows=2, cols=4)
    table.autofit = False
    for idx, header in enumerate(["Brand", "Price", "Plane", "Status"]):
        cell = table.cell(0, idx)
        cell.text = header
        cell.width = Inches(0.28)
    for idx, value in enumerate(["Vector", "$169.50", "Plane", "Context"]):
        cell = table.cell(1, idx)
        cell.text = value
        cell.width = Inches(0.28)
    doc.save(bad_path)

    proc = subprocess.run(
        [sys.executable, str(ROOT / "quality_gates" / "run_quality_gate.py"), str(bad_path), "--no-render"],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "quality gate allowed short words/prices to visually break")
    _assert("table_body_word_narrow" in proc.stdout, "short word/price break error not reported")

    _run(
        [
            sys.executable,
            str(ROOT / "quality_gates" / "run_quality_gate.py"),
            str(bad_path),
            "--no-render",
            "--auto-fix",
        ]
    )


def test_quality_gate_blocks_stale_output_after_generation_failure(output_dir: Path) -> None:
    stale_path = output_dir / "stale_output.docx"
    doc = Document()
    doc.add_paragraph("This file existed before a failed generation command.")
    doc.save(stale_path)
    cutoff = stale_path.stat().st_mtime + 5
    proc = subprocess.run(
        [
            sys.executable,
            str(ROOT / "quality_gates" / "run_quality_gate.py"),
            str(stale_path),
            "--no-render",
            "--created-after",
            str(cutoff),
        ],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "quality gate allowed stale output to pass")
    _assert("stale_output" in proc.stdout, "stale output error not reported")


def test_safe_save_reports_locked_destination(output_dir: Path) -> None:
    sys.path.insert(0, str(ROOT))
    from pro_docx_gen.engine import renderer

    target = output_dir / "locked_destination.docx"
    doc = Document()
    doc.add_paragraph("safe save")
    original_replace = renderer.os.replace

    def fake_replace(src, dst):
        raise PermissionError("simulated lock")

    renderer.os.replace = fake_replace
    try:
        try:
            renderer._safe_save_document(doc, str(target))
        except RuntimeError as exc:
            _assert("Target DOCX is locked" in str(exc), "locked-save error was not actionable")
        else:
            raise AssertionError("safe save did not fail on simulated locked destination")
    finally:
        renderer.os.replace = original_replace


def test_text_encoding_gate(output_dir: Path) -> None:
    bad_path = output_dir / "mojibake_should_fail.md"
    bad_path.write_text("Broken text: \ufffd\n", encoding="utf-8")
    try:
        proc = subprocess.run(
            [sys.executable, str(ROOT / "quality_gates" / "check_text_encoding.py"), str(bad_path)],
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
        )
        _assert(proc.returncode != 0, "encoding gate allowed mojibake to pass")
        _assert("mojibake" in proc.stdout, "encoding gate did not report mojibake")
    finally:
        bad_path.unlink(missing_ok=True)

    _run([sys.executable, str(ROOT / "quality_gates" / "check_text_encoding.py"), str(ROOT)])


def test_zip_layout_checker(output_dir: Path) -> None:
    bad_zip = output_dir / "bad_layout.zip"
    with zipfile.ZipFile(bad_zip, "w") as zf:
        zf.writestr("pro_docx_gen/__init__.py", "__version__='bad'")
    proc = subprocess.run(
        [sys.executable, str(ROOT / "quality_gates" / "check_zip_layout.py"), str(bad_zip)],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "zip layout checker allowed duplicate-root layout")
    _assert("forbidden duplicate root" in proc.stdout, "zip layout error not reported")

    polluted_zip = output_dir / "polluted_layout.zip"
    with zipfile.ZipFile(polluted_zip, "w") as zf:
        zf.writestr("pro-docx-gen/pro_docx_gen/__init__.py", "__version__='1.6.6'")
        zf.writestr("pro-docx-gen/pro_docx_gen/SKILL.md", "---\nname: pro-docx-gen\n---\n")
        zf.writestr("pro-docx-gen/quality_gates/run_quality_gate.py", "")
        zf.writestr("pro-docx-gen/quality_gates/check_zip_layout.py", "")
        zf.writestr("pro-docx-gen/smoke_tests/run_smoke_tests.py", "")
        zf.writestr("pro-docx-gen/README.md", "path=C:\\Users\\k1832\\Desktop\\bad.docx")
        zf.writestr("pro-docx-gen/CHANGELOG.md", "# Changelog")
        zf.writestr("pro-docx-gen/smoke_tests/_output/result.docx", b"not a real docx")
        zf.writestr("pro-docx-gen/_work/jobs/leaked/manifest.json", "{}")
    proc = subprocess.run(
        [sys.executable, str(ROOT / "quality_gates" / "check_zip_layout.py"), str(polluted_zip)],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    _assert(proc.returncode != 0, "zip layout checker allowed generated outputs/local paths")
    _assert("generated test output" in proc.stdout, "generated output pollution not reported")
    _assert("runtime work directory" in proc.stdout, "runtime work directory pollution not reported")
    _assert("local machine path leaked" in proc.stdout, "local path leak not reported")


GATE_TESTS = [
    test_imports,
    test_package_only_import,
    test_agent_deliver_rejects_no_render,
    test_portability_backend_guidance,
    test_shared_compat_wrappers,
    test_quality_gate_blocks_warning,
    test_quality_gate_blocks_stale_output_after_generation_failure,
    test_safe_save_reports_locked_destination,
    test_text_encoding_gate,
    test_zip_layout_checker,
]

RENDER_TESTS = [
    test_docx_quality_gate,
    test_agent_entrypoint_workflow,
    test_sampled_four_up_review_bundle,
]

TARGET_TESTS = {
    "gate": GATE_TESTS,
    "render": RENDER_TESTS,
    "chart": [test_chart_assets_do_not_leak, test_chart_readability_helpers, test_preinsertion_figure_gate],
    "table": [
        test_quality_gate_auto_fixes_narrow_table_header,
        test_quality_gate_auto_fixes_splittable_table_rows,
        test_quality_gate_auto_fixes_callout_split_risk,
        test_quality_gate_auto_fixes_prose_table_orphan_risk,
        test_quality_gate_auto_fixes_narrow_table_body_word,
        test_quality_gate_blocks_short_word_and_price_breaks,
    ],
    "save": [test_safe_save_reports_locked_destination],
    "package": [test_text_encoding_gate, test_zip_layout_checker],
}


def _run_test(test_func, output_dir: Path) -> None:
    if inspect.signature(test_func).parameters:
        test_func(output_dir)
    else:
        test_func()


def _run_tests(tests: list, output_dir: Path) -> None:
    seen = set()
    for test_func in tests:
        if test_func in seen:
            continue
        seen.add(test_func)
        _run_test(test_func, output_dir)


def main() -> int:
    parser = argparse.ArgumentParser(description="Run PRO-DOCX v1.6.6 smoke tests.")
    parser.add_argument("--output-dir", type=Path, default=ROOT / "smoke_tests" / "_output")
    parser.add_argument(
        "--profile",
        choices=("gate-render", "full"),
        default="gate-render",
        help="Default runs only gate and render layers. Use full only before release or after broad changes.",
    )
    parser.add_argument(
        "--target",
        action="append",
        choices=sorted(TARGET_TESTS),
        help="Run a focused smoke target after a related feature issue is found. May be repeated.",
    )
    args = parser.parse_args()

    args.output_dir.mkdir(parents=True, exist_ok=True)
    if args.target:
        selected = []
        for target in args.target:
            selected.extend(TARGET_TESTS[target])
        _run_tests(selected, args.output_dir)
        print(f"PASS targeted smoke tests targets={','.join(args.target)}")
        return 0

    if args.profile == "full":
        selected = []
        for tests in TARGET_TESTS.values():
            selected.extend(tests)
        _run_tests(selected, args.output_dir)
        print("PASS full smoke tests")
        return 0

    _run_tests(GATE_TESTS + RENDER_TESTS, args.output_dir)
    print("PASS gate-render smoke tests")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
