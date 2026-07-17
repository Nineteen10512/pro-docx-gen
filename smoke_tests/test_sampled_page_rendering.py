import importlib.util
import tempfile
import unittest
from pathlib import Path
from unittest import mock

import agent_docx
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
GATE_PATH = ROOT / "quality_gates" / "run_quality_gate.py"
SPEC = importlib.util.spec_from_file_location("pro_docx_quality_gate", GATE_PATH)
gate = importlib.util.module_from_spec(SPEC)
assert SPEC and SPEC.loader
SPEC.loader.exec_module(gate)


class SamplePageNumberTests(unittest.TestCase):
    def test_select_sample_pages_uses_one_based_source_pages(self):
        self.assertEqual(gate.select_sample_page_numbers(1), [1])
        self.assertEqual(gate.select_sample_page_numbers(2), [1])
        self.assertEqual(gate.select_sample_page_numbers(3), [1])
        self.assertEqual(gate.select_sample_page_numbers(4), [1, 4])
        self.assertEqual(gate.select_sample_page_numbers(7), [1, 4, 7])
        self.assertEqual(gate.select_sample_page_numbers(60), list(range(1, 61, 3)))

    def test_preview_mapping_preserves_source_page_numbers(self):
        preview = {
            "page_pngs": ["p1.png", "p4.png", "p7.png"],
            "page_numbers": [1, 4, 7],
        }

        self.assertEqual(
            agent_docx._page_previews_from_result(preview),
            [(1, "p1.png"), (4, "p4.png"), (7, "p7.png")],
        )

    def test_preview_mapping_rejects_mismatched_page_metadata(self):
        preview = {
            "page_pngs": ["p1.png", "p4.png"],
            "page_numbers": [1],
        }

        with self.assertRaisesRegex(ValueError, "page_numbers"):
            agent_docx._page_previews_from_result(preview)


class SelectiveComExportTests(unittest.TestCase):
    def test_selective_script_computes_pages_and_exports_only_page_ranges(self):
        script = gate._selective_export_script("KWPS.Application")

        self.assertIn("ComputeStatistics(2)", script)
        self.assertIn("$page += $GroupSize", script)
        self.assertIn("$page, $page", script)
        self.assertIn("ExportAsFixedFormat", script)
        self.assertNotIn("ExportAsFixedFormat($PdfPath, 17)", script)
        self.assertIn("try { $app.Quit() } catch {}", script)

    def test_selective_com_receives_absolute_document_and_output_paths(self):
        completed = mock.Mock(returncode=1, stdout="", stderr="failed")
        with mock.patch.object(gate.subprocess, "run", return_value=completed) as run:
            gate._render_sampled_with_com(
                Path("relative-report.docx"),
                Path("relative-quality"),
                3,
                "KWPS.Application",
                "wps_com_selective",
            )

        command = run.call_args.args[0]
        self.assertTrue(Path(command[-3]).is_absolute())
        self.assertTrue(Path(command[-2]).is_absolute())

    def test_strict_render_tries_only_selective_com_backends(self):
        expected = {
            "status": "passed",
            "engine": "word_com_selective",
            "page_count": 4,
            "sample_group_size": 3,
            "sampled_page_numbers": [1, 4],
            "sampled_pdfs": ["p1.pdf", "p4.pdf"],
        }
        with mock.patch.object(
            gate,
            "_render_sampled_with_wps_com",
            return_value={"status": "failed", "engine": "wps_com_selective"},
        ) as wps, mock.patch.object(
            gate,
            "_render_sampled_with_word_com",
            return_value=expected,
        ) as word, mock.patch.object(
            gate,
            "_render_with_office_cli",
        ) as office_full, mock.patch.object(
            gate,
            "_render_with_wps_com",
        ) as wps_full, mock.patch.object(
            gate,
            "_render_with_word_com",
        ) as word_full:
            result = gate.try_render_sampled(Path("report.docx"), Path("quality"), 3)

        self.assertEqual(result["sampled_page_numbers"], [1, 4])
        wps.assert_called_once()
        word.assert_called_once()
        office_full.assert_not_called()
        wps_full.assert_not_called()
        word_full.assert_not_called()

    def test_strict_render_fails_without_selective_backend(self):
        failure = {"status": "failed", "reason": "unavailable"}
        with mock.patch.object(gate, "_render_sampled_with_wps_com", return_value=failure), mock.patch.object(
            gate, "_render_sampled_with_word_com", return_value=failure
        ):
            result = gate.try_render_sampled(Path("report.docx"), Path("quality"), 3)

        self.assertEqual(result["status"], "failed")
        self.assertEqual(result["reason"], "no selective page render backend succeeded")

    def test_targeted_review_script_exports_only_explicit_pages(self):
        script = gate._targeted_export_script("KWPS.Application")

        self.assertIn("$PagesCsv.Split(',')", script)
        self.assertIn("foreach ($page in $pages)", script)
        self.assertIn("$page, $page", script)
        self.assertNotIn("$page += $GroupSize", script)

    def test_targeted_review_tries_only_selective_com_backends(self):
        expected = {
            "status": "passed",
            "engine": "word_com_targeted",
            "page_count": 40,
            "sampled_page_numbers": [24, 26],
            "sampled_pdfs": ["p24.pdf", "p26.pdf"],
        }
        with mock.patch.object(
            gate,
            "_render_targeted_with_wps_com",
            return_value={"status": "failed", "engine": "wps_com_targeted"},
        ) as wps, mock.patch.object(
            gate,
            "_render_targeted_with_word_com",
            return_value=expected,
        ) as word, mock.patch.object(gate, "try_render_sampled") as regular_sampling:
            result = gate.try_render_specific_pages(Path("report.docx"), Path("quality"), [24, 26])

        self.assertEqual(result["sampled_page_numbers"], [24, 26])
        wps.assert_called_once()
        word.assert_called_once()
        regular_sampling.assert_not_called()


class SampledQualityGateTests(unittest.TestCase):
    def test_sampled_preview_uses_short_pdf_aliases_for_windows_path_safety(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            original = root / ("very-long-report-name-" * 6 + "page-0001.pdf")
            original.write_bytes(b"%PDF-1.4\n")
            generated_png = root / "p0001-1.png"
            generated_png.write_bytes(b"png")

            def fake_preview(pdf_path, out_dir, max_pages=200):
                self.assertEqual(pdf_path.name, "p0001.pdf")
                self.assertTrue(pdf_path.exists())
                return {
                    "status": "passed",
                    "page_pngs": [str(generated_png)],
                    "pngs": [str(generated_png)],
                }

            with mock.patch.object(gate, "_render_pdf_previews", side_effect=fake_preview):
                result = gate._render_sampled_pdf_previews(
                    {
                        "sampled_page_numbers": [1],
                        "sampled_pdfs": [str(original)],
                    },
                    root / "quality",
                )

        self.assertEqual(result["status"], "passed")
        self.assertEqual(result["page_numbers"], [1])

    def test_inspect_file_uses_sampled_render_and_preserves_source_pages(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            root = Path(temp_dir)
            docx_path = root / "report.docx"
            document = Document()
            document.add_heading("Sample report", level=1)
            document.add_paragraph("Evidence-backed content for sampled rendering.")
            document.save(docx_path)
            render_result = {
                "status": "passed",
                "engine": "wps_com_selective",
                "page_count": 7,
                "sample_group_size": 3,
                "sampled_page_numbers": [1, 4, 7],
                "sampled_pdfs": [str(root / "p1.pdf"), str(root / "p4.pdf"), str(root / "p7.pdf")],
            }
            preview_result = {
                "status": "passed",
                "page_numbers": [1, 4, 7],
                "page_pngs": [str(root / "p1.png"), str(root / "p4.png"), str(root / "p7.png")],
                "pngs": [str(root / "p1.png"), str(root / "p4.png"), str(root / "p7.png")],
                "page_count_previewed": 3,
            }
            with mock.patch.object(gate, "try_render_sampled", return_value=render_result) as sampled, mock.patch.object(
                gate,
                "_render_sampled_pdf_previews",
                return_value=preview_result,
                create=True,
            ) as previews, mock.patch.object(gate, "try_render") as full_render:
                report = gate.inspect_file(
                    docx_path,
                    render=True,
                    output_dir=root / "quality",
                    sample_group_size=3,
                )

        self.assertEqual(report["render"]["sampled_page_numbers"], [1, 4, 7])
        self.assertNotIn("pdf", report["render"])
        self.assertEqual(report["preview_png"]["page_numbers"], [1, 4, 7])
        sampled.assert_called_once()
        previews.assert_called_once()
        full_render.assert_not_called()


class AgentSampledDeliveryTests(unittest.TestCase):
    def test_quality_preview_pages_are_used_without_second_sampling(self):
        data = {
            "files": [
                {
                    "metrics": {},
                    "render": {"page_count": 7},
                    "preview_png": {
                        "page_numbers": [1, 4, 7],
                        "page_pngs": ["p1.png", "p4.png", "p7.png"],
                    },
                }
            ]
        }

        pages = agent_docx._sampled_page_previews_from_quality_data(data)

        self.assertEqual(pages, [(1, "p1.png"), (4, "p4.png"), (7, "p7.png")])

    def test_submission_manifest_declares_sampled_only_scope(self):
        manifest = agent_docx._submission_manifest(
            target=Path("report.docx"),
            quality_report=Path("quality.json"),
            document_page_count=7,
            sampled_pages=[(1, "p1.png"), (4, "p4.png"), (7, "p7.png")],
            review_bundle={
                "pdf": "review.pdf",
                "sheet_pngs": ["review-1.png"],
                "pages_per_sheet": 4,
                "sampled_page_numbers": [1, 4, 7],
            },
        )

        self.assertEqual(manifest["render_scope"], "sampled_pages_only")
        self.assertEqual(manifest["document_page_count"], 7)
        self.assertEqual(manifest["visual_review_sampling"]["sampled_page_numbers"], [1, 4, 7])
        self.assertNotIn("full_document_pdf", manifest)


if __name__ == "__main__":
    unittest.main()
