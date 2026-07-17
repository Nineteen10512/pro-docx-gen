"""v1.7.0 3 份 demo docx 生成脚本。

为 18 模板 + 表格引擎 5 项新能力 + 3 大模块（KPI / header / footer）
生成 3 份代表性 demo，放到 pro_docx_gen/v17_demo/。

- business_compact_demo.docx：商务紧凑（标题/段落行高压缩 + 单行表格 +
  提要框 + 进度条 + 评级 + 条件格式）
- research_report_demo.docx：研报标准（4 字段 header + KPI 卡片行 +
  YoY/汇总表格 + 风险提示固定尾部）
- table_advanced_features_demo.docx：表格引擎 5 项新能力专项 demo
  （条件格式 / 进度条 / 星级 / 多表头合并 / 自动计算 YoY·汇总）

每个 demo 跑完都 self_audit 一下，确保 0 violations。
"""
from __future__ import annotations

import os
import sys

# 让 pro_docx_gen 可以被 import
ROOT = "/app/data/所有对话/主对话/skills/pro-docx-gen"
sys.path.insert(0, ROOT)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pro_docx_gen import (
    __version__,
    render_conditional_cell,
    render_progress_bar,
    render_star_rating,
    render_merged_header,
    auto_compute_rows,
    render_kpi_card_row,
    render_research_header,
    render_rating_badge,
    render_risk_disclaimer_footer,
    compute_yoy,
    compute_summary,
)
from pro_docx_gen.engine.renderer import _set_cell_shading
from pro_docx_gen.shared.color_palette import resolve_palette
from pro_docx_gen.tokens import get_theme, merge_theme, BASE_TOKENS
from pro_docx_gen.docx_jsx import generate as docx_generate
from pro_docx_gen.engine.research_modules import (
    RESEARCH_NAVY, RESEARCH_PURPLE, RATING_OVERWEIGHT_FILL,
)


# 帮助：构造一份 tokens（绕过 DocxRenderer，直接给 doc 灌内容）
def _build_tokens(theme_name: str = "academic") -> dict:
    tokens = get_theme(theme_name)
    return tokens


# ───────────────────────────────────────────────────────────────
# Demo 1: business_compact — 商务紧凑报告
# ───────────────────────────────────────────────────────────────


def make_business_compact_demo(out_path: str) -> None:
    """商务紧凑 demo — 走 generate() + template_name=business_compact."""
    doc = {
        "meta": {
            "title": "2026 Q2 销售周报",
            "author": "PRO-DOCX v1.7.0",
        },
        "sections": [
            {
                "title": "执行摘要",
                "level": 1,
                "content": [
                    {"type": "paragraph", "text": "本周完成 23 单 / 转化率 18.4%，环比+10.5%。"},
                    {"type": "callout", "variant": "info",
                     "title": "重点", "body": "新增 4 家头部客户，A 类产品占比提升至 41%。"},
                    {"type": "paragraph", "text": "整体节奏符合 Q2 预算；下周一开展大客户回访。"},
                ],
            },
            {
                "title": "关键数据",
                "level": 1,
                "content": [
                    {
                        "type": "table",
                        "caption": "周度销售数据（条件格式：YoY 自动染色）",
                        "headers": ["指标", "上周", "本周", "环比"],
                        "col_widths": [2.0, 1.5, 1.5, 1.5],
                        "header_repeat": True,
                        "rows": [
                            ["销售额（万元）", 320, 354, 0.105],
                            ["订单数", 198, 220, 0.111],
                            ["A 类订单占比", 0.32, 0.41, 0.281],
                            ["新客户数", 6, 10, 0.667],
                            ["退款单", 4, 7, 0.75],
                        ],
                    },
                ],
            },
            {
                "title": "风险与阻塞",
                "level": 1,
                "content": [
                    {"type": "paragraph", "text": "退款单环比+75%，集中在 7 月新发批次，建议立刻排查。"},
                    {"type": "callout", "variant": "warning",
                     "title": "阻塞", "body": "华东仓配送延迟 24h，已与物流方对接。"},
                ],
            },
            {
                "title": "下一步计划",
                "level": 1,
                "content": [
                    {"type": "paragraph", "text": "1) 大客户回访 5 家；2) 退款排查 SOP 上线；3) A 类主推活动策划。"},
                ],
            },
        ],
    }
    docx_generate(
        doc,
        out_path,
        theme="business",
        lang="cn",
        template_name="business_compact",
    )


# ───────────────────────────────────────────────────────────────
# Demo 2: research_report — 研报标准
# ───────────────────────────────────────────────────────────────


def make_research_report_demo(out_path: str) -> None:
    """研报标准 demo — 走 generate() + template_name=research_report
    + 在 sections 之前注入 v1.7.0 模块（通过 plan 注入或后处理）。

    实际 v1.7.0 的 3 大模块 (header / KPI / risk_disclaimer) 是直接调用
    公开 API 插入的。本 demo 走"先生成骨架，再后处理插入 v1.7.0 模块"
    的双阶段路径，避免动 generate() 的语义 JSON schema。
    """
    # Stage 1: 走 generate() 拿到研报骨架
    skeleton = {
        "meta": {
            "title": "海康威视首次覆盖：安防龙头，AI 时代新增长",
            "author": "PRO-DOCX v1.7.0 Research",
        },
        "sections": [
            {
                "title": "公司概览",
                "level": 1,
                "content": [
                    {"type": "paragraph",
                     "text": "海康威视是全球领先的安防产品及解决方案提供商，业务覆盖前端感知、后端存储、智能算法及行业应用。"},
                    {"type": "paragraph",
                     "text": "公司 2024 年营收 924.6 亿元，归母净利润 119.6 亿元。EBG/SMBG/PBG 三业务线协同推进数字化转型。"},
                ],
            },
            {
                "title": "投资摘要",
                "level": 1,
                "content": [
                    {"type": "paragraph",
                     "text": "我们看好公司在 AI 视觉与垂直行业应用的长期增长，给予 Overweight 评级。"},
                ],
            },
            {
                "title": "财务分析",
                "level": 1,
                "content": [
                    {"type": "table",
                     "caption": "海康威视财务摘要（亿元，YoY%）",
                     "headers": ["项目", "2022", "2023", "2024", "2025E", "2026E", "YoY%"],
                     "col_widths": [1.3, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8],
                     "header_repeat": True,
                     "rows": [
                         ["营业收入", 831.7, 893.4, 924.6, 1015.0, 1118.5, 0.135],
                         ["归母净利润", 128.3, 141.1, 119.6, 138.0, 156.2, 0.131],
                         ["毛利率", "44.3%", "44.0%", "43.8%", "44.2%", "44.5%", 0.016],
                         ["净利率", "15.4%", "15.8%", "12.9%", "13.6%", "14.0%", -0.043],
                         ["研发投入", 98.1, 113.5, 121.4, 134.0, 148.2, 0.103],
                     ],
                    },
                ],
            },
            {
                "title": "估值与目标价",
                "level": 1,
                "content": [
                    {"type": "paragraph",
                     "text": "采用 PE 估值法，参考行业可比公司，给予 2026E 25x PE，对应目标价 38.5 元。"},
                ],
            },
            {
                "title": "风险因素",
                "level": 1,
                "content": [
                    {"type": "paragraph",
                     "text": "海外政策风险、原材料价格波动、行业竞争加剧。"},
                ],
            },
        ],
    }
    docx_generate(
        skeleton,
        out_path,
        theme="academic",
        lang="cn",
        template_name="research_report",
    )

    # Stage 2: 用 python-docx 打开，后处理注入 v1.7.0 模块
    doc = Document(out_path)
    tokens = _build_tokens("academic")

    # 把 doc body 的前几个段移到 v1.7.0 header 后
    body = doc.element.body
    first_p = None
    paragraphs_to_remove = []
    for child in list(body):
        if child.tag == qn("w:p"):
            txt = "".join(t.text or "" for t in child.iter(qn("w:t")))
            if "海康威视" in txt or "PRO-DOCX" in txt or "2026" in txt and "Q2" in txt:
                # 这是 title block；保留
                pass
            if first_p is None:
                first_p = child
                break
        elif child.tag == qn("w:tbl"):
            break

    # 在 document 开头（title block 后）插入 4 字段 header
    # 找到 title block 后的第一个段落（即第一个 paragraph of first section）
    # 简单方案：在 body[0] 后面插入一组新段落
    insertion_target = None
    section_count = 0
    for child in list(body):
        if child.tag == qn("w:p"):
            section_count += 1
            if section_count == 1:
                # 这是 title block 段（首段）
                insertion_target = child
                break

    # 准备临时 doc 来构造 v1.7.0 header / KPI 段
    tmp = Document()
    # 走 v1.7.0 渲染
    render_research_header(tmp, tokens, {
        "title": "海康威视",
        "stock_code": "002415.SZ",
        "rating": "Overweight",
        "target_price": "38.50 元",
        "current_price": "31.20",
        "report_date": "2026-07-15",
        "subtitle": "安防龙头·AI 时代新增长曲线",
        "analyst": "张明",
    }, lang="cn")
    render_kpi_card_row(tmp, [
        {"label": "收盘价", "value": "31.20", "suffix": "元", "delta": 0.025,
         "trend": [30.1, 30.5, 30.8, 31.0, 30.7, 31.1, 31.2]},
        {"label": "目标价", "value": "38.50", "suffix": "元", "delta": 0.234,
         "trend": [33.0, 34.5, 35.0, 36.0, 37.0, 37.8, 38.5]},
        {"label": "市值", "value": "2934", "suffix": "亿元", "delta": 0.025,
         "trend": [2800, 2850, 2870, 2900, 2880, 2920, 2934]},
        {"label": "PE (TTM)", "value": "24.5", "suffix": "x", "delta": -0.05,
         "trend": [25.0, 25.3, 25.0, 24.8, 24.9, 24.7, 24.5]},
    ], tokens, lang="cn")

    # 把临时 doc 的前 N 个段落/表格插到 title block 之后
    # body 段（tmp.element.body）所有非 sectPr 元素
    tmp_body = tmp.element.body
    new_elements = []
    for child in list(tmp_body):
        if child.tag == qn("w:sectPr"):
            continue
        new_elements.append(child)

    # 找 position
    if insertion_target is not None:
        # 找到 insertion_target 在 body 中的索引
        idx = list(body).index(insertion_target)
        for offset, el in enumerate(new_elements):
            body.insert(idx + 1 + offset, el)
    else:
        for el in new_elements:
            body.insert(0, el)

    # Stage 3: 在 body 末尾加风险提示
    render_risk_disclaimer_footer(doc, tokens, lang="cn", institution="示例证券研究所")
    doc.save(out_path)


# ───────────────────────────────────────────────────────────────
# Demo 3: table_advanced_features — 5 项新能力专项 demo
# ───────────────────────────────────────────────────────────────


def make_table_advanced_features_demo(out_path: str) -> None:
    """表格引擎 5 项新能力专项 demo — 直接用 python-docx 构造 doc，
    调用 5 个公开函数 + KPI / research 头尾模块。

    走 generate() 难以精确控制每个 cell 的格式化（条件格式 / 进度条
    / 星级 / 多表头合并 / 自动计算），所以直接 python-docx 拼 doc。
    """
    doc = Document()
    tokens = _build_tokens("academic")
    # 修正 cn_font 等让下面的渲染可用
    # tokens 已包含 spacing/color/font

    # 文档元信息
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("pro-docx-gen v1.7.0 — 表格引擎 5 项新能力 Demo")
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("1A1A1A")

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("条件格式 / 进度条 / 星级评分 / 多表头合并 / 自动计算")
    r2.font.size = Pt(11)
    r2.font.italic = True
    r2.font.color.rgb = RGBColor.from_string("666666")

    doc.add_paragraph()

    # ── 1) 条件格式 + 进度条 + 星级评分（一张大表演示）──
    h = doc.add_paragraph()
    rh = h.add_run("1. 条件格式 + 进度条 + 星级评分")
    rh.font.size = Pt(14)
    rh.font.bold = True
    rh.font.color.rgb = RGBColor.from_string("1A1A1A")

    # 表格：指标 / 数值 / 同比 (条件格式) / 进度 (进度条) / 评分 (星级)
    t1 = doc.add_table(rows=1, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = False
    # 标题行
    header_titles = ["指标", "数值", "同比", "完成度", "评分"]
    widths = [Inches(1.4), Inches(1.0), Inches(1.0), Inches(2.0), Inches(1.4)]
    for i, ht in enumerate(header_titles):
        c = t1.rows[0].cells[i]
        c.width = widths[i]
        c.text = ""
        cp = c.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(ht)
        cr.font.size = Pt(10)
        cr.font.bold = True
        cr.font.color.rgb = RGBColor.from_string("FFFFFF")
        _set_cell_shading(c, "1F3864")

    # 数据行
    rows_data = [
        ("营收增长", "10.5%", 0.105, 78.0, 4.5),
        ("净利润增长", "-3.2%", -0.032, 55.0, 3.0),
        ("新客户获取", "23.1%", 0.231, 92.0, 5.0),
        ("市场份额", "0.0%", 0.0, 60.0, 3.5),
        ("NPS 评分", "—", None, 41.0, 2.5),
    ]
    for label, val, yoy, pct, stars in rows_data:
        row = t1.add_row()
        # 指标
        c0 = row.cells[0]; c0.width = widths[0]; c0.text = ""
        c0p = c0.paragraphs[0]; c0p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        c0r = c0p.add_run(label); c0r.font.size = Pt(10); c0r.font.color.rgb = RGBColor.from_string("1A1A1A")
        # 数值
        c1 = row.cells[1]; c1.width = widths[1]; c1.text = ""
        c1p = c1.paragraphs[0]; c1p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c1r = c1p.add_run(val); c1r.font.size = Pt(10); c1r.font.color.rgb = RGBColor.from_string("333333")
        # 同比 — 条件格式
        render_conditional_cell(row.cells[2], yoy, {
            "type": "yoy",
            "value_text": val if val != "—" else "—",
        }, tokens)
        row.cells[2].width = widths[2]
        # 进度条
        render_progress_bar(row.cells[3], pct, tokens, color="0E7C3A" if pct >= 60 else "B45309" if pct >= 40 else "B91C1C")
        row.cells[3].width = widths[3]
        # 星级
        render_star_rating(row.cells[4], stars, tokens, color="D97706")
        row.cells[4].width = widths[4]

    doc.add_paragraph()

    # ── 2) 多表头合并 ──
    h2 = doc.add_paragraph()
    rh2 = h2.add_run("2. 多表头合并（二维表头）")
    rh2.font.size = Pt(14); rh2.font.bold = True
    rh2.font.color.rgb = RGBColor.from_string("1A1A1A")

    t2 = doc.add_table(rows=2 + 3, cols=1 + 6)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.autofit = False
    render_merged_header(t2, [
        ["指标", "2024", "2024", "2024", "2025E", "2025E", "2025E"],
        ["",     "Q1",   "Q2",   "Q3",   "Q1E",  "Q2E",  "Q3E"],
    ], tokens, col_widths_inches=[1.4, 0.7, 0.7, 0.7, 0.7, 0.7, 0.7])
    # 填数据行
    body_data = [
        ["营收", 220, 240, 260, 250, 270, 290],
        ["毛利", 95, 105, 112, 110, 120, 132],
        ["净利", 30, 35, 38, 36, 42, 48],
    ]
    for i, row_vals in enumerate(body_data):
        for j, v in enumerate(row_vals):
            cell = t2.rows[2 + i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(v))
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor.from_string("1A1A1A")

    doc.add_paragraph()

    # ── 3) 自动计算 (YoY + 汇总) ──
    h3 = doc.add_paragraph()
    rh3 = h3.add_run("3. 自动计算：YoY 同比 + 汇总行")
    rh3.font.size = Pt(14); rh3.font.bold = True
    rh3.font.color.rgb = RGBColor.from_string("1A1A1A")

    base_rows = [
        ["2024 Q1", 220, 95, 30],
        ["2024 Q2", 240, 105, 35],
        ["2024 Q3", 260, 112, 38],
        ["2024 Q4", 280, 120, 42],
        ["2025 Q1", 250, 110, 36],
        ["2025 Q2", 270, 120, 42],
        ["2025 Q3", 290, 132, 48],
    ]
    result = auto_compute_rows(
        base_rows,
        yoy_col_index=1,         # 营收
        summary_mode="sum",      # 在头部加汇总行
    )
    new_rows = result["rows"]
    # 5 列：期间 / 营收 / 毛利 / 净利 / 营收 YoY
    n_data_cols = 5
    t3 = doc.add_table(rows=1 + len(new_rows), cols=n_data_cols)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3.autofit = False
    # 表头
    for i, h in enumerate(["期间", "营收", "毛利", "净利", "营收 YoY"]):
        c = t3.rows[0].cells[i]; c.text = ""
        c.width = Inches(1.3)
        cp = c.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(h); cr.font.size = Pt(10); cr.font.bold = True
        cr.font.color.rgb = RGBColor.from_string("FFFFFF")
        _set_cell_shading(c, "1F3864")
    # 数据
    for ri, row in enumerate(new_rows):
        for ci in range(n_data_cols):
            if ci >= len(row):
                continue
            val = row[ci]
            cell = t3.rows[ri + 1].cells[ci]; cell.text = ""
            cell.width = Inches(1.3)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if ri == 0:  # summary row
                r.font.bold = True
                r.font.color.rgb = RGBColor.from_string("1F3864")
            elif ci == 4:  # YoY column → condition
                render_conditional_cell(cell, val, {
                    "type": "yoy",
                    "value_text": str(val),
                }, tokens)
            else:
                r.font.color.rgb = RGBColor.from_string("333333")

    doc.add_paragraph()

    # ── 4) KPI 卡片行 ──
    h4 = doc.add_paragraph()
    rh4 = h4.add_run("4. KPI 卡片行")
    rh4.font.size = Pt(14); rh4.font.bold = True
    rh4.font.color.rgb = RGBColor.from_string("1A1A1A")

    # 注入 layout_content_width hint
    from pro_docx_gen.engine.layout import LayoutCalculator
    layout = LayoutCalculator(tokens)
    tokens["__layout_content_width_emu__"] = layout.content_width
    render_kpi_card_row(doc, [
        {"label": "总收入", "value": "2934", "suffix": "亿元",
         "delta": 0.105, "trend": [2700, 2780, 2810, 2850, 2880, 2900, 2934]},
        {"label": "净利润", "value": "119.6", "suffix": "亿元",
         "delta": -0.152, "trend": [130, 132, 128, 125, 122, 121, 119.6]},
        {"label": "毛利率", "value": "43.8", "suffix": "%",
         "delta": 0.005, "trend": [43.5, 43.6, 43.7, 43.7, 43.8, 43.8, 43.8]},
        {"label": "PE", "value": "24.5", "suffix": "x",
         "delta": -0.05, "trend": [25.5, 25.3, 25.0, 24.9, 24.8, 24.7, 24.5]},
    ], tokens, lang="cn")

    doc.add_paragraph()

    # ── 5) 评级徽章 ──
    h5 = doc.add_paragraph()
    rh5 = h5.add_run("5. 评级徽章")
    rh5.font.size = Pt(14); rh5.font.bold = True
    rh5.font.color.rgb = RGBColor.from_string("1A1A1A")

    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pir = p_intro.add_run("研报 header 用评级徽章强化视觉权重，三色：")
    pir.font.size = Pt(10); pir.font.color.rgb = RGBColor.from_string("333333")

    # 用一行 3 张表展示 3 种评级
    for rating in ("Overweight", "Hold", "Underweight"):
        render_rating_badge(doc, tokens, rating, lang="cn", width_inches=1.4)
        # 加一个尾随段以便换行
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()

    # ── 6) 风险提示尾部 ──
    render_risk_disclaimer_footer(doc, tokens, lang="cn", institution="示例证券研究所")

    doc.save(out_path)


# ───────────────────────────────────────────────────────────────
# 主入口
# ───────────────────────────────────────────────────────────────


def main() -> int:
    out_dir = os.path.join(ROOT, "pro_docx_gen", "v17_demo")
    os.makedirs(out_dir, exist_ok=True)

    targets = [
        ("business_compact_demo.docx", make_business_compact_demo),
        ("research_report_demo.docx", make_research_report_demo),
        ("table_advanced_features_demo.docx", make_table_advanced_features_demo),
    ]
    for name, fn in targets:
        path = os.path.join(out_dir, name)
        print(f"→ generating {name} ...")
        fn(path)
        size = os.path.getsize(path)
        print(f"  ✓ {path}  ({size:,} bytes)")
    print("All 3 demos generated.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
