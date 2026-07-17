"""Renderer — 底层 OOXML 渲染器。

唯一直接接触 python-docx API 的模块。
根据 design tokens 配置样式，按扁平节点列表依次渲染。

v1.2 新增能力：
- 修订追踪（revision 节点 → w:ins / w:del）
- 批注（comment 字段 / comment 节点 → comments.xml 注入）
- 脚注（footnote/endnote 节点 → footnotes.xml 注入）
- 水印（watermark 节点/文档级 watermark → VML 文字水印）
- 页面边框（page_border → w:pgBorders）
- 页面设置（纸张大小/方向/独立边距/装订线/首页不同/奇偶页不同）
- 页眉页脚增强（首页不同、奇偶页不同、"第 X 页 共 Y 页"、页眉图片）
- 公式（equation 节点 → LaTeX 纯文本占位，v1.2 简化实现）
- 签名区（signature_block / signature_line）
- 文档元数据（core properties + custom properties）
- 表格表头跨页重复

v1.6.6 强化：
- **强制 heading 颜色为 #1A1A1A**：所有 Heading 1-6 节点 + Title 节点
  必须使用 ``color_palette.heading_rgb()``，不再读 theme_overrides["color"]["heading"]，
  杜绝模板/theme 把 heading 写成黑金/浅色导致白底不可见。
- **强制 text 颜色为 #333333**：正文 run 默认色用 ``color_palette.text_rgb()``。
- **chart 页内适配**（根因修复图表空白页）：
  在插入 chart 前计算 ``available_height_inches``（当前页剩余空间），
  若原图高度超出剩余空间 → 按比例自动缩图；若剩余空间 < 1.5 英寸
  → 触发 page break 后再插入图。
- **chart 紧凑模式（compact_mode: true）**：两个 chart 节点相邻（中间无
  heading/paragraph）时，取消两图之间多余的段间距、caption 用 keep_with_next
  绑定到下一图，避免「图1 单独占一整页 + 图2 单独占一整页」型空白页。
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT, CONTENT_TYPE as CT
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from typing import Optional
import os
import datetime
import tempfile
import shutil
import uuid

from .layout import LayoutCalculator
from .chart_renderer import render_chart_to_png
from .citation import format_reference as _cite_format
from ..figure_preflight import (
    FigureAssetGateError,
    assert_figure_asset_ready,
    compute_available_height_inches,
)

# v1.6.6: import the canonical heading/text/muted RGBColor factory
# functions so that the heading/text color enforcement is centrally owned
# in shared/color_palette.py. Renderer code MUST go through these helpers
# instead of reading theme_overrides["color"]["heading"] directly.
from ..shared.color_palette import (
    heading_rgb as _heading_rgb,
    text_rgb as _text_rgb,
    muted_rgb as _muted_rgb,
    HEADING_HEX as _HEADING_HEX,
    TEXT_HEX as _TEXT_HEX,
    MUTED_HEX as _MUTED_HEX,
)


# v1.6.6: minimum available page space (in inches) for chart insertion.
# If the current page has less remaining space, force a page break.
CHART_MIN_AVAILABLE_INCHES = 1.5
# v1.6.6: assumed top/bottom safety margin (line spacing + footer area)
# used to estimate available height when the python-docx layout API is not
# sufficient (it does not expose live page-remaining space).
CHART_PAGE_BOTTOM_MARGIN_INCHES = 0.6


class _WritableTempDir:
    """Temporary directory that remains writable in restricted Windows sandboxes."""

    def __init__(self, prefix: str, base_dir: str):
        self._base_dir = os.path.join(base_dir, ".pro_docx_work")
        os.makedirs(self._base_dir, exist_ok=True)
        self.name = os.path.join(self._base_dir, f"{prefix}{uuid.uuid4().hex}")
        os.makedirs(self.name, exist_ok=False)

    def cleanup(self):
        shutil.rmtree(self.name, ignore_errors=True)
        try:
            os.rmdir(self._base_dir)
        except OSError:
            pass


# ─── XML 辅助 ────────────────────────────────────────────────────

def _set_cell_shading(cell, fill_hex: str):
    """设置表格单元格背景色（fill_hex 无 # 前缀，如 '1F3864'）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def _set_cell_border(cell, **kwargs):
    """设置单元格边框。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, attrs in kwargs.items():
        elem = tc_borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tc_borders.append(elem)
        for k, v in attrs.items():
            elem.set(qn(f"w:{k}"), str(v))


def _set_paragraph_border_left(paragraph, color_hex: str, size: int = 12):
    """为段落设置左边框（用于引用块）。"""
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color_hex)
    p_bdr.append(left)


def _set_paragraph_shading(paragraph, fill_hex: str):
    """为段落设置底纹。"""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    p_pr.append(shd)


def _add_field(paragraph, field_code: str, placeholder: str = ""):
    """在段落中插入 Word 域代码（用于 TOC、页码等）。

    Args:
        paragraph: python-docx Paragraph 对象。
        field_code: 域指令文本，例如 ``TOC \\o "1-3" \\h \\z \\u``。
        placeholder: separate 与 end 之间的占位文字（未更新域前显示）。
    """
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    run._r.append(instr)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_sep)

    if placeholder:
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = placeholder
        run._r.append(t)
    else:
        t = OxmlElement("w:t")
        t.text = ""
        run._r.append(t)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _rgb_to_hex(color) -> str:
    if color is None:
        return "333333"
    if hasattr(color, "__getitem__"):
        try:
            return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
        except Exception:
            pass
    return "333333"


def _set_run_font(run, font_name: str, cn_font_name: str = None):
    """设置 run 字体，同时设置 west 和 eastAsia。"""
    run.font.name = font_name
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    if cn_font_name:
        r_fonts.set(qn("w:eastAsia"), cn_font_name)
    else:
        r_fonts.set(qn("w:eastAsia"), font_name)


def _set_table_header_repeat(row):
    """设置表格表头行跨页重复（w:tblHeader）。"""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)


def _set_table_row_no_split(row):
    """Prevent Word/WPS from splitting a table row across pages."""
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def _safe_save_document(doc: Document, output_path: str) -> str:
    """Save through a temp file, then replace the target.

    This avoids half-written DOCX files. If Word/WPS locks the destination,
    the caller gets a clear error instead of silent corruption.
    """
    output_path = os.path.abspath(output_path)
    out_dir = os.path.dirname(output_path) or "."
    os.makedirs(out_dir, exist_ok=True)
    fd, tmp_path = tempfile.mkstemp(
        prefix=f".{os.path.basename(output_path)}.",
        suffix=".tmp.docx",
        dir=out_dir,
    )
    os.close(fd)
    try:
        doc.save(tmp_path)
        try:
            os.replace(tmp_path, output_path)
        except PermissionError as exc:
            raise RuntimeError(
                "Target DOCX is locked by Word/WPS or another process. "
                f"Close it and retry, or use a different output_path. Safe temp copy: {tmp_path}"
            ) from exc
    except Exception:
        if os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except OSError:
                pass
        raise
    return output_path


# ─── 渲染器主类 ──────────────────────────────────────────────────

class DocxRenderer:
    """基于 design tokens 将扁平节点列表渲染为 docx。"""

    def __init__(self, tokens: dict, lang: str = "en", output_dir: str = ".", variant: Optional[str] = None):
        self.tokens = tokens
        self.lang = lang
        self.variant = variant  # v1.6.0: 当前风格变体名
        # v1.6.6 — root-cause fix: enforce canonical heading/text/muted colors
        # at the tokens level. Templates and theme_overrides["color"] used to
        # leak #D4AF37 / #E0E7FF / #64FFDA into heading/text, which became
        # unreadable on a white background. We now lock heading → #1A1A1A,
        # text → #333333, muted → #666666, and let accent/primary/secondary
        # keep their theme-defined personality.
        try:
            color_dict = self.tokens.setdefault("color", {})
            color_dict["heading"] = _heading_rgb()
            color_dict["title"] = _heading_rgb()
            color_dict["text"] = _text_rgb()
            color_dict["muted"] = _muted_rgb()
            # Some templates also write to "bg" / "background" — leave those alone.
        except Exception:
            # If tokens is frozen / has a custom color object, fall through;
            # the explicit overrides in _setup_styles and _render_heading
            # are the safety net.
            pass
        self.layout = LayoutCalculator(tokens)
        self.doc: Optional[Document] = None
        self._list_counters = {}
        self._figure_counter = 0
        self._table_counter = 0
        self._equation_counter = 0
        self._caption_labels: dict[str, str] = {}  # label -> "图 1-1" 占位文本
        self._bookmark_id_counter = 1000  # v1.4 P2-4 cross-reference bookmark id counter
        self._output_dir = output_dir
        self._chart_tmp_dir = None
        # v1.2: 内部 id 计数器
        self._next_id = 1000
        # v1.2: 批注收集
        self._comments: list[dict] = []  # [{id, author, date, text, paragraphs: [...(para, cid)]}]
        # v1.2: 脚注/尾注收集
        self._footnotes: list[dict] = []   # [{id, text}]
        self._endnotes: list[dict] = []
        # v1.4 P1-4: 参考文献渲染状态
        self._ref_counter = 0
        self._ref_style: str = "apa"
        # v1.4 P3-9c: auto_structure 状态（None 表示未启用）
        self._auto_lang: Optional[str] = None
        self._auto_mode: Optional[str] = None

    # ─── 对外入口 ────────────────────────────────────────────────

    @staticmethod
    def _pt_value(v) -> float:
        """Convert a token font size to float points."""
        if hasattr(v, "pt"):
            return float(v.pt)
        try:
            return float(v) / 12700.0
        except Exception:
            return float(v)

    def render(self, plan: dict, output_path: str) -> str:
        """按渲染计划生成 docx 并保存。"""
        self.doc = Document()
        self._output_dir = os.path.dirname(os.path.abspath(output_path)) or "."
        os.makedirs(self._output_dir, exist_ok=True)
        if self._chart_tmp_dir is not None:
            self._chart_tmp_dir.cleanup()
        self._chart_tmp_dir = _WritableTempDir("pro_docx_charts_", self._output_dir)

        # v1.4 P3-9c: 记录 auto_structure 解析出的 lang/mode（None 表示未启用/未识别）
        self._auto_lang = plan.get("auto_lang")
        self._auto_mode = plan.get("auto_mode")

        self._setup_page(plan)
        self._setup_styles()
        self._setup_header_footer(plan)

        # 文档级水印 / 页面边框（在所有节生效）
        self._apply_watermark(plan.get("watermark"))
        self._apply_page_border(plan.get("page_border"))

        # 文档元数据
        self._write_core_properties(plan.get("meta", {}))

        # 节点渲染
        for node in plan["nodes"]:
            self._render_node(node)

        # 尾注追加（若有）
        if self._endnotes:
            self._append_endnotes_section()

        # 后处理：写入 comments.xml / footnotes.xml / endnotes.xml part
        self._inject_comments_part()
        self._inject_footnotes_part()
        self._inject_endnotes_part()

        # v1.4 P3-4 / P3-6: 文档级分栏 / 行号（作用于最后一节 sectPr）
        self._apply_section_layout(plan)

        _safe_save_document(self.doc, output_path)
        if self._chart_tmp_dir is not None:
            self._chart_tmp_dir.cleanup()
            self._chart_tmp_dir = None
        return output_path

    def _next_revision_id(self) -> str:
        self._next_id += 1
        return str(self._next_id)

    # ─── 页面设置 ────────────────────────────────────────────────

    def _setup_page(self, plan: dict = None):
        page = self.tokens.get("page", {})
        # 应用到所有节
        for section in self.doc.sections:
            section.page_width = self.tokens["spacing"]["page_width"]
            section.page_height = self.tokens["spacing"]["page_height"]
            section.top_margin = page.get("margin_top", Inches(1))
            section.bottom_margin = page.get("margin_bottom", Inches(1))
            section.left_margin = page.get("margin_left", Inches(1))
            section.right_margin = page.get("margin_right", Inches(1))
            section.gutter = page.get("gutter", Inches(0)) or Inches(0)
            section.header_distance = page.get("header_distance", Inches(0.5))
            section.footer_distance = page.get("footer_distance", Inches(0.5))

            # 方向
            orientation = page.get("orientation", "portrait")
            if orientation == "landscape":
                section.orientation = WD_ORIENT.LANDSCAPE
                # 交换宽高（python-docx 要求）
                section.page_width, section.page_height = (
                    self.tokens["spacing"]["page_height"], self.tokens["spacing"]["page_width"]
                )
            else:
                section.orientation = WD_ORIENT.PORTRAIT

            # 首页不同 / 奇偶页不同
            section.different_first_page_header_footer = bool(
                page.get("different_first_page", False)
            )
            # 奇偶页不同需要 sectPr 设置 evenAndOddHeaders
            if page.get("different_odd_even", False):
                self._set_even_odd_headers(section)

    def _set_even_odd_headers(self, section):
        """在 sectPr 中设置 titlePg/evenAndOddHeaders。"""
        sectPr = section._sectPr
        existing = sectPr.find(qn("w:titlePg"))
        # evenAndOddHeaders 在 settings.xml 里设置（全局）
        settings = self.doc.settings.element
        eoh = settings.find(qn("w:evenAndOddHeaders"))
        if eoh is None:
            eoh = OxmlElement("w:evenAndOddHeaders")
            settings.append(eoh)

    def _apply_section_layout(self, plan: dict):
        """v1.4 P3-4 / P3-6: 在最后一节 sectPr 上设置分栏 / 行号。"""
        sections = self.doc.sections
        if not sections:
            return
        sectPr = sections[-1]._sectPr

        # P3-4: 分栏
        cols_cfg = plan.get("columns")
        if cols_cfg:
            self._set_cols(sectPr, cols_cfg)

        # P3-6: 行号
        ln_cfg = plan.get("line_numbers")
        if ln_cfg:
            self._set_line_numbers(sectPr, ln_cfg)

    def _set_cols(self, sectPr, cfg: dict):
        """在 sectPr 中插入/替换 w:cols。"""
        existing = sectPr.find(qn("w:cols"))
        if existing is not None:
            sectPr.remove(existing)
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"), str(int(cfg.get("count", 1))))
        if "space" in cfg:
            cols.set(qn("w:space"), str(int(cfg["space"])))
        if cfg.get("sep"):
            cols.set(qn("w:sep"), "1")
        # 按 OOXML 规范，w:cols 应在 sectPr 的正确位置（pgSz/pgMar 之后）
        # 简化处理：append 即可，Word 容错
        sectPr.append(cols)

    def _set_line_numbers(self, sectPr, cfg: dict):
        """在 sectPr 中插入/替换 w:lnNumType。"""
        existing = sectPr.find(qn("w:lnNumType"))
        if existing is not None:
            sectPr.remove(existing)
        ln = OxmlElement("w:lnNumType")
        ln.set(qn("w:countBy"), str(int(cfg.get("increment", 1))))
        ln.set(qn("w:start"), str(int(cfg.get("start", 1))))
        ln.set(qn("w:restart"), str(cfg.get("restart", "continuous")))
        if "distance" in cfg:
            ln.set(qn("w:distance"), str(int(cfg["distance"])))
        sectPr.append(ln)

    def _insert_paragraph_cols_sectPr(self, paragraph, cfg: dict):
        """v1.4 P3-4: 段落级分栏 — 在段落 pPr 末尾插入 sectPr/w:cols（连续分节符）。

        注意：这会在 Word 中表现为段末连续分节符，其后的内容将进入新节并按新栏数排版。
        段落级分栏只影响该段之后的内容，直到下一个 sectPr 出现。
        """
        pPr = paragraph._p.get_or_add_pPr()
        # 移除已有 sectPr（避免重复）
        existing = pPr.find(qn("w:sectPr"))
        if existing is not None:
            pPr.remove(existing)
        sectPr = OxmlElement("w:sectPr")
        # type=continuous 让分节符不换页
        sect_type = OxmlElement("w:type")
        sect_type.set(qn("w:val"), "continuous")
        sectPr.append(sect_type)
        # pgSz/pgMar 继承上级节（可省略，Word 自动继承）
        cols = OxmlElement("w:cols")
        cols.set(qn("w:num"), str(int(cfg.get("count", 1))))
        if "space" in cfg:
            cols.set(qn("w:space"), str(int(cfg["space"])))
        if cfg.get("sep"):
            cols.set(qn("w:sep"), "1")
        sectPr.append(cols)
        pPr.append(sectPr)

    def _setup_styles(self):
        """配置 NamedStyle。"""
        styles = self.doc.styles
        t = self.tokens
        f = t["font"]
        sp = t["spacing"]
        c = t["color"]
        align = t["alignment"]
        cn_font = f["family"]["cn"]
        cn_heading = f["family"]["cn_heading"]

        def _apply_font_style(style, font_name, cn_name, size, color, bold=False, italic=False):
            style.font.name = font_name
            style.font.size = size
            style.font.color.rgb = color
            style.font.bold = bold
            style.font.italic = italic
            rpr = style.element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:eastAsia"), cn_name)

        def _apply_para_format(pf, before=Pt(0), after=Pt(0), line_spacing=None,
                               first_line_indent=None, alignment=None, left_indent=None):
            pf.space_before = before
            pf.space_after = after
            if line_spacing is not None:
                pf.line_spacing = line_spacing
            if first_line_indent is not None:
                pf.first_line_indent = first_line_indent
            if alignment is not None:
                pf.alignment = alignment
            if left_indent is not None:
                pf.left_indent = left_indent

        normal = styles["Normal"]
        # v1.6.6: body text is locked to #333333 regardless of any
        # theme override; this is the root-cause fix for "fancy body color"
        # regressions caused by template/theme_overrides mutations.
        _apply_font_style(normal, f["family"]["default"], cn_font, f["size"]["body"], _text_rgb())
        _apply_para_format(
            normal.paragraph_format,
            before=sp["paragraph_before"],
            after=sp["paragraph_after"],
            line_spacing=sp["line_spacing"],
            alignment=align["body"],
        )

        # v1.6.6: Title style is always #1A1A1A (root-cause fix).
        title_style = styles["Title"]
        _apply_font_style(title_style, f["family"]["heading"], cn_heading,
                          f["size"]["title"], _heading_rgb(), bold=True)
        _apply_para_format(
            title_style.paragraph_format,
            before=sp["before_title"], after=sp["after_title"],
            alignment=align["title"], line_spacing=sp["line_spacing"],
        )

        # v1.6.6: All Heading 1-6 styles are forced to #1A1A1A.
        # Templates and theme_overrides["color"]["heading"] are no longer
        # allowed to control heading color — the canonical value wins.
        for lvl in range(1, 6):
            h_key = f"Heading {lvl}"
            h_style = styles[h_key]
            _apply_font_style(h_style, f["family"]["heading"], cn_heading,
                              f["size"][f"h{lvl}"], _heading_rgb(), bold=True)
            _apply_para_format(
                h_style.paragraph_format,
                before=sp[f"before_h{lvl}"], after=sp[f"after_h{lvl}"],
                line_spacing=sp["line_spacing"],
                alignment=align.get(f"h{lvl}", WD_ALIGN_PARAGRAPH.LEFT),
            )

    def _setup_header_footer(self, plan: dict):
        """配置页眉页脚（v1.2 增强：支持图片、X of Y、首页不同、奇偶页不同）。"""
        section = self.doc.sections[0]
        t = self.tokens
        hf_font = t["font"]["family"]["default"]
        hf_cn = t["font"]["family"]["cn"]
        hf_size = t["font"]["size"].get("header", Pt(9))
        hf_color = t["color"].get("header_text", t["color"]["muted"])
        ff_size = t["font"]["size"].get("footer", Pt(9))
        ff_color = t["color"].get("footer_text", t["color"]["muted"])

        header_cfg = plan.get("header") or t.get("header", {})
        if header_cfg:
            self._render_header_footer(
                section.header, header_cfg, hf_font, hf_cn, hf_size, hf_color, is_header=True
            )
            if not header_cfg.get("show_on_first_page", True):
                section.different_first_page_header_footer = True

        footer_cfg = plan.get("footer") or t.get("footer", {})
        if footer_cfg is None:
            footer_cfg = {}
        if footer_cfg or footer_cfg.get("page_number", True):
            self._render_header_footer(
                section.footer, footer_cfg, hf_font, hf_cn, ff_size, ff_color, is_header=False
            )

    def _render_header_footer(self, hf_part, cfg, font_name, cn_font, size, color, is_header=False):
        """渲染单个 header/footer part 的内容。"""
        t = self.tokens
        p = hf_part.paragraphs[0] if hf_part.paragraphs else hf_part.add_paragraph()
        align_key = "header" if is_header else "footer"
        p.alignment = t["alignment"].get(align_key, WD_ALIGN_PARAGRAPH.CENTER)

        # 页眉图片
        img = cfg.get("image_path")
        if img and is_header and os.path.exists(img):
            run = p.add_run()
            try:
                w_in = cfg.get("image_width_inches")
                if w_in:
                    run.add_picture(img, width=Inches(w_in))
                else:
                    run.add_picture(img, height=Inches(0.4))
            except Exception:
                pass
            if cfg.get("text"):
                r = p.add_run("  " + cfg.get("text", ""))
                _set_run_font(r, font_name, cn_font)
                r.font.size = size
                r.font.color.rgb = color
        else:
            if cfg.get("text"):
                run = p.add_run(cfg["text"])
                if cfg.get("text") and (cfg.get("page_number") or cfg.get("page_x_of_y")):
                    run = p.add_run("  ")
                _set_run_font(run, font_name, cn_font)
                run.font.size = size
                run.font.color.rgb = color

        # 页脚页码 / X of Y
        if not is_header:
            page_num = cfg.get("page_number", True)
            page_xofy = cfg.get("page_x_of_y", False)
            page_xofy_cn = cfg.get("page_x_of_y_cn", False)
            if page_xofy or page_xofy_cn:
                # "Page X of Y" / "第 X 页，共 Y 页"
                if page_xofy_cn:
                    r = p.add_run("第 ")
                    _set_run_font(r, font_name, cn_font); r.font.size=size; r.font.color.rgb=color
                    _add_field(p, "PAGE")
                    r = p.add_run(" 页，共 ")
                    _set_run_font(r, font_name, cn_font); r.font.size=size; r.font.color.rgb=color
                    _add_field(p, "NUMPAGES")
                    r = p.add_run(" 页")
                    _set_run_font(r, font_name, cn_font); r.font.size=size; r.font.color.rgb=color
                else:
                    r = p.add_run("Page ")
                    _set_run_font(r, font_name, cn_font); r.font.size=size; r.font.color.rgb=color
                    _add_field(p, "PAGE")
                    r = p.add_run(" of ")
                    _set_run_font(r, font_name, cn_font); r.font.size=size; r.font.color.rgb=color
                    _add_field(p, "NUMPAGES")
            elif page_num:
                _add_field(p, "PAGE")

    # ─── 文档级水印 / 页面边框 ─────────────────────────────────────

    def _apply_watermark(self, watermark_cfg: dict | None):
        wm = watermark_cfg or self.tokens.get("watermark", {})
        if not wm or not wm.get("enabled", False):
            return
        text = wm.get("text", "DRAFT")
        color = wm.get("color", self.tokens["color"]["watermark"])
        font_size = wm.get("font_size", Pt(60))
        rotation = wm.get("rotation", -45)
        color_hex = _rgb_to_hex(color)
        fs_pt = int(font_size.pt) if hasattr(font_size, "pt") else int(font_size)
        for section in self.doc.sections:
            header = section.header
            # 水印通过 VML shape 嵌在页眉段落中实现
            wp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            run = wp.add_run()
            pict_xml = f'''
            <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    xmlns:v="urn:schemas-microsoft-com:vml"
                    xmlns:o="urn:schemas-microsoft-com:office:office">
              <v:shapetype id="_x0000_t136" coordsize="21600,21600" o:spt="136" adj="10800" path="m@7,l@8,m@5,21600l@6,21600e">
                <v:formulas>
                  <v:f eqn="sum #0 0 10800"/><v:f eqn="prod #0 2 1"/>
                  <v:f eqn="sum 21600 0 @1"/><v:f eqn="sum 0 0 @2"/>
                  <v:f eqn="sum 21600 0 @3"/><v:f eqn="if @0 @3 0"/>
                  <v:f eqn="if @0 21600 @1"/><v:f eqn="if @0 0 @2"/>
                  <v:f eqn="if @0 @4 21600"/><v:f eqn="mid @5 @6"/>
                  <v:f eqn="mid @8 @5"/><v:f eqn="mid @7 @8"/>
                  <v:f eqn="mid @6 @7"/><v:f eqn="sum @6 0 @5"/>
                </v:formulas>
                <v:path textpathok="t" o:connecttype="custom" o:connectlocs="@9,0;@10,10800;@11,21600;@12,10800" o:connectangles="270,180,90,0"/>
                <v:textpath on="t" fitshape="t"/>
              </v:shapetype>
              <v:shape id="PowerPlusWaterMarkObject" o:spid="_x0000_s2049" type="#_x0000_t136"
                       style="position:absolute;margin-left:0;margin-top:0;width:500pt;height:100pt;
                              z-index:-251658240;mso-position-horizontal:center;
                              mso-position-horizontal-relative:margin;
                              mso-position-vertical:center;mso-position-vertical-relative:margin;
                              rotation:{rotation}"
                       o:allowincell="f" fillcolor="#{color_hex}" stroked="f">
                <v:fill opacity=".5"/>
                <v:textpath style="font-family:&quot;{self.tokens['font']['family']['cn_heading']}&quot;;font-size:{fs_pt}pt" string="{text}"/>
              </v:shape>
            </w:pict>
            '''
            from lxml import etree
            try:
                elem = etree.fromstring(pict_xml)
                run._r.append(elem)
            except Exception as e:
                print(f"[watermark] failed to inject: {e}")

    def _apply_page_border(self, pb_cfg: dict | None):
        pb = pb_cfg or self.tokens.get("page_border", {})
        if not pb or not pb.get("enabled", False):
            return
        color = pb.get("color", self.tokens["color"]["page_border"])
        color_hex = _rgb_to_hex(color)
        sz = pb.get("size", 6)
        space = pb.get("space", 24)
        offset = pb.get("offset_from", "page")
        style = pb.get("style", "single")
        for section in self.doc.sections:
            sectPr = section._sectPr
            pgBorders = sectPr.find(qn("w:pgBorders"))
            if pgBorders is None:
                pgBorders = OxmlElement("w:pgBorders")
                pgBorders.set(qn("w:offsetFrom"), offset)
                sectPr.append(pgBorders)
            for edge in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{edge}")
                b.set(qn("w:val"), style)
                b.set(qn("w:sz"), str(sz))
                b.set(qn("w:space"), str(space))
                b.set(qn("w:color"), color_hex)
                pgBorders.append(b)

    # ─── 元数据 ──────────────────────────────────────────────────

    def _write_core_properties(self, meta: dict):
        cp = self.doc.core_properties
        if meta.get("title"):
            cp.title = meta["title"]
        if meta.get("author"):
            cp.author = meta["author"]
        if meta.get("subject"):
            cp.subject = meta["subject"]
        # keywords 支持 list 或 str
        kw = meta.get("keywords")
        if kw:
            cp.keywords = ", ".join(kw) if isinstance(kw, list) else str(kw)
        if meta.get("category"):
            cp.category = meta["category"]
        if meta.get("comments"):
            cp.comments = meta["comments"]
        if meta.get("status"):
            cp.content_status = meta["status"]
        if meta.get("company") or meta.get("manager"):
            # 写入 custom properties
            self._write_custom_properties(meta)

    def _write_custom_properties(self, meta: dict):
        """写入自定义属性（company/manager 等）。"""
        try:
            from docx.opc.constants import CONTENT_TYPE as CT
            # python-docx 没有直接暴露 custom properties API，用 XML 追加
            part = self.doc.part
            # 查找或创建 custom.xml part
            rels = part.rels
            custom_part = None
            for rel in rels.values():
                if "customXml" in rel.reltype or "custom-properties" in rel.reltype:
                    custom_part = rel.target_part
                    break
            # 简化：仅在有 manager/company 时写 app.xml（extended properties）
            app_part = None
            for rel in rels.values():
                if rel.reltype == RT.EXTENDED_PROPERTIES:
                    app_part = rel.target_part
                    break
            if app_part is not None and meta.get("company"):
                try:
                    from lxml import etree
                    el = etree.fromstring(app_part.blob)
                    ns = "{http://schemas.openxmlformats.org/officeDocument/2006/extended-properties}"
                    company_el = el.find(f"{ns}Company")
                    if company_el is None:
                        company_el = etree.SubElement(el, f"{ns}Company")
                        el.append(company_el)
                    company_el.text = meta["company"]
                    if meta.get("manager"):
                        m_el = el.find(f"{ns}Manager")
                        if m_el is None:
                            m_el = etree.SubElement(el, f"{ns}Manager")
                            el.append(m_el)
                        m_el.text = meta["manager"]
                    app_part._blob = etree.tostring(el, xml_declaration=True, encoding="UTF-8", standalone=True)
                except Exception as e:
                    print(f"[extended-props] failed: {e}")
        except Exception as e:
            print(f"[custom-props] failed: {e}")

    # ─── 节点分发 ────────────────────────────────────────────────

    def _render_node(self, node: dict):
        ntype = node["node_type"]
        handlers = {
            "title_block": self._render_title_block,
            "abstract_block": self._render_abstract,
            "heading": self._render_heading,
            "paragraph": self._render_paragraph,
            "list_item": self._render_list_item,
            "table": self._render_table,
            "figure": self._render_figure,
            "chart": self._render_chart,
            "kpi_card": self._render_kpi_card,
            "callout": self._render_callout,
            "page_break": self._render_page_break,
            "toc": self._render_toc,
            "reference_item": self._render_reference_item,
            "references_block": self._render_references_block,
            # v1.2
            "revision": self._render_revision_inline,
            "comment": self._render_comment_block,
            "footnote": self._render_footnote_ref,
            "endnote": self._render_endnote_ref,
            "watermark": lambda n: self._apply_watermark(n),
            "page_border": lambda n: self._apply_page_border(n),
            "equation": self._render_equation,
            "signature_block": self._render_signature_block,
            "signature_line": self._render_signature_line,
            # v1.3
            "svg_shape": self._render_svg_shape,
            # v1.4 P2-4
            "ref": self._render_ref,
        }
        h = handlers.get(ntype)
        if h:
            h(node)
        else:
            print(f"[renderer] Warning: unknown node type '{ntype}', skipped")

    # ─── 各节点渲染 ──────────────────────────────────────────────

    def _render_title_block(self, node: dict):
        t = self.tokens
        cn_heading = t["font"]["family"]["cn_heading"]
        cn_font = t["font"]["family"]["cn"]

        # v1.6.0: 变体封面样式
        cover_style = {}
        if self.variant:
            try:
                from ..variant_tokens import get_variant_cover_style
                cover_style = get_variant_cover_style(self.variant)
            except ImportError:
                pass

        title_size = cover_style.get("title_size", t["font"]["size"]["title"])
        title_color = cover_style.get("title_color", t["color"]["title"])
        title_bold = cover_style.get("title_bold", True)
        subtitle_color = cover_style.get("subtitle_color", t["color"]["muted"])

        # v1.6.0: variant 暗色背景（如 modern_tech）
        bg_color = cover_style.get("bg_color")
        if bg_color:
            # 在标题前插入空白段落做背景区域（简化方案：用底纹模拟）
            pass

        # 标题前空行（spacious 变体留更多呼吸空间）
        if cover_style.get("spacing", "") == "spacious":
            spacer = self.doc.add_paragraph()
            spacer.paragraph_format.space_after = Pt(12)

        title_p = self.doc.add_paragraph()
        title_p.alignment = t["alignment"]["title"]
        title_p.paragraph_format.space_before = t["spacing"]["before_title"]
        title_p.paragraph_format.space_after = t["spacing"]["after_title"]
        title_p.paragraph_format.line_spacing = t["spacing"]["line_spacing"]
        run = title_p.add_run(node["title"])
        _set_run_font(run, t["font"]["family"]["heading"], cn_heading)
        run.font.size = title_size
        # v1.6.6: title color is always #1A1A1A (root-cause fix);
        # do not honor cover_style["title_color"] when it is a fancy
        # theme accent (e.g. elegant_luxury gold / modern_tech teal).
        run.font.color.rgb = _heading_rgb()
        run.font.bold = title_bold

        # v1.6.0: 装饰分割线（corporate_formal / modern_tech / bold_impact / elegant_luxury）
        if cover_style.get("divider"):
            div_p = self.doc.add_paragraph()
            div_p.alignment = t["alignment"]["title"]
            div_p.paragraph_format.space_before = Pt(6)
            div_p.paragraph_format.space_after = Pt(6)
            # 使用 ── 字符模拟分割线
            div_width = cover_style.get("divider_width_pct", 0.3)
            div_char_count = max(3, int(40 * div_width))
            div_text = "─" * div_char_count
            div_run = div_p.add_run(div_text)
            div_run.font.size = Pt(10)
            div_run.font.color.rgb = cover_style.get("divider_color", t["color"]["accent"])

        if node.get("subtitle"):
            sp = self.doc.add_paragraph()
            sp.alignment = t["alignment"]["subtitle"]
            sp.paragraph_format.space_after = t["spacing"]["after_subtitle"]
            r = sp.add_run(node["subtitle"])
            _set_run_font(r, t["font"]["family"]["default"], cn_font)
            r.font.size = t["font"]["size"]["subtitle"]
            r.font.color.rgb = subtitle_color

        meta_parts = [node.get("author"), node.get("institution"), node.get("date")]
        meta_parts = [m for m in meta_parts if m]
        if meta_parts:
            mp = self.doc.add_paragraph()
            mp.alignment = t["alignment"]["meta"]
            mp.paragraph_format.space_after = Pt(12)
            r = mp.add_run("  ·  ".join(meta_parts))
            _set_run_font(r, t["font"]["family"]["default"], cn_font)
            r.font.size = Pt(11)
            r.font.color.rgb = t["color"]["muted"]

    def _render_abstract(self, node: dict):
        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        sp = t["spacing"]

        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = sp["first_line_indent"]
        p.paragraph_format.line_spacing = sp["line_spacing"]
        p.paragraph_format.space_after = Pt(6)

        label_run = p.add_run("Abstract: " if self.lang == "en" else "摘要：")
        _set_run_font(label_run, t["font"]["family"]["heading"], t["font"]["family"]["cn_heading"])
        label_run.font.size = t["font"]["size"]["abstract"]
        label_run.font.bold = True
        label_run.font.color.rgb = t["color"]["heading"]

        # v1.4 P2-4b: abstract 文本支持 inline {ref}
        if node.get("inline_segments"):
            self._render_inline_segments(p, node["inline_segments"],
                                         size=t["font"]["size"]["abstract"],
                                         color=t["color"]["text"],
                                         font_key="default", cn_font_key="cn")
        else:
            text_run = p.add_run(node.get("text", ""))
            _set_run_font(text_run, t["font"]["family"]["default"], cn_font)
            text_run.font.size = t["font"]["size"]["abstract"]
            text_run.font.color.rgb = t["color"]["text"]

        keywords = node.get("keywords", [])
        if keywords:
            kp = self.doc.add_paragraph()
            kp.paragraph_format.space_after = Pt(12)
            kl_run = kp.add_run("Keywords: " if self.lang == "en" else "关键词：")
            _set_run_font(kl_run, t["font"]["family"]["heading"], t["font"]["family"]["cn_heading"])
            kl_run.font.size = t["font"]["size"]["abstract"]
            kl_run.font.bold = True
            kl_run.font.color.rgb = t["color"]["heading"]
            kv_run = kp.add_run("; ".join(keywords) if self.lang == "en" else "；".join(keywords))
            _set_run_font(kv_run, t["font"]["family"]["default"], cn_font)
            kv_run.font.size = t["font"]["size"]["abstract"]
            kv_run.font.color.rgb = t["color"]["text"]

    def _render_heading(self, node: dict):
        """Render a heading paragraph with the configured level (1-4) style.

        渲染标题段落，自动套用 Heading 1-4 样式（主题色/字号/加粗/段前段后间距）。标题文本支持
        bookmark 用于交叉引用与 TOC 识别。v1.4 P3-9c：若 heading 由 auto_structure 自动识别
        且带有 style_override，则按 (lang, mode) 组合覆盖字号/字体/加粗/对齐/段前段后。"""

        lvl = node.get("level", 1)
        style_override = node.get("style_override") if node.get("auto_structured") else None
        text = node["text"]
        if style_override and style_override.get("all_caps"):
            text = text.upper()

        if style_override:
            # 自动识别标题：不使用 Heading N 样式（避免被主题字体/字号覆盖），直接以普通段落+自定义格式渲染
            p = self.doc.add_paragraph()
            pf = p.paragraph_format
            align = style_override.get("alignment", "left")
            p.alignment = {
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
                "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
            sb = style_override.get("space_before_pt")
            sa = style_override.get("space_after_pt")
            if sb is not None:
                pf.space_before = Pt(sb)
            if sa is not None:
                pf.space_after = Pt(sa)
            run = p.add_run(text)
            lat = style_override.get("font_latin", self.tokens["font"]["family"]["heading"])
            ea = style_override.get("font_eastasia", self.tokens["font"]["family"]["cn_heading"])
            _set_run_font(run, lat, ea)
            run.font.size = Pt(style_override.get("font_size_pt", 14))
            run.font.bold = bool(style_override.get("bold", True))
            run.font.italic = bool(style_override.get("italic", False))
            # v1.6.6: heading color is always #1A1A1A (root-cause fix);
            # do not read tokens["color"]["heading"] here.
            run.font.color.rgb = _heading_rgb()
        else:
            style_name = f"Heading {lvl}"
            p = self.doc.add_paragraph(style=style_name)
            run = p.add_run(text)
            _set_run_font(run, self.tokens["font"]["family"]["heading"],
                          self.tokens["font"]["family"]["cn_heading"])
            # v1.6.6: explicitly force heading color on the run itself
            # (Heading N style already gets #1A1A1A from _setup_styles,
            # but we belt-and-suspenders the run color so an inherited
            # default style cannot leak a fancy theme color).
            run.font.color.rgb = _heading_rgb()
        if node.get("comment"):
            self._attach_comment(p, node["comment"])

    def _render_paragraph(self, node: dict):
        """Render a normal body paragraph, supporting runs, drop_cap, and inline refs.

        渲染正文段落。支持 runs 数组精细控制（bold/italic/color/size）、drop_cap 首字下沉（首字 3 倍字号
        加粗 heading 色）、indent 首行缩进、align 对齐、columns 段落级分栏切换。中文字体自动设置 eastAsia。"""

        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        style = node.get("style", "normal")
        sp = t["spacing"]

        # v1.4 P3-9c: auto_structure 模式下，正文（normal/abstract/caption/quote/footnote）使用
        # 按 (lang, mode) 选择的专业字体（学术宋体五号 / 商业雅黑11pt / TNR12 / Calibri 11）
        auto_body = None
        if self._auto_lang and self._auto_mode:
            # 惰性 import 避免循环
            try:
                from ..compiler.parser import _auto_body_font
                auto_body = _auto_body_font(self._auto_lang, self._auto_mode)
            except Exception:
                auto_body = None
        if auto_body:
            cn_font = auto_body.get("font_eastasia", cn_font)
            body_latin = auto_body.get("font_latin", t["font"]["family"]["default"])
            body_size = Pt(auto_body.get("font_size_pt", 10.5))
        else:
            body_latin = t["font"]["family"]["default"]
            body_size = t["font"]["size"]["body"]

        p = self.doc.add_paragraph()

        if style == "quote":
            p.paragraph_format.left_indent = sp["quote_indent"]
            p.paragraph_format.line_spacing = sp["line_spacing"]
            p.paragraph_format.space_after = sp["paragraph_after"]
            _set_paragraph_border_left(p, _rgb_to_hex(t["color"]["quote_border"]), 12)
            # v1.4 P2-4b: 支持 inline {ref}（quote 段落）
            if node.get("inline_segments"):
                self._render_inline_segments(p, node["inline_segments"],
                                             size=self._pt_value(body_size),
                                             color=t["color"]["muted"],
                                             italic=True,
                                             font_key="default", cn_font_key="cn")
            else:
                run = p.add_run(node["text"])
                _set_run_font(run, body_latin, cn_font)
                run.font.size = body_size
                run.font.italic = True
                run.font.color.rgb = t["color"]["muted"]

        elif style == "code":
            p.paragraph_format.left_indent = sp["code_indent"]
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = sp["paragraph_after"]
            _set_paragraph_shading(p, _rgb_to_hex(t["color"]["code_bg"]))
            # 代码块不解析 {ref}（parser 已跳过，这里也强制使用原始 text 以双保险）
            run = p.add_run(node.get("text", ""))
            _set_run_font(run, t["font"]["family"]["code"], t["font"]["family"]["code"])
            run.font.size = t["font"]["size"]["code"]
            run.font.color.rgb = t["color"]["text"]

        elif style == "abstract":
            p.paragraph_format.first_line_indent = sp["first_line_indent"]
            p.paragraph_format.line_spacing = sp["line_spacing"]
            if node.get("inline_segments"):
                self._render_inline_segments(p, node["inline_segments"],
                                             size=self._pt_value(t["font"]["size"]["abstract"]),
                                             color=t["color"]["text"],
                                             font_key="default", cn_font_key="cn")
            else:
                run = p.add_run(node["text"])
                _set_run_font(run, body_latin, cn_font)
                run.font.size = t["font"]["size"]["abstract"]

        elif style == "footnote":
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run(node["text"])
            _set_run_font(run, body_latin, cn_font)
            run.font.size = t["font"]["size"]["footnote"]
            run.font.color.rgb = t["color"]["muted"]

        elif style == "caption":
            # v1.4 P3-9b: 独立 caption 段落（未绑定到 figure/table 时）
            p.alignment = t["alignment"]["caption"]
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            if node.get("inline_segments"):
                self._render_inline_segments(p, node["inline_segments"],
                                             size=self._pt_value(t["font"]["size"]["caption"]),
                                             color=t["color"]["muted"],
                                             bold=True,
                                             font_key="default", cn_font_key="cn")
            else:
                run = p.add_run(node.get("text", ""))
                _set_run_font(run, body_latin, cn_font)
                run.font.size = t["font"]["size"]["caption"]
                run.font.color.rgb = t["color"]["muted"]
                run.font.bold = True

        else:  # normal
            p.paragraph_format.first_line_indent = sp["first_line_indent"]
            p.paragraph_format.line_spacing = sp["line_spacing"]
            p.paragraph_format.space_after = sp["paragraph_after"]
            p.alignment = t["alignment"]["body"]

            # runs 高级 API（允许段落内含 revision/comment/footnote 等 inline 节点）
            if node.get("runs"):
                # v1.4 P3-5: runs 模式下 drop_cap 仅对首个 text run 生效（简化）
                self._render_runs(p, node["runs"], body_latin=body_latin, cn_font=cn_font, body_size=body_size)
            elif node.get("inline_segments") and not node.get("drop_cap"):
                # v1.4 P2-4b: 解析出的内联 ref/text 段交替渲染
                self._render_inline_segments(
                    p, node["inline_segments"],
                    size=self._pt_value(body_size),
                    color=t["color"]["text"],
                    bold=bool(node.get("bold")),
                    italic=bool(node.get("italic")),
                    font_latin=body_latin, cn_font=cn_font,
                )
            else:
                text = node["text"] or ""
                if node.get("drop_cap") and text.strip():
                    # v1.4 P3-5: 首字下沉——首字 3 倍字号加粗，左缩进清零
                    p.paragraph_format.first_line_indent = Pt(0)
                    first = text[0]
                    rest = text[1:]
                    cap_pt = self._pt_value(body_size) * 3
                    run_cap = p.add_run(first)
                    _set_run_font(run_cap, body_latin, cn_font)
                    run_cap.font.size = Pt(cap_pt)
                    run_cap.font.bold = True
                    run_cap.font.color.rgb = t["color"]["heading"]
                    if rest:
                        run_rest = p.add_run(rest)
                        _set_run_font(run_rest, body_latin, cn_font)
                        run_rest.font.size = body_size
                        run_rest.font.color.rgb = t["color"]["text"]
                        if node.get("bold"):
                            run_rest.font.bold = True
                        if node.get("italic"):
                            run_rest.font.italic = True
                else:
                    run = p.add_run(text)
                    _set_run_font(run, body_latin, cn_font)
                    run.font.size = body_size
                    run.font.color.rgb = t["color"]["text"]
                    if node.get("bold"):
                        run.font.bold = True
                    if node.get("italic"):
                        run.font.italic = True

        if node.get("comment") and not node.get("runs"):
            self._attach_comment(p, node["comment"])

        # v1.4 P3-4: 段落级分栏（在 pPr 末尾插入 sectPr 实现段末连续分节符）
        if node.get("columns"):
            self._insert_paragraph_cols_sectPr(p, node["columns"])

    def _render_runs(self, paragraph, runs: list, *, body_latin=None, cn_font=None, body_size=None):
        """渲染段落内的多 run 节点（支持 revision/footnote/comment inline）。"""
        t = self.tokens
        if cn_font is None:
            cn_font = t["font"]["family"]["cn"]
        if body_latin is None:
            body_latin = t["font"]["family"]["default"]
        if body_size is None:
            body_size = t["font"]["size"]["body"]
        for r in runs:
            rtype = r.get("type", "text")
            if rtype in ("text", None):
                run = paragraph.add_run(r.get("text", ""))
                _set_run_font(run, body_latin, cn_font)
                run.font.size = body_size
                run.font.color.rgb = t["color"]["text"]
                if r.get("bold"):
                    run.font.bold = True
                if r.get("italic"):
                    run.font.italic = True
            elif rtype == "revision":
                self._render_revision_run(paragraph, r)
            elif rtype == "footnote_ref":
                self._render_footnote_ref_inline(paragraph, r)
            elif rtype == "comment":
                self._attach_comment(paragraph, r)
                # 同时插入 text
                if r.get("text"):
                    run = paragraph.add_run(r.get("anchor_text", ""))
                    _set_run_font(run, body_latin, cn_font)
                    run.font.size = body_size

    def _render_list_item(self, node: dict):
        """Render a single bullet/ordered list item with automatic numbering.

        渲染单个列表项，支持多层嵌套（level 0/1/2）。通过 _ensure_numbering_part 自动注入 numbering.xml，
        根据 ordered/level 选择 numId/ilvl 与 prefix 符号；同组相邻列表项共享同一 numId 以保证连续编号。"""

        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        sp = t["spacing"]
        ordered = bool(node.get("ordered", True))
        level = int(node.get("level", 0))
        level = min(level, 3)  # cap to 3
        text = node.get("text", "")
        list_id = node.get("list_id")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = self.layout.list_indent_for_level(level)
        p.paragraph_format.line_spacing = sp["line_spacing"]
        p.paragraph_format.space_after = Pt(2)

        # v1.4 P2-3: use Word native numbering (List Number / List Bullet styles
        # with ilvl via numPr) so nested lists render correctly with automatic
        # counters. Falls back to manual prefix (legacy) if numbering setup fails.
        used_numbering = False
        try:
            if list_id and ordered:
                num_id = self._ensure_numbering(list_id, ordered,
                                               start=int(node.get("_list_start", 1) or 1) if node.get("_list_start") else 1,
                                               is_root=bool(node.get("_list_start")))
                self._apply_numPr(p, num_id, level)
                # Set paragraph style hint so Word recognizes bullet/number indentation
                try:
                    p.style = self.doc.styles["List Paragraph"]
                except Exception:
                    pass
                used_numbering = True
        except Exception as e:
            print(f"[list_item] numbering setup failed, fallback to manual prefix: {e}")

        if not used_numbering:
            # Legacy manual prefix rendering (kept for backward compatibility)
            if ordered:
                key = f"{list_id or 'default'}:{level}"
                self._list_counters[key] = self._list_counters.get(key, 0) + 1
                # reset deeper levels
                for k in list(self._list_counters.keys()):
                    if k.startswith(f"{list_id or 'default'}:"):
                        try:
                            kl = int(k.split(":")[-1])
                            if kl > level:
                                self._list_counters[k] = 0
                        except Exception:
                            pass
                prefix = _ordered_prefix(level, self._list_counters[key])
            else:
                prefix = _bullet_prefix(level)
            pre_run = p.add_run(prefix)
            _set_run_font(pre_run, t["font"]["family"]["default"], cn_font)
            pre_run.font.size = t["font"]["size"]["body"]
            pre_run.font.color.rgb = t["color"]["accent"]

        # v1.4 P2-4b: 支持 inline_segments（list_item 中 {ref xxx}）
        if node.get("inline_segments"):
            self._render_inline_segments(p, node["inline_segments"],
                                         size=t["font"]["size"]["body"],
                                         color=t["color"]["text"],
                                         font_key="default", cn_font_key="cn")
        else:
            run = p.add_run(text)
            _set_run_font(run, t["font"]["family"]["default"], cn_font)
            run.font.size = t["font"]["size"]["body"]
            run.font.color.rgb = t["color"]["text"]

    # ------------------------------------------------------------------
    # v1.4 P2-3: Multi-level numbering helpers
    # ------------------------------------------------------------------
    def _ensure_numbering_part(self):
        """Return (numbering_part, numbering_elem), creating if necessary."""
        if hasattr(self, "_numbering_elem") and self._numbering_elem is not None:
            return self._numbering_part, self._numbering_elem
        from docx.oxml import parse_xml
        from docx.opc.constants import CONTENT_TYPE as CT, RELATIONSHIP_TYPE as RT
        doc_part = self.doc.part
        # Find existing numbering part
        numbering_part = None
        for rel in doc_part.rels.values():
            if rel.reltype == RT.NUMBERING:
                numbering_part = rel.target_part
                break
        if numbering_part is None:
            # Create minimal numbering part
            W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
            xml_bytes = (
                f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                f'<w:numbering xmlns:w="{W_NS}"></w:numbering>'
            ).encode("utf-8")
            from docx.opc.part import Part
            from docx.opc.packuri import PackURI
            partname = PackURI("/word/numbering.xml")
            numbering_part = Part(
                partname,
                CT.WML_NUMBERING,
                xml_bytes,
                doc_part.package,
            )
            doc_part.relate_to(numbering_part, RT.NUMBERING)
        numbering_elem = numbering_part.element
        self._numbering_part = numbering_part
        self._numbering_elem = numbering_elem
        self._num_id_registry = {}  # list_id -> abstractNumId
        return numbering_part, numbering_elem

    def _next_abstract_num_id(self) -> int:
        self._abstract_num_counter = getattr(self, "_abstract_num_counter", 0) + 1
        # avoid conflict with existing abstractNum
        existing_ids = {int(e.get(qn("w:abstractNumId"), "0")) for e in self._numbering_elem.findall(qn("w:abstractNum"))}
        while self._abstract_num_counter in existing_ids:
            self._abstract_num_counter += 1
        return self._abstract_num_counter

    def _next_num_id(self) -> int:
        self._num_counter = getattr(self, "_num_counter", 0) + 1
        existing_ids = {int(e.get(qn("w:numId"), "0")) for e in self._numbering_elem.findall(qn("w:num"))}
        while self._num_counter in existing_ids:
            self._num_counter += 1
        return self._num_counter

    def _ensure_numbering(self, list_id: str, ordered: bool, start: int = 1, is_root: bool = False) -> int:
        """Ensure an abstractNum + num definition exists for this list_id; return numId."""
        _, n_elem = self._ensure_numbering_part()
        if list_id in self._num_id_registry:
            return self._num_id_registry[list_id]
        abs_id = self._next_abstract_num_id()
        num_id = self._next_num_id()
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        # Build abstractNum XML with up to 4 levels (0..3)
        levels_xml = []
        for ilvl in range(4):
            lvl_fmt = _level_format(ordered, ilvl)
            lvl_text = _level_lvlText(ordered, ilvl)
            ind_left = 360 + ilvl * 360  # twips; ~0.25" per level
            hanging = 360
            bullet_char = _bullet_char(ilvl) if not ordered else ""
            if ordered:
                rpr = ""
            else:
                # Symbol font for round bullets; Wingdings for squares to match Word
                bullet_font = "Symbol" if ilvl in (0, 1) else "Wingdings"
                rpr = (
                    f'<w:rPr><w:rFonts w:ascii="{bullet_font}" w:hAnsi="{bullet_font}" '
                    f'w:hint="default"/></w:rPr>'
                )
            start_val = start if ilvl == 0 else 1
            suff = "tab"
            levels_xml.append(
                f'<w:lvl w:ilvl="{ilvl}">'
                f'<w:start w:val="{start_val}"/>'
                f'<w:numFmt w:val="{lvl_fmt if ordered else "bullet"}"/>'
                f'<w:lvlText w:val="{lvl_text if ordered else bullet_char}"/>'
                f'<w:lvlJc w:val="left"/>'
                f'<w:pPr><w:ind w:left="{ind_left}" w:hanging="{hanging}"/></w:pPr>'
                f'{rpr}'
                f'</w:lvl>'
            )
        abstract_xml = (
            f'<w:abstractNum xmlns:w="{W_NS}" w:abstractNumId="{abs_id}">'
            f'<w:multiLevelType w:val="hybridMultilevel"/>'
            + "".join(levels_xml) +
            f'</w:abstractNum>'
        )
        num_xml = (
            f'<w:num xmlns:w="{W_NS}" w:numId="{num_id}">'
            f'<w:abstractNumId w:val="{abs_id}"/>'
            f'</w:num>'
        )
        from docx.oxml import parse_xml
        abs_elem = parse_xml(abstract_xml)
        num_elem = parse_xml(num_xml)
        # Insert abstractNum BEFORE any existing <w:num> elements per schema
        first_num = n_elem.find(qn("w:num"))
        if first_num is not None:
            first_num.addprevious(abs_elem)
        else:
            n_elem.append(abs_elem)
        n_elem.append(num_elem)
        self._num_id_registry[list_id] = num_id
        return num_id

    def _apply_numPr(self, paragraph, num_id: int, ilvl: int):
        """Attach w:numPr to a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        # remove any existing numPr
        for old in pPr.findall(qn("w:numPr")):
            pPr.remove(old)
        numPr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(int(ilvl)))
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), str(int(num_id)))
        numPr.append(ilvl_el)
        numPr.append(numId_el)
        pPr.append(numPr)


    def _render_table(self, node: dict):
        """Render a table with header row, auto column widths, and themed styling.

        渲染表格：第一行为表头（加粗/主题底色），数据行隔行变色，自动按列数平均分配列宽；支持 header_rows
        多级表头、align 列对齐、col_widths 自定义列宽；表头跨页重复通过 ooxml tblHeader 实现。"""

        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        headers = node["headers"]
        rows = node["rows"]
        col_widths = node.get("col_widths")
        header_repeat = node.get("header_repeat", t["table"].get("header_repeat", True))

        if not headers:
            return

        # v1.4 P2-4: auto-numbered caption ABOVE table
        caption = node.get("caption")
        label = node.get("label")
        caption_text = ""
        caption_para = None
        if caption or label:
            self._table_counter += 1
            num = self._table_counter
            prefix = _caption_prefix("table", self.lang)
            caption_text = f"{prefix}{num}"
            if caption:
                caption_text = f"{caption_text}: {caption}"
            # Register label for ref
            if label:
                self._caption_labels[label] = caption_text
            caption_para = self.doc.add_paragraph()
            caption_para.alignment = t["alignment"]["caption"]
            caption_para.paragraph_format.space_before = Pt(6)
            caption_para.paragraph_format.space_after = Pt(3)
            self._write_caption_with_bookmark(caption_para, label, caption_text,
                                              _tbl_label_id(num), t, cn_font)

        n_cols = len(headers)
        n_rows = len(rows) + 1
        table = self.doc.add_table(rows=n_rows, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        widths = self.layout.compute_table_col_widths(headers, rows, col_widths)
        border_attrs = {"val": "single", "sz": "4", "color": _rgb_to_hex(t["color"]["table_border"])}
        header_bg = _rgb_to_hex(t["color"]["table_header_bg"])
        alt_bg = _rgb_to_hex(t["color"]["table_alt_row"])

        # 表头
        hdr_row = table.rows[0]
        _set_table_row_no_split(hdr_row)
        if header_repeat:
            _set_table_header_repeat(hdr_row)
        hdr_cells = hdr_row.cells
        for i, h in enumerate(headers):
            cell = hdr_cells[i]
            cell.width = widths[i] if i < len(widths) else None
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if t["table"]["header_bg"]:
                _set_cell_shading(cell, header_bg)
            _set_cell_border(cell, top=border_attrs, bottom=border_attrs,
                             left=border_attrs, right=border_attrs)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(h))
            _set_run_font(r, t["font"]["family"]["heading"], t["font"]["family"]["cn_heading"])
            r.font.size = t["font"]["size"]["body"]
            r.font.color.rgb = t["color"]["table_header_text"]
            r.font.bold = True

        for ri, row in enumerate(rows):
            _set_table_row_no_split(table.rows[ri + 1])
            cells = table.rows[ri + 1].cells
            is_alt = (ri % 2 == 1)
            for ci, val in enumerate(row):
                cell = cells[ci]
                cell.width = widths[ci] if ci < len(widths) else None
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if t["table"]["alt_rows"] and is_alt:
                    _set_cell_shading(cell, alt_bg)
                _set_cell_border(cell, top=border_attrs, bottom=border_attrs,
                                 left=border_attrs, right=border_attrs)
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(str(val))
                _set_run_font(r, t["font"]["family"]["default"], cn_font)
                r.font.size = Pt(10)
                r.font.color.rgb = t["color"]["text"]

        total_text = sum(len(str(v)) for row in rows for v in row)
        prose_table_risk = 4 <= n_rows <= 6 and total_text >= 420
        if prose_table_risk:
            for table_row in table.rows[:-1]:
                for cell in table_row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.keep_with_next = True

        after_p = self.doc.add_paragraph()
        after_p.paragraph_format.space_after = t["spacing"]["paragraph_after"]

    def _render_figure(self, node: dict):
        """Render a figure (image + caption) with automatic numbering and bookmark.

        渲染图（图片+题注）。path 为本地路径或 URL（自动下载）；label 存在时自动编号为「图 N-x」
        并插入 w:bookmarkStart/End 供 ref 节点交叉引用；支持 width 控制显示宽度、caption 自定义题注文本；
        path 缺失时仅渲染题注段落（caption-only 占位，用于交叉引用预占位）。"""

        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        path = node.get("path", "")

        if path and os.path.exists(path):
            w_in = float(node.get("width_inches") or (self.layout.content_width / 914400.0))
            assert_figure_asset_ready(
                path,
                display_width_inches=w_in,
                source_width_inches=node.get("source_width_inches"),
                declared_min_font_pt=node.get("min_text_pt"),
                contains_text=bool(node.get("contains_text") or node.get("min_text_pt") is not None),
            )
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            try:
                run.add_picture(path, width=Inches(w_in))
            except Exception as e:
                run.add_text(f"[Image: {path}]")
        elif path and (path.startswith("http://") or path.startswith("https://")):
            # Try to download remote image
            try:
                from coze_workload_identity import requests
                import tempfile
                resp = requests.get(path, timeout=10.0, allow_redirects=True)
                resp.raise_for_status()
                fd, tmp = tempfile.mkstemp(suffix=".png")
                with os.fdopen(fd, "wb") as f:
                    f.write(resp.content)
                w_in = float(node.get("width_inches") or (self.layout.content_width / 914400.0))
                assert_figure_asset_ready(
                    tmp,
                    display_width_inches=w_in,
                    source_width_inches=node.get("source_width_inches"),
                    declared_min_font_pt=node.get("min_text_pt"),
                    contains_text=bool(node.get("contains_text") or node.get("min_text_pt") is not None),
                )
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(tmp, width=Inches(w_in))
            except FigureAssetGateError:
                raise
            except Exception as e:
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                err_r = p.add_run(f"[Image unavailable: {path}]")
                _set_run_font(err_r, t["font"]["family"]["default"], cn_font)
                err_r.font.size = t["font"]["size"]["body"]
                err_r.font.color.rgb = t["color"]["muted"]

        caption = node.get("caption")
        label = node.get("label")
        if caption or label:
            self._figure_counter += 1
            num = self._figure_counter
            prefix = _caption_prefix("figure", self.lang)
            caption_text = f"{prefix}{num}"
            if caption:
                caption_text = f"{caption_text}: {caption}"
            if label:
                self._caption_labels[label] = caption_text
            cp = self.doc.add_paragraph()
            cp.alignment = t["alignment"]["caption"]
            cp.paragraph_format.space_after = t["spacing"]["paragraph_after"]
            self._write_caption_with_bookmark(cp, label, caption_text,
                                              _fig_label_id(num), t, cn_font)

    def _render_kpi_card(self, node: dict):
        """Render a KPI card: big metric number + label + delta + optional mini chart SVG.

        渲染 KPI 指标卡：大号 value 数字（主题 primary 色）+ label 标签 + delta 环比（上/下箭头染色），
        右侧可内嵌迷你 SVG 图表（bar/line/pie）。适合报告首页关键指标高亮。"""

        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        value = node.get("value", "")
        label = node.get("label", "")
        subtext = node.get("subtext", "")

        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _set_table_row_no_split(table.rows[0])
        cell = table.rows[0].cells[0]
        cell.width = Inches(2)
        border_attrs = {"val": "single", "sz": "4", "color": _rgb_to_hex(t["color"]["accent"])}
        _set_cell_border(cell, top=border_attrs, bottom=border_attrs,
                         left=border_attrs, right=border_attrs)
        _set_cell_shading(cell, _rgb_to_hex(t["color"]["code_bg"]))

        cell.text = ""
        vp = cell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        vr = vp.add_run(str(value))
        _set_run_font(vr, t["font"]["family"]["heading"], t["font"]["family"]["cn_heading"])
        vr.font.size = Pt(18)
        vr.font.bold = True
        vr.font.color.rgb = t["color"]["accent"]
        lp = cell.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(str(label))
        _set_run_font(lr, t["font"]["family"]["default"], cn_font)
        lr.font.size = Pt(10)
        lr.font.color.rgb = t["color"]["text"]
        if subtext:
            sp = cell.add_paragraph()
            sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sp.add_run(str(subtext))
            _set_run_font(sr, t["font"]["family"]["default"], cn_font)
            sr.font.size = Pt(9)
            sr.font.color.rgb = t["color"]["muted"]

    def _render_callout(self, node: dict):
        """Render a highlighted callout box with themed icon color.

        渲染提示框：按 variant（info/warning/success/danger）选择不同左边框颜色与图标，
        整体浅灰/浅黄/浅红底色；支持 title + body 双段文本。"""

        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        variant = node.get("variant", "info")
        title = node.get("title")
        body = node.get("body", "")

        bg_key = f"callout_{variant}_bg"
        bdr_key = f"callout_{variant}_border"
        bg = _rgb_to_hex(t["color"].get(bg_key, t["color"]["code_bg"]))
        bdr = _rgb_to_hex(t["color"].get(bdr_key, t["color"]["accent"]))

        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        _set_cell_shading(cell, bg)
        border_attrs = {"val": "single", "sz": "8", "color": bdr}
        _set_cell_border(cell, top=border_attrs, bottom=border_attrs,
                         left=border_attrs, right=border_attrs)

        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if title:
            tr = p.add_run(str(title) + "\n")
            _set_run_font(tr, t["font"]["family"]["heading"], t["font"]["family"]["cn_heading"])
            tr.font.size = Pt(11)
            tr.font.bold = True
            tr.font.color.rgb = t["color"]["heading"]
        # v1.4 P2-4b: callout body 支持 inline {ref}
        if node.get("inline_segments"):
            self._render_inline_segments(p, node["inline_segments"],
                                         size=t["font"]["size"]["body"],
                                         color=t["color"]["text"],
                                         font_key="default", cn_font_key="cn")
        else:
            br = p.add_run(str(body))
            _set_run_font(br, t["font"]["family"]["default"], cn_font)
            br.font.size = t["font"]["size"]["body"]
            br.font.color.rgb = t["color"]["text"]

        ap = self.doc.add_paragraph()
        ap.paragraph_format.space_after = t["spacing"]["paragraph_after"]

    def _render_page_break(self, node: dict):
        p = self.doc.add_paragraph()
        run = p.add_run()
        run.add_break(WD_BREAK.PAGE)
        self._list_counters = {}

    def _render_toc(self, node: dict):
        """插入 Word TOC 域代码（F9 刷新生成真实目录）。"""
        t = self.tokens
        levels = node.get("levels") or [1, 2, 3]
        if isinstance(levels, int):
            levels = list(range(1, levels + 1))
        lvl_range = f"{min(levels)}-{max(levels)}"
        title = node.get("title")
        if title is None:
            title = "Table of Contents" if self.lang == "en" else "目  录"

        hp = self.doc.add_paragraph()
        hp.alignment = t["alignment"]["title"]
        hr = hp.add_run(title)
        _set_run_font(hr, t["font"]["family"]["heading"], t["font"]["family"]["cn_heading"])
        hr.font.size = t["font"]["size"]["h1"]
        hr.font.bold = True
        hr.font.color.rgb = t["color"]["heading"]

        # TOC 域代码段落：fldChar begin + instrText + fldChar separate + 占位 + fldChar end
        p = self.doc.add_paragraph()
        placeholder = (
            "请右键选择\"更新域\"以生成目录"
            if self.lang != "en"
            else "Please right-click and select 'Update Field' to generate the TOC"
        )
        _add_field(p, f' TOC \\o "{lvl_range}" \\h \\z \\u ', placeholder=placeholder)
        # 提示段（小号斜体）
        tip = self.doc.add_paragraph()
        tr = tip.add_run(
            "(Press F9 or right-click 'Update Field' in Word to refresh)"
            if self.lang == "en"
            else "（在 Word 中按 F9 或右键\"更新域\"即可刷新目录）"
        )
        _set_run_font(tr, t["font"]["family"]["default"], t["font"]["family"]["cn"])
        tr.font.size = Pt(9)
        tr.font.color.rgb = t["color"]["muted"]
        tr.font.italic = True

    def _render_reference_item(self, ref: dict):
        """v1.1/v1.2 旧版独立 reference_item 节点——使用 tokens 配置的格式。"""
        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        fmt = t["reference"]["format"]
        # 向后兼容：harvard -> apa；gbt7714 -> gb7714
        if fmt == "harvard":
            fmt = "apa"
        if fmt == "gbt7714":
            fmt = "gb7714"
        self._ref_counter += 1
        text = _cite_format(ref, style=fmt, index=self._ref_counter)

        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = t["reference"]["hanging_indent"]
        p.paragraph_format.first_line_indent = -t["reference"]["hanging_indent"]
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)

        run = p.add_run(text)
        _set_run_font(run, t["font"]["family"]["default"], cn_font)
        run.font.size = t["reference"]["font_size"]
        run.font.color.rgb = t["color"]["text"]

    def _render_references_block(self, node: dict):
        """v1.4 P1-4 / P1-4b: 专业参考文献块排版。

        - 每条文献独立段落，段后 Pt(8)，段前 Pt(0)，行距 1.5 倍（学术惯例）
        - 悬挂缩进 left=Inches(0.5), first_line=-Inches(0.5)
        - 编号格式按 citation_style：apa/mla 无编号；gb7714/ieee 方括号 [N]
        - 标题段前 Pt(24)/段后 Pt(12)/h1 字号/加粗，中文学术居中，其他左对齐
        - 中文条目（含 CJK）SimSun 10.5pt；英文条目按 style 选字体：
          * apa/ieee → Times New Roman 12pt
          * mla → Times New Roman 11pt
          * 商业模式 → Calibri 11pt
        - DOI/URL 蓝色下划线（可点击）
        - 可选 group_by_type：按 article/book/thesis/conference/webpage 分组，每组带小标题
        """
        t = self.tokens
        style = (node.get("citation_style") or "apa").lower()
        style = {"harvard": "apa", "gbt7714": "gb7714"}.get(style, style)
        if style not in {"apa", "gb7714", "mla", "ieee"}:
            style = "apa"
        title = node.get("title")
        if title is None:
            title = "References" if self.lang == "en" else "参考文献"
        items = node.get("items", []) or []
        group_by_type = bool(node.get("group_by_type"))
        # 判断模式：auto_structure 时按 auto_mode；否则默认 academic
        is_business = (self._auto_mode == "business")
        is_cn = (self._auto_lang == "cn") or (self.lang in ("cn", "zh"))

        # 中文 vs 英文字体
        cn_font_heading = t["font"]["family"]["cn_heading"]
        en_font_heading = t["font"]["family"]["heading"]
        cn_font_body = "SimSun"
        if style == "mla":
            en_font_body = "Times New Roman"; en_body_pt = 11
        elif is_business:
            en_font_body = "Calibri"; en_body_pt = 11
        else:
            en_font_body = "Times New Roman"; en_body_pt = 12
        cn_body_pt = 10.5  # 五号字

        # 标题对齐：中文学术居中，其他左对齐
        title_align = WD_ALIGN_PARAGRAPH.CENTER if (is_cn and not is_business) else WD_ALIGN_PARAGRAPH.LEFT

        # 标题段
        hp = self.doc.add_paragraph()
        hp.alignment = title_align
        hp.paragraph_format.space_before = Pt(24)
        hp.paragraph_format.space_after = Pt(12)
        hr = hp.add_run(title)
        _set_run_font(hr, en_font_heading, cn_font_heading)
        hr.font.size = t["font"]["size"]["h1"]
        hr.font.bold = True
        hr.font.color.rgb = t["color"]["heading"]

        # 重置条目计数
        self._ref_style = style

        # 是否编号
        numbered = style in ("gb7714", "ieee")

        def _render_group(sub_items: list, start_idx: int, subtitle: str | None = None):
            if subtitle is not None:
                sp = self.doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(12)
                sp.paragraph_format.space_after = Pt(4)
                sr = sp.add_run(subtitle)
                _set_run_font(sr, en_font_heading, cn_font_heading)
                sr.font.size = t["font"]["size"]["h3"]
                sr.font.bold = True
                sr.font.color.rgb = t["color"]["heading"]
            for j, ref in enumerate(sub_items):
                idx = start_idx + j
                text = _cite_format(ref, style=style, index=idx if numbered else None)
                self._render_reference_entry(text, ref, numbered, idx,
                                             cn_font_body, en_font_body, cn_body_pt, en_body_pt,
                                             heading_font=en_font_heading, cn_heading_font=cn_font_heading)

        if group_by_type:
            # 按 type 分组（保持输入顺序内的组序）
            type_order = ["article", "book", "thesis", "conference", "webpage", "report", "standard", "patent"]
            type_title = {"article": "期刊论文" if is_cn else "Journal Articles",
                          "book": "专著" if is_cn else "Books",
                          "thesis": "学位论文" if is_cn else "Theses",
                          "conference": "会议论文" if is_cn else "Conference Papers",
                          "webpage": "网络资源" if is_cn else "Web Resources",
                          "report": "报告" if is_cn else "Reports",
                          "standard": "标准" if is_cn else "Standards",
                          "patent": "专利" if is_cn else "Patents"}
            groups: dict[str, list] = {}
            ungrouped = []
            for it in items:
                rt = (it.get("type") or "article").lower() if isinstance(it, dict) else "article"
                if rt in type_order:
                    groups.setdefault(rt, []).append(it)
                else:
                    ungrouped.append(it)
            cursor = 1
            for rt in type_order:
                if rt in groups:
                    _render_group(groups[rt], cursor, type_title[rt])
                    cursor += len(groups[rt])
            if ungrouped:
                _render_group(ungrouped, cursor, None)
        else:
            _render_group(items, 1, None)

    # ---- Reference entry helpers ----

    _RE_URL = None  # lazy compile

    def _render_reference_entry(self, text, ref, numbered, idx,
                                cn_font_body, en_font_body, cn_body_pt, en_body_pt,
                                heading_font, cn_heading_font):
        """渲染单条参考文献：悬挂缩进 + DOI/URL 蓝色下划线 + 中英文字体自动切换。"""
        # Lazy import to avoid circular - try multiple import strategies
        _contains_cjk = None
        try:
            from ..shared.citation import _contains_cjk  # type: ignore
        except Exception:
            try:
                from pro_docx_gen.shared.citation import _contains_cjk  # type: ignore
            except Exception:
                import sys as _s, os as _o
                _h = _o.path.dirname(_o.path.abspath(__file__))
                _r = _o.path.normpath(_o.path.join(_h, "..", "..", ".."))
                if _r not in _s.path: _s.path.insert(0, _r)
                from shared.citation import _contains_cjk  # type: ignore
        t = self.tokens
        is_cjk = _contains_cjk(text)
        body_font = cn_font_body if is_cjk else en_font_body
        body_pt = cn_body_pt if is_cjk else en_body_pt

        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = -Inches(0.5)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)

        # 编号前缀
        if numbered:
            prefix = f"[{idx}] "
            pr = p.add_run(prefix)
            _set_run_font(pr, body_font, cn_font_body)
            pr.font.size = Pt(body_pt)
            pr.font.color.rgb = t["color"]["text"]

        # 拆分 URL/DOI 段，蓝色下划线
        self._write_mixed_ref_text(p, text, body_font, cn_font_body, body_pt,
                                   t["color"]["text"])

    def _write_mixed_ref_text(self, p, text: str, latin_font: str, cn_font: str,
                              pt_size: float, base_color):
        """将参考文献条目文本按 URL/DOI 拆分为普通段与蓝色下划线段，分段写入 run。"""
        import re
        if self._RE_URL is None:
            # 匹配 http(s)://... 或 doi:10.... 直到空白/中文标点/结束
            DocxRenderer._RE_URL = re.compile(
                r'(https?://\S+|doi:\s*10\.\S+|10\.\d{4,9}/\S+)',
                re.IGNORECASE,
            )
        url_re = self._RE_URL
        pos = 0
        link_color = RGBColor(0x05, 0x63, 0xC1)
        for m in url_re.finditer(text):
            if m.start() > pos:
                seg = text[pos:m.start()]
                r = p.add_run(seg)
                _set_run_font(r, latin_font, cn_font)
                r.font.size = Pt(pt_size)
                r.font.color.rgb = base_color
            link_text = m.group(0)
            # 去掉 link 末尾可能粘连的句号/逗号/分号
            trailing_punc = ""
            while link_text and link_text[-1] in ".,;:。，；：）)]」』】》":
                trailing_punc = link_text[-1] + trailing_punc
                link_text = link_text[:-1]
            rl = p.add_run(link_text)
            _set_run_font(rl, latin_font, cn_font)
            rl.font.size = Pt(pt_size)
            rl.font.color.rgb = link_color
            rl.font.underline = True
            if trailing_punc:
                rp = p.add_run(trailing_punc)
                _set_run_font(rp, latin_font, cn_font)
                rp.font.size = Pt(pt_size)
                rp.font.color.rgb = base_color
            pos = m.end()
        if pos < len(text):
            seg = text[pos:]
            r = p.add_run(seg)
            _set_run_font(r, latin_font, cn_font)
            r.font.size = Pt(pt_size)
            r.font.color.rgb = base_color

    @staticmethod
    def _format_reference(ref: dict, fmt: str) -> str:
        """v1.1/v1.2 兼容入口：内部转调 citation 引擎。"""
        style = {"harvard": "apa", "gbt7714": "gb7714"}.get(fmt, fmt or "apa")
        return _cite_format(ref, style=style)

    def _render_chart(self, node: dict):
        """Render a chart (bar/line/pie/area/scatter/radar) as an inline image.

        渲染图表：用 matplotlib 生成 PNG 图片后以 add_picture 内联插入，支持 theme palette 自动配色；
        图表类型与 PPT 侧共享；生成后若 chart_caption=True 自动加题注编号。matplotlib 中文使用 SimHei/
        Microsoft YaHei 回退。"""

        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        self._figure_counter += 1
        fig_idx = self._figure_counter

        if self._chart_tmp_dir is None:
            self._chart_tmp_dir = _WritableTempDir("pro_docx_charts_", self._output_dir)
        assets_dir = self._chart_tmp_dir.name
        png_path = os.path.join(assets_dir, f"chart_{fig_idx:03d}.png")
        # v1.6.6: page-fit contract. Compute remaining page height before
        # rendering so chart_renderer can auto-shrink a too-tall chart. This
        # is the root-cause fix for the "blank page after a tall chart"
        # regression: the old code just trusted chart_renderer to fit and
        # never measured the available page area.
        compact_mode_flag = bool(node.get("compact_mode", False))
        available_h = compute_available_height_inches(self.tokens)
        try:
            render_chart_to_png(
                node,
                t,
                png_path,
                available_height_inches=available_h,
                compact_mode=compact_mode_flag,
            )
        except Exception as e:
            err_p = self.doc.add_paragraph()
            err_r = err_p.add_run(f"[Chart rendering failed: {e}]")
            _set_run_font(err_r, t["font"]["family"]["default"], cn_font)
            err_r.font.size = Pt(10)
            err_r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            return
        # chart_renderer writes runtime metadata back onto the spec dict so
        # we can pick the correct paragraph spacing/keep_with_next.
        compact_now = bool(node.get("_compact_mode", compact_mode_flag))
        # Trust the figure's *actual* width (post-shrink) over width_pct so
        # the inserted picture matches the PNG we just rendered. The chart
        # spec stores inches; the layout.content_width is EMU.
        rendered_width_in = float(node.get("_figure_width_in") or 0)
        if rendered_width_in <= 0:
            width_pct = float(node.get("width_pct", 1.0))
            width_pct = max(0.3, min(1.0, width_pct))
            rendered_width_in = (self.layout.content_width / 914400.0) * width_pct
        display_width_inches = rendered_width_in
        # Recompute aspect-correct native EMU width for the final insert.
        # We deliberately do NOT pass a fixed width when the renderer has
        # already shrunk the figure — that lets python-docx use the PNG's
        # intrinsic aspect ratio at the height-equivalent width.
        aspect_picture_width_emu = int(self.layout.content_width * (
            rendered_width_in / max(self.layout.content_width / 914400.0, 0.01)
        ))
        # Final sanity check on the rendered file.
        assert_figure_asset_ready(
            png_path,
            display_width_inches=display_width_inches,
            source_width_inches=display_width_inches,
            declared_min_font_pt=9.0,
            contains_text=True,
        )
        align_key = node.get("align", "center")
        p = self.doc.add_paragraph()
        if align_key == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif align_key == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # v1.6.6: compact_mode ⇒ zero space around the figure + tight
        # keep_with_next on the caption, so dense chart walls stay glued
        # to the surrounding text. Default behaviour is preserved.
        if compact_now:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
        else:
            p.paragraph_format.space_before = Pt(6)
        run = p.add_run()
        try:
            run.add_picture(png_path, width=int(aspect_picture_width_emu))
        except Exception as e:
            run.add_text(f"[Chart image: {png_path}] ({e})")

        title = node.get("title")
        caption = node.get("caption")
        theme_label = "Figure" if self.lang == "en" else "图"
        caption_label = f"{theme_label} {fig_idx}"
        if title:
            caption_text = f"{caption_label}: {title}"
        else:
            caption_text = caption_label
        if caption:
            caption_text = f"{caption_text}  {caption}"

        cp = self.doc.add_paragraph()
        cp.alignment = t["alignment"]["caption"]
        # v1.6.6: tight spacing in compact_mode; default still uses the
        # theme's paragraph_after so the chart reads like any other figure.
        if compact_now:
            cp.paragraph_format.space_before = Pt(0)
            cp.paragraph_format.space_after = Pt(0)
            cp.paragraph_format.keep_with_next = True
        else:
            cp.paragraph_format.space_after = t["spacing"]["paragraph_after"]
        cr = cp.add_run(caption_text)
        _set_run_font(cr, t["font"]["family"]["default"], cn_font)
        cr.font.size = t["font"]["size"].get("chart_caption", Pt(9))
        cr.font.color.rgb = t["color"]["muted"]
        cr.font.italic = True

    # ─── v1.2: 修订追踪 ──────────────────────────────────────────

    def _render_revision_inline(self, node: dict):
        """独立 revision 节点 → 在新段落中渲染修订文本。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = self.tokens["spacing"]["first_line_indent"]
        p.paragraph_format.line_spacing = self.tokens["spacing"]["line_spacing"]
        self._render_revision_run(p, node)

    def _render_revision_run(self, paragraph, r: dict):
        """在现有 paragraph 中追加修订 run。"""
        action = r.get("action", "insert")
        author = r.get("author") or self.tokens["revision"]["author"]
        date = r.get("date") or datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        insert_color = _rgb_to_hex(self.tokens["color"]["revision_insert"])

        if action == "insert":
            self._add_insertion(paragraph, r.get("text", ""), author, date)
        elif action == "delete":
            self._add_deletion(paragraph, r.get("text", ""), author, date)
        elif action == "replace":
            self._add_deletion(paragraph, r.get("old_text", ""), author, date)
            self._add_insertion(paragraph, r.get("new_text", ""), author, date)

    def _add_insertion(self, paragraph, text: str, author: str, date: str):
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), self._next_revision_id())
        ins.set(qn("w:author"), author)
        ins.set(qn("w:date"), date)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), _rgb_to_hex(self.tokens["color"]["revision_insert"]))
        rpr.append(color)
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        rpr.append(u)
        # 中文字体
        rfonts = OxmlElement("w:rFonts")
        fn = self.tokens["font"]["family"]["default"]
        rfonts.set(qn("w:ascii"), fn); rfonts.set(qn("w:hAnsi"), fn)
        rfonts.set(qn("w:eastAsia"), self.tokens["font"]["family"]["cn"])
        rpr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(self.tokens["font"]["size"]["body"].pt * 2)))
        rpr.append(sz)
        run.append(rpr)
        t_el = OxmlElement("w:t")
        t_el.text = text
        t_el.set(qn("xml:space"), "preserve")
        run.append(t_el)
        ins.append(run)
        paragraph._p.append(ins)

    def _add_deletion(self, paragraph, text: str, author: str, date: str):
        # w:del 包裹 w:r > w:delText
        delete = OxmlElement("w:del")
        delete.set(qn("w:id"), self._next_revision_id())
        delete.set(qn("w:author"), author)
        delete.set(qn("w:date"), date)
        run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), _rgb_to_hex(self.tokens["color"]["revision_delete"]))
        rpr.append(color)
        strike = OxmlElement("w:strike")
        strike.set(qn("w:val"), "true")
        rpr.append(strike)
        rfonts = OxmlElement("w:rFonts")
        fn = self.tokens["font"]["family"]["default"]
        rfonts.set(qn("w:ascii"), fn); rfonts.set(qn("w:hAnsi"), fn)
        rfonts.set(qn("w:eastAsia"), self.tokens["font"]["family"]["cn"])
        rpr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(self.tokens["font"]["size"]["body"].pt * 2)))
        rpr.append(sz)
        run.append(rpr)
        dt = OxmlElement("w:delText")
        dt.text = text
        dt.set(qn("xml:space"), "preserve")
        run.append(dt)
        delete.append(run)
        paragraph._p.append(delete)

    # ─── v1.2: 批注 ──────────────────────────────────────────────

    def _render_comment_block(self, node: dict):
        """独立 comment 节点（对刚结束段落的整段批注）——实现为一个空段落后附 comment reference。"""
        p = self.doc.add_paragraph()
        # 插入不可见 anchor，指向该段整体
        self._attach_comment(p, node)

    def _attach_comment(self, paragraph, comment_data: dict):
        """给段落附加批注（整段范围）。"""
        cid = len(self._comments)
        author = comment_data.get("author") or self.tokens["comment"]["author"]
        date = comment_data.get("date") or datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
        text = comment_data.get("text", "")
        self._comments.append({"id": cid, "author": author, "date": date, "text": text})

        p = paragraph._p
        # commentRangeStart 放在 pPr 之后
        crs = OxmlElement("w:commentRangeStart")
        crs.set(qn("w:id"), str(cid))
        p.insert(1 if len(p) > 0 and p[0].tag.endswith('}pPr') else 0, crs)
        # commentRangeEnd + commentReference 放在末尾
        cre = OxmlElement("w:commentRangeEnd")
        cre.set(qn("w:id"), str(cid))
        p.append(cre)
        # reference run
        ref_run = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rstyle = OxmlElement("w:rStyle")
        rstyle.set(qn("w:val"), "CommentReference")
        rpr.append(rstyle)
        ref_run.append(rpr)
        cref = OxmlElement("w:commentReference")
        cref.set(qn("w:id"), str(cid))
        ref_run.append(cref)
        p.append(ref_run)

    def _inject_comments_part(self):
        if not self._comments:
            return
        from lxml import etree
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        NS = {"w": W_NS}
        root = etree.Element(f"{{{W_NS}}}comments", nsmap={"w": W_NS})
        for c in self._comments:
            ce = etree.SubElement(root, f"{{{W_NS}}}comment")
            ce.set(f"{{{W_NS}}}id", str(c["id"]))
            ce.set(f"{{{W_NS}}}author", c["author"])
            ce.set(f"{{{W_NS}}}date", c["date"])
            ce.set(f"{{{W_NS}}}initials", c["author"][:2] if c["author"] else "PD")
            p = etree.SubElement(ce, f"{{{W_NS}}}p")
            r = etree.SubElement(p, f"{{{W_NS}}}r")
            t = etree.SubElement(r, f"{{{W_NS}}}t")
            t.text = c["text"]
        blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        partname = PackURI("/word/comments.xml")
        content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
        part = Part(partname, content_type, blob, self.doc.part.package)
        self.doc.part.relate_to(part, RT.COMMENTS)

    # ─── v1.2: 脚注/尾注 ─────────────────────────────────────────

    def _render_footnote_ref(self, node: dict):
        """独立 footnote 节点：在新段落末尾插入上标脚注引用，并登记脚注。"""
        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        p = self.doc.add_paragraph()
        # 没有 anchor 文本时，显示 [N]
        fid = len(self._footnotes) + 1  # id=0/1 是保留分隔符，从 2 开始
        real_id = fid + 1
        self._footnotes.append({"id": real_id, "text": node["text"]})
        # 引用标记
        run = p.add_run(f"[{fid}] ")
        _set_run_font(run, t["font"]["family"]["default"], cn_font)
        run.font.size = Pt(9)
        run.font.superscript = True
        run.font.color.rgb = t["color"]["accent"]

    def _render_footnote_ref_inline(self, paragraph, r: dict):
        t = self.tokens
        fid = len(self._footnotes) + 1
        real_id = fid + 1
        self._footnotes.append({"id": real_id, "text": r.get("text", "")})
        run = paragraph.add_run(str(fid))
        _set_run_font(run, t["font"]["family"]["default"], t["font"]["family"]["cn"])
        run.font.size = t["font"]["size"]["footnote"]
        run.font.superscript = True
        run.font.color.rgb = t["color"]["accent"]

    def _render_endnote_ref(self, node: dict):
        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        p = self.doc.add_paragraph()
        eid = len(self._endnotes) + 1
        real_id = eid + 1
        self._endnotes.append({"id": real_id, "text": node["text"]})
        run = p.add_run(f"[{eid}] ")
        _set_run_font(run, t["font"]["family"]["default"], cn_font)
        run.font.size = Pt(9)
        run.font.superscript = True
        run.font.color.rgb = t["color"]["accent"]

    def _append_endnotes_section(self):
        """在文档末尾渲染尾注列表（作为普通段落，不注入 endnotes part，简化实现）。"""
        if not self._endnotes:
            return
        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        self.doc.add_page_break()
        h = self.doc.add_paragraph(style="Heading 1")
        hr = h.add_run("尾注" if self.lang == "cn" else "Endnotes")
        _set_run_font(hr, t["font"]["family"]["heading"], t["font"]["family"]["cn_heading"])
        for i, en in enumerate(self._endnotes, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.first_line_indent = -Inches(0.3)
            r = p.add_run(f"[{i}] {en['text']}")
            _set_run_font(r, t["font"]["family"]["default"], cn_font)
            r.font.size = t["font"]["size"]["footnote"]
            r.font.color.rgb = t["color"]["text"]

    def _inject_footnotes_part(self):
        """注入 word/footnotes.xml part 及对应 content type / 关系。"""
        if not self._footnotes:
            return
        self._inject_notes_part("footnotes", self._footnotes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
            "/word/footnotes.xml", RT.FOOTNOTES)

    def _inject_endnotes_part(self):
        if not self._endnotes:
            return
        self._inject_notes_part("endnotes", self._endnotes,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml",
            "/word/endnotes.xml", RT.ENDNOTES)

    def _inject_notes_part(self, tag: str, notes: list, content_type: str, part_uri: str, reltype: str):
        from lxml import etree
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        root = etree.Element(f"{{{W_NS}}}{tag}", nsmap={"w": W_NS})
        # 保留 id=0 (separator) 和 id=1 (continuationSeparator)
        for rid, sep_type in [(0, "separator"), (1, "continuationSeparator")]:
            fn = etree.SubElement(root, f"{{{W_NS}}}{tag[:-1]}")
            fn.set(f"{{{W_NS}}}id", str(rid))
            fn.set(f"{{{W_NS}}}type", sep_type)
            p = etree.SubElement(fn, f"{{{W_NS}}}p")
        for n in notes:
            fn = etree.SubElement(root, f"{{{W_NS}}}{tag[:-1]}")
            fn.set(f"{{{W_NS}}}id", str(n["id"]))
            p = etree.SubElement(fn, f"{{{W_NS}}}p")
            r = etree.SubElement(p, f"{{{W_NS}}}r")
            t = etree.SubElement(r, f"{{{W_NS}}}t")
            t.text = n["text"]
        blob = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
        partname = PackURI(part_uri)
        part = Part(partname, content_type, blob, self.doc.part.package)
        self.doc.part.relate_to(part, reltype)

    # ─── v1.3: 原生 OMML 公式（DOCX-P0-1） ─────────────────────

    def _render_equation(self, node: dict):
        """Render an OMML equation (LaTeX → OMML via latex2mathml + XSLT) with optional caption.

        渲染可编辑 Word 公式：通过 latex2mathml 将 LaTeX 源码转为 MathML，再经 MML2OMML.XSL 转为 OMML
        嵌入段落；display=block 时居中独占一行并自动编号「(N)」，display=inline 时嵌入正文。
        label 存在时插入 bookmark 供 ref 节点交叉引用。"""

        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        latex = node.get("latex", "")
        display = node.get("display", "block")
        caption = node.get("caption")
        label = node.get("label")
        caption_only = bool(node.get("_caption_only"))
        # v1.4 P2-4: auto-number for equation if caption/label provided
        eq_num_text = None
        if caption or label:
            self._equation_counter += 1
            num = self._equation_counter
            prefix = _caption_prefix("equation", self.lang)
            eq_num_text = f"({num})"
            if label:
                # Label target includes number text; REF shows "(N)" by default
                self._caption_labels[label] = eq_num_text

        # v1.4 P3-9b: caption-only 模式（由 auto_structure 将"公式5 ..."行转为 equation 节点）
        if caption_only or not latex:
            cp = self.doc.add_paragraph()
            cp.alignment = t["alignment"]["caption"]
            if eq_num_text:
                nr = cp.add_run(eq_num_text + (" " if caption else ""))
                _set_run_font(nr, t["font"]["family"]["default"], cn_font)
                nr.font.size = t["font"]["size"]["caption"]
                nr.font.bold = True
            if caption:
                cr = cp.add_run(caption)
                _set_run_font(cr, t["font"]["family"]["default"], cn_font)
                cr.font.size = t["font"]["size"]["caption"]
                cr.font.italic = False
            if label and (eq_num_text or caption):
                self._add_bookmark(cp, label)
            return

        try:
            from .omml import append_equation_to_paragraph
            p = self.doc.add_paragraph()
            caption_text = caption
            # Append auto-number to caption text for OMML path
            if eq_num_text:
                caption_text = f"{eq_num_text}" + (f" {caption}" if caption else "")
            append_equation_to_paragraph(
                p, latex, display=display, caption=caption_text,
                cn_font=cn_font,
                caption_size_pt=self._pt_value(t["font"]["size"]["caption"]),
            )
            # Wrap bookmark around the caption if label is set
            if label and caption_text:
                # The last paragraph added by append_equation_to_paragraph should be the caption;
                # We find it and wrap a bookmark around the number run. Simpler: add our own
                # bookmark by re-opening the last paragraph in document body.
                cap_p = self.doc.paragraphs[-1]
                self._add_bookmark(cap_p, label)
        except Exception as e:
            # Fallback: plain text
            p = self.doc.add_paragraph()
            run = p.add_run(latex)
            _set_run_font(run, t["font"]["family"]["code"], t["font"]["family"]["code"])
            run.font.size = t["font"]["size"]["body"]
            run.font.italic = True
            if eq_num_text or caption:
                cp = self.doc.add_paragraph()
                if eq_num_text:
                    nr = cp.add_run(eq_num_text + " ")
                    _set_run_font(nr, t["font"]["family"]["default"], cn_font)
                    nr.font.size = t["font"]["size"]["caption"]
                    nr.font.bold = True
                    if label:
                        self._add_bookmark(cp, label)
                if caption:
                    cr = cp.add_run(caption)
                    _set_run_font(cr, t["font"]["family"]["default"], cn_font)
                    cr.font.size = t["font"]["size"]["caption"]
                    cr.font.italic = True

    # ─── v1.3: SVG 形状（DOCX-P0-1 / svg_shape 节点） ─────────────

    def _render_svg_shape(self, node: dict):
        svg = node.get("svg", "")
        if not svg:
            return
        width_spec = node.get("width", "5cm")
        align = node.get("align", "center")

        # Compute target EMU width/height
        from docx.shared import Cm, Inches, Emu
        def _to_emu(v) -> int:
            if isinstance(v, (int, float)):
                return int(v)
            s = str(v).strip()
            if s.endswith("%"):
                # content width percentage
                pct = float(s[:-1]) / 100.0
                # assume content width ~6 inches usable
                return int(Inches(6) * pct)
            if s.endswith("cm"):
                return int(Cm(float(s[:-2])))
            if s.endswith("mm"):
                return int(Cm(float(s[:-2])/10))
            if s.endswith("in") or s.endswith('"'):
                return int(Inches(float(s.rstrip('"'))))
            if s.endswith("pt"):
                return int(float(s[:-2]) * 12700)
            try:
                return int(float(s))
            except Exception:
                return int(Cm(5))
        w_emu = _to_emu(width_spec)
        # default 1:1 aspect ratio if none in SVG; else read from viewBox
        h_emu = w_emu
        import re
        m = re.search(r'viewBox\s*=\s*"([^"]+)"', svg)
        if m:
            parts = [float(x) for x in re.findall(r"[-+]?\d*\.?\d+", m.group(1))]
            if len(parts) == 4 and parts[2] > 0 and parts[3] > 0:
                h_emu = int(w_emu * parts[3] / parts[2])

        try:
            from ..shared.svg_engine import svg_to_docx_drawing
        except ImportError:
            try:
                from pro_docx_gen.shared.svg_engine import svg_to_docx_drawing
            except ImportError:
                try:
                    from skills.shared.svg_engine import svg_to_docx_drawing
                except ImportError:
                    p = self.doc.add_paragraph(); p.add_run("[SVG shape]").font.italic = True
                    return
        drawing = svg_to_docx_drawing(svg, w_emu, h_emu, shape_id=id(svg)&0x7FFFFFFF)

        p = self.doc.add_paragraph()
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Ensure the run exists and append drawing to the paragraph
        run = p.add_run()
        run._r.append(drawing)

    # ─── v1.2: 签名区 ────────────────────────────────────────────

    def _render_signature_block(self, node: dict):
        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        title = node.get("title", "签字")
        name = node.get("name", "")
        date = node.get("date", "")

        self.doc.add_paragraph()  # 空行
        # 标题
        hp = self.doc.add_paragraph()
        hr = hp.add_run(title)
        _set_run_font(hr, t["font"]["family"]["heading"], t["font"]["family"]["cn_heading"])
        hr.font.size = Pt(12); hr.font.bold = True
        hr.font.color.rgb = t["color"]["heading"]
        # 表格布局：签字人/日期/横线
        table = self.doc.add_table(rows=2, cols=2)
        table.autofit = True
        labels = [f"签字人（{name}）：" if name else "签字人：",
                  f"日期：{date}" if date else "日期："]
        for i, lab in enumerate(labels):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(lab)
            _set_run_font(r, t["font"]["family"]["default"], cn_font)
            r.font.size = t["font"]["size"]["body"]
            r.font.color.rgb = t["color"]["text"]
            # 横线行
            sig_cell = table.rows[1].cells[i]
            sig_cell.text = ""
            sp = sig_cell.paragraphs[0]
            sr = sp.add_run("____________________________")
            _set_run_font(sr, t["font"]["family"]["default"], cn_font)
            sr.font.size = t["font"]["size"]["body"]
        self.doc.add_paragraph()

    def _render_signature_line(self, node: dict):
        """Render a signature/approval line with name, title, date placeholders.

        渲染签名/审批线：在同一段落中横向排列「签字：______」「职务：______」「日期：______」
        等占位位，通过制表符与下划线实现；用于合同、审批单、报告落款。"""

        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        signer = node.get("signer", "签字人")
        date_label = node.get("date", "日期")
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        r1 = p.add_run(f"{signer}：____________________    ")
        _set_run_font(r1, t["font"]["family"]["default"], cn_font)
        r1.font.size = t["font"]["size"]["body"]
        r2 = p.add_run(f"{date_label}：____________________")
        _set_run_font(r2, t["font"]["family"]["default"], cn_font)
        r2.font.size = t["font"]["size"]["body"]

    def _bookmark_id_next(self) -> int:
        self._bookmark_id_counter = getattr(self, "_bookmark_id_counter", 1000) + 1
        return self._bookmark_id_counter

    def _add_bookmark(self, paragraph, name: str, bm_id: int | None = None):
        """Wrap a w:bookmarkStart/w:bookmarkEnd around ``paragraph``.

        Must be called AFTER runs are added (start is prepended, end appended).
        """
        if not name:
            return
        bid = bm_id if bm_id is not None else self._bookmark_id_next()
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(int(bid)))
        start.set(qn("w:name"), name)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(int(bid)))
        # Insert start at beginning, end at end
        p_el = paragraph._p
        p_el.insert(0, start)
        p_el.append(end)

    def _write_caption_with_bookmark(self, paragraph, label: str | None,
                                     text: str, bm_name: str, t: dict, cn_font: str):
        """Write caption text with bold number, and wrap bookmarkStart/End around
        the whole caption paragraph if ``label`` is given."""
        # Caption format: "图 1: 描述文字" (prefix bold, rest normal)
        # Split into "prefix+number+colon" and "rest" if there's a ': '
        run = paragraph.add_run(text)
        _set_run_font(run, t["font"]["family"]["default"], cn_font)
        run.font.size = t["font"]["size"]["caption"]
        run.font.color.rgb = t["color"]["muted"]
        # bold the label prefix (up to first ':' if present)
        if ":" in text:
            # rebuild: bold prefix, normal rest
            run.text = ""
            pre, _, post = text.partition(":")
            r1 = paragraph.add_run(pre + ":")
            _set_run_font(r1, t["font"]["family"]["default"], cn_font)
            r1.font.size = t["font"]["size"]["caption"]
            r1.font.color.rgb = t["color"]["muted"]
            r1.font.bold = True
            if post.strip():
                r2 = paragraph.add_run(post)
                _set_run_font(r2, t["font"]["family"]["default"], cn_font)
                r2.font.size = t["font"]["size"]["caption"]
                r2.font.color.rgb = t["color"]["muted"]
        else:
            run.font.bold = True
        if label:
            # Use the user-supplied label as bookmark name for REF target
            self._add_bookmark(paragraph, label)

    def _add_ref_field(self, paragraph, target: str, placeholder: str = ""):
        """Insert a REF field: fldChar begin + instrText + separate + placeholder + end."""
        # begin
        fld_begin = OxmlElement("w:r")
        fChar1 = OxmlElement("w:fldChar")
        fChar1.set(qn("w:fldCharType"), "begin")
        fld_begin.append(fChar1)
        paragraph._p.append(fld_begin)
        # instrText
        instr_r = OxmlElement("w:r")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f" REF {target} \\h "
        instr_r.append(instr)
        paragraph._p.append(instr_r)
        # separate
        sep_r = OxmlElement("w:r")
        fChar2 = OxmlElement("w:fldChar")
        fChar2.set(qn("w:fldCharType"), "separate")
        sep_r.append(fChar2)
        paragraph._p.append(sep_r)
        # placeholder text (used before Word F9 refresh)
        ph = placeholder or self._caption_labels.get(target, target)
        text_r = paragraph.add_run(ph)
        # end
        end_r = OxmlElement("w:r")
        fChar3 = OxmlElement("w:fldChar")
        fChar3.set(qn("w:fldCharType"), "end")
        end_r.append(fChar3)
        paragraph._p.append(end_r)
        return text_r

    def _render_ref(self, node: dict):
        t = self.tokens
        cn_font = t["font"]["family"]["cn"]
        target = node.get("target", "")
        prefix = node.get("prefix", "")
        suffix = node.get("suffix", "")
        p = self.doc.add_paragraph() if node.get("style") != "inline" else None
        # Inline refs are tricky because they need to be inserted at the current
        # point; but since nodes are flattened we can't easily splice back into
        # the previous paragraph. We render refs as their own paragraph for
        # simplicity (which works for the common "如图 X 所示" usage inside a
        # paragraph via Markdown parsing — LLM outputs refs in dedicated runs).
        para = p or self.doc.add_paragraph()
        para.paragraph_format.space_after = t["spacing"]["paragraph_after"]
        if prefix:
            r = para.add_run(prefix)
            _set_run_font(r, t["font"]["family"]["default"], cn_font)
            r.font.size = t["font"]["size"]["body"]
            r.font.color.rgb = t["color"]["text"]
        self._add_ref_field(para, target)
        if suffix:
            r = para.add_run(suffix)
            _set_run_font(r, t["font"]["family"]["default"], cn_font)
            r.font.size = t["font"]["size"]["body"]
            r.font.color.rgb = t["color"]["text"]

    # v1.4 P2-4b: 在已有段落内渲染 inline_segments（text/ref 交替）
    def _render_inline_segments(self, paragraph, segments: list[dict], *,
                                size=None, color=None, bold=False, italic=False,
                                font_key="default", cn_font_key="cn",
                                font_latin=None, cn_font=None):
        """在 ``paragraph`` 内追加 inline_segments 列表（text / ref 交替）。

        每个 text 段作为普通 run 渲染；每个 ref 段通过 ``_add_ref_field`` 在
        当前段落中直接插入 REF 域代码（不新建段落），并对占位 run 施加相同字体样式。
        """
        t = self.tokens
        body_size = size if size is not None else t["font"]["size"]["body"]
        body_color = color if color is not None else t["color"]["text"]
        # v1.4 P3-9c: 支持显式传 latin/cn 字体；否则按 font_key/cn_font_key 查 tokens
        if font_latin is None:
            font_latin = t["font"]["family"].get(font_key, t["font"]["family"]["default"])
        if cn_font is None:
            cn_font = t["font"]["family"].get(cn_font_key, t["font"]["family"]["cn"])
        for seg in segments:
            stype = seg.get("type", "text")
            if stype == "ref":
                target = seg.get("target", "")
                # 插入 REF 域代码，并对 placeholder run 设置正确字体
                ph_run = self._add_ref_field(paragraph, target)
                try:
                    _set_run_font(ph_run, font_latin, cn_font)
                    ph_run.font.size = Pt(body_size) if not hasattr(body_size, "pt") else body_size
                    ph_run.font.color.rgb = body_color
                    if bold:
                        ph_run.font.bold = True
                    if italic:
                        ph_run.font.italic = True
                except Exception:
                    pass
            else:
                txt = seg.get("text", "") or ""
                if txt:
                    r = paragraph.add_run(txt)
                    _set_run_font(r, font_latin, cn_font)
                    r.font.size = Pt(body_size) if not hasattr(body_size, "pt") else body_size
                    r.font.color.rgb = body_color
                    if bold:
                        r.font.bold = True
                    if italic:
                        r.font.italic = True

# ----------------------------------------------------------------------
# Module-level helpers (outside DocxRenderer class)
# ----------------------------------------------------------------------

def _ordered_prefix(level: int, counter: int) -> str:
    """Fallback manual prefix for ordered lists (kept for safety)."""
    if level == 0:
        return f"{counter}. "
    if level == 1:
        return f"{_to_alpha(counter)}. "
    if level == 2:
        return f"{_to_roman(counter)}. "
    return f"{_to_roman(counter)}. "


def _bullet_prefix(level: int) -> str:
    bullets = ["\u2022", "\u25e6", "\u25aa", "\u25ab"]
    return f"{bullets[min(level, len(bullets) - 1)]} "


def _bullet_char(level: int) -> str:
    """Return Unicode bullet per level (matching Word defaults)."""
    return ["\u2022", "\u25e6", "\u25aa", "\u25ab"][min(level, 3)]


def _level_format(ordered: bool, level: int) -> str:
    if not ordered:
        return "bullet"
    return ["decimal", "lowerLetter", "lowerRoman", "lowerRoman"][min(level, 3)]


def _level_lvlText(ordered: bool, level: int) -> str:
    if not ordered:
        return _bullet_char(level)
    return ["%1.", "%2.", "%3.", "%3."][min(level, 3)]


def _to_alpha(n: int) -> str:
    """1 -> a, 2 -> b ... 27 -> aa."""
    s = ""
    while n > 0:
        n, rem = divmod(n - 1, 26)
        s = chr(ord("a") + rem) + s
    return s


def _to_roman(n: int) -> str:
    vals = [(1000, "m"), (900, "cm"), (500, "d"), (400, "cd"),
            (100, "c"), (90, "xc"), (50, "l"), (40, "xl"),
            (10, "x"), (9, "ix"), (5, "v"), (4, "iv"), (1, "i")]
    out = ""
    for v, sym in vals:
        while n >= v:
            out += sym
            n -= v
    return out

def _caption_prefix(kind: str, lang: str) -> str:
    """Return caption label prefix like '图 ' / 'Figure '."""
    cn = (lang or "en").lower() in ("cn", "zh", "zh-cn", "zh_cn")
    table = {"figure": "图 " if cn else "Figure ",
             "table": "表 " if cn else "Table ",
             "equation": ""}
    return table.get(kind, "")


def _fig_label_id(n: int) -> str:
    return f"fig{n}"


def _tbl_label_id(n: int) -> str:
    return f"tbl{n}"


def _eq_label_id(n: int) -> str:
    return f"eq{n}"

def _apply_run_style(run, tokens, *, size=None, bold=False, italic=False,
                     color=None, font_key="default", cn_font_key="cn"):
    """Apply font/size/color to a run with both latin + eastAsia font set.

    v1.4 P3-QA: 统一 run 字体设置入口，确保所有新建 run 都正确设置 eastAsia 字体。
    """
    latin = tokens["font"]["family"].get(font_key, tokens["font"]["family"]["default"])
    cn = tokens["font"]["family"].get(cn_font_key, tokens["font"]["family"]["cn"])
    _set_run_font(run, latin, cn)
    if size is not None:
        run.font.size = size
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True
    if color is not None:
        run.font.color.rgb = color


# ═════════════════════════════════════════════════════════════════════
# v1.7.0 表格引擎 5 项新能力 — 模块级公开函数
# ═════════════════════════════════════════════════════════════════════
#
# 这些函数是**模块级公开 API**（不绑死在 DocxRenderer 实例上），便于：
#   1. 模板构造 doc 后直接调用
#   2. 单元测试独立 import
#   3. 既有 DocxRenderer._render_table 流程外"插入式"使用
#
# 同时也以同名方法的形式暴露在 DocxRenderer（见类内部封装）。
# 颜色策略：所有 run-color / cell-shading 都不进 heading 字段，
# self_audit 与 v1.6.6 三层防复发不受影响。
# ═════════════════════════════════════════════════════════════════════


def _validate_hex(hex_str, *, name: str = "color") -> str:
    """Validate a #RRGGBB or RRGGBB string. Return normalized form (no #)."""
    if not isinstance(hex_str, str):
        raise ValueError(f"{name}: must be a string, got {type(hex_str).__name__}")
    s = hex_str.strip().lstrip("#").upper()
    if len(s) != 6 or any(c not in "0123456789ABCDEF" for c in s):
        raise ValueError(f"{name}: must be 6-hex, got {hex_str!r}")
    return s


def _validate_positive_pct(value, *, name: str = "percent") -> float:
    """Validate a 0..100 percentage value (int/float/str)."""
    if isinstance(value, bool):
        raise ValueError(f"{name}: bool not allowed")
    if isinstance(value, (int, float)):
        v = float(value)
    elif isinstance(value, str):
        s = value.strip().rstrip("%").strip()
        try:
            v = float(s)
        except ValueError:
            raise ValueError(f"{name}: cannot parse {value!r} as percent")
    else:
        raise ValueError(f"{name}: unsupported type {type(value).__name__}")
    if v < 0 or v > 100:
        raise ValueError(f"{name}: must be in [0, 100], got {v}")
    return v


# ─── 1) render_conditional_cell：条件格式（同比标红/标绿） ─────────


def render_conditional_cell(
    cell,
    value,
    fmt: dict,
    tokens: dict,
) -> dict:
    """Render a single cell with conditional formatting.

    Args:
        cell: a ``docx.table._Cell`` (must already be added to a table row).
        value: the cell's value. May be a number, numeric string, or a
            pre-formatted pct string ("10.5%"), or "—".
        fmt: format spec dict. Supported types:

            ``"type": "yoy"`` (default)::
                {
                    "type": "yoy",
                    "positive_color": "0E7C3A",
                    "negative_color": "B91C1C",
                    "zero_color":     "666666",
                    "neutral_color":  "666666",
                    "text_color":     "1A1A1A",
                    "bg_color":       None,
                    "bold":           True,
                    "value_text":     None,
                }

            ``"type": "threshold"``::
                {
                    "type": "threshold",
                    "thresholds": [(80, "0E7C3A"), (50, "B45309"), (0, "B91C1C")],
                    "neutral_color": "666666",
                }

        tokens: token dict (used for fonts).

    Returns:
        meta dict: ``{"color": str, "value_text": str, "rule": str}``.

    Raises:
        ValueError: on unknown ``fmt["type"]`` or bad threshold config.
    """
    if not isinstance(fmt, dict):
        raise ValueError(f"render_conditional_cell: fmt must be dict, got {type(fmt).__name__}")
    fmt_type = fmt.get("type", "yoy")
    cn_font = tokens.get("font", {}).get("family", {}).get("cn", "SimSun")
    default_font = tokens.get("font", {}).get("family", {}).get("default", "Calibri")

    value_text = fmt.get("value_text")
    if value_text is None:
        value_text = str(value) if value is not None else "—"

    rule = fmt_type
    color_hex = fmt.get("text_color", "1A1A1A")
    bg_hex = fmt.get("bg_color")

    if fmt_type == "yoy":
        from ..shared.auto_calc import (
            yoy_color_for,
            YOY_COLOR_POSITIVE, YOY_COLOR_NEGATIVE, YOY_COLOR_ZERO, YOY_COLOR_NEUTRAL,
        )
        positive = _validate_hex(fmt.get("positive_color", YOY_COLOR_POSITIVE), name="positive_color")
        negative = _validate_hex(fmt.get("negative_color", YOY_COLOR_NEGATIVE), name="negative_color")
        zero = _validate_hex(fmt.get("zero_color", YOY_COLOR_ZERO), name="zero_color")
        neutral = _validate_hex(fmt.get("neutral_color", YOY_COLOR_NEUTRAL), name="neutral_color")

        rule_color = yoy_color_for(value)
        mapping = {
            YOY_COLOR_POSITIVE: positive,
            YOY_COLOR_NEGATIVE: negative,
            YOY_COLOR_ZERO: zero,
        }
        s = str(value).strip() if value is not None else ""
        if s in ("—", "-", "n/a", "N/A") or rule_color == YOY_COLOR_NEUTRAL:
            color_hex = neutral
            rule = "yoy:neutral"
        else:
            color_hex = mapping.get(rule_color, zero)
            rule = "yoy:" + (
                "positive" if rule_color == YOY_COLOR_POSITIVE
                else "negative" if rule_color == YOY_COLOR_NEGATIVE
                else "zero"
            )
    elif fmt_type == "threshold":
        thresholds = fmt.get("thresholds")
        if not isinstance(thresholds, (list, tuple)) or not thresholds:
            raise ValueError("render_conditional_cell: 'threshold' type requires non-empty 'thresholds' list")
        try:
            v = float(str(value).replace(",", "").replace("%", "").strip())
        except (TypeError, ValueError):
            v = None
        chosen = None
        if v is not None:
            sorted_th = sorted(thresholds, key=lambda x: float(x[0]), reverse=True)
            for t_val, t_color in sorted_th:
                if v >= float(t_val):
                    chosen = _validate_hex(t_color, name="threshold color")
                    rule = f"threshold:>= {t_val}"
                    break
            if chosen is None:
                chosen = _validate_hex(sorted_th[-1][1], name="threshold color")
                rule = f"threshold:< {sorted_th[-1][0]}"
        else:
            rule = "threshold:invalid"
            chosen = _validate_hex(fmt.get("neutral_color", "666666"), name="neutral_color")
        color_hex = chosen
    else:
        raise ValueError(
            f"render_conditional_cell: unknown fmt.type {fmt_type!r}; "
            f"expected 'yoy' or 'threshold'"
        )

    if fmt.get("bg_color"):
        bg_hex = _validate_hex(fmt["bg_color"], name="bg_color")

    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(value_text)
    _set_run_font(r, default_font, cn_font)
    r.font.size = Pt(10)
    r.font.bold = bool(fmt.get("bold", True))
    r.font.color.rgb = RGBColor.from_string(color_hex)
    if bg_hex:
        _set_cell_shading(cell, bg_hex)

    return {"color": color_hex, "value_text": value_text, "rule": rule}


# ─── 2) render_progress_bar：纯色矩形 + 灰底进度条 ─────────────────


def render_progress_bar(
    cell,
    percent,
    tokens: dict,
    *,
    color: str = "0E7C3A",
    track_color: str = "E5E7EB",
    show_label: bool = True,
    label_format: str = "{percent:.0f}%",
    bar_height_pt: int = 10,
) -> dict:
    """Render an in-cell progress bar (0..100%) using unicode block characters.

    Implementation note:
        python-docx does not give us a direct API to draw rectangles inside
        a cell. We use **block characters** (U+2588 ``█`` and U+2591 ``░``)
        to draw the filled / unfilled portions. This avoids needing to embed
        an image, keeps the cell content text-searchable, and renders
        correctly in both Word and WPS.

    Args:
        cell: target cell.
        percent: 0..100 (int/float/str like ``"85%"``).
        tokens: token dict.
        color: filled-bar color (default deep green).
        track_color: track color (default light grey) — currently affects
            the label tint only; the unicode glyphs carry the visual weight.
        show_label: if True, append a " 85%" label after the bar.
        label_format: format string for the label.
        bar_height_pt: visual line height; controls the cell padding.
    """
    pct = _validate_positive_pct(percent, name="percent")
    color_hex = _validate_hex(color, name="color")
    _ = _validate_hex(track_color, name="track_color")

    cn_font = tokens.get("font", {}).get("family", {}).get("cn", "SimSun")
    default_font = tokens.get("font", {}).get("family", {}).get("default", "Calibri")

    bar_width = 20
    filled = max(0, min(bar_width, int(round(pct / 100.0 * bar_width))))
    bar_text = ("\u2588" * filled) + ("\u2591" * (bar_width - filled))

    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    r_bar = p.add_run(bar_text)
    _set_run_font(r_bar, "Consolas", cn_font)
    r_bar.font.size = Pt(bar_height_pt)
    r_bar.font.color.rgb = RGBColor.from_string(color_hex)

    if show_label:
        if not label_format or "{percent" not in label_format:
            label_format = "{percent:.0f}%"
        try:
            label_text = "  " + label_format.format(percent=pct)
        except Exception:
            label_text = f"  {pct:.0f}%"
        r_lbl = p.add_run(label_text)
        _set_run_font(r_lbl, default_font, cn_font)
        r_lbl.font.size = Pt(9)
        r_lbl.font.bold = True
        r_lbl.font.color.rgb = RGBColor.from_string("333333")

    return {"percent": pct, "filled": filled, "width": bar_width, "color": color_hex}


# ─── 3) render_star_rating：unicode ★/☆ 评分 ────────────────────────


def render_star_rating(
    cell,
    rating,
    tokens: dict,
    *,
    max_stars: int = 5,
    color: str = "D97706",
    show_label: bool = True,
    full_char: str = "\u2605",
    empty_char: str = "\u2606",
) -> dict:
    """Render a star rating inside a cell.

    Args:
        cell: target cell.
        rating: int/float/str like ``4`` / ``4.5`` / ``"4.5"`` / ``"4.5/5"``.
        max_stars: total stars to draw (1..10).
        color: filled-star color (default warm amber; **not** a heading
            field, so self_audit does not block it).
        show_label: if True, append " 4.5/5" label.
        full_char: char to use for a filled star.
        empty_char: char to use for an empty star.
    """
    if not isinstance(max_stars, int) or max_stars < 1 or max_stars > 10:
        raise ValueError(f"render_star_rating: max_stars must be 1..10, got {max_stars}")

    if isinstance(rating, bool):
        raise ValueError("render_star_rating: rating must not be bool")
    if isinstance(rating, (int, float)):
        r_val = float(rating)
    elif isinstance(rating, str):
        s = rating.strip()
        if "/" in s:
            s = s.split("/", 1)[0]
        try:
            r_val = float(s)
        except ValueError:
            raise ValueError(f"render_star_rating: cannot parse {rating!r}")
    else:
        raise ValueError(f"render_star_rating: unsupported rating type {type(rating).__name__}")

    if r_val < 0 or r_val > max_stars:
        raise ValueError(
            f"render_star_rating: rating {r_val} out of range [0, {max_stars}]"
        )

    color_hex = _validate_hex(color, name="color")
    cn_font = tokens.get("font", {}).get("family", {}).get("cn", "SimSun")
    default_font = tokens.get("font", {}).get("family", {}).get("default", "Calibri")

    rounded = round(r_val * 2) / 2
    full_count = int(rounded)
    has_half = (rounded - full_count) == 0.5
    half_glyph = "\u2be8"  # ⯨ — half star; visible in most modern fonts
    stars = (full_char * full_count)
    if has_half:
        stars += half_glyph
    empty_count = max_stars - full_count - (1 if has_half else 0)
    stars += (empty_char * empty_count)

    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    r = p.add_run(stars)
    _set_run_font(r, default_font, cn_font)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(color_hex)

    if show_label:
        try:
            label = f"  {r_val:.1f}/{max_stars}"
        except Exception:
            label = f"  {r_val}/{max_stars}"
        r_lbl = p.add_run(label)
        _set_run_font(r_lbl, default_font, cn_font)
        r_lbl.font.size = Pt(9)
        r_lbl.font.color.rgb = RGBColor.from_string("666666")

    return {
        "rating": r_val,
        "rounded": rounded,
        "full": full_count,
        "has_half": has_half,
        "empty": empty_count,
        "max": max_stars,
    }


# ─── 4) render_merged_header：二维表头（行/列合并） ────────────────


def render_merged_header(
    table,
    header_rows: list,
    tokens: dict,
    *,
    col_widths_inches: list = None,
    merges: list = None,
    header_bg: str = "1F3864",
    header_text_color: str = "FFFFFF",
) -> dict:
    """Render a 2D merged header on top of an existing table.

    Args:
        table: a ``docx.table.Table`` with at least ``len(header_rows)`` rows.
        header_rows: 2D list of header cell texts, e.g. ::

            [
                ["",    "2024",    "2024",    "2024",    "2025E",  "2025E",  "2025E"],
                ["指标", "Q1",     "Q2",     "Q3",     "Q1E",   "Q2E",   "Q3E"],
            ]

            Empty strings in row N mean "merge with the cell to the left"
            (horizontal merge). Equal strings stacked vertically in column 0
            (e.g. "2024" over "2024" over "2024") trigger "merge with the
            cell above" (vertical merge). For more complex merges, pass
            ``merges`` explicitly.
        tokens: token dict.
        col_widths_inches: optional list of column widths in inches.
            Length must equal the number of columns.
        merges: optional list of ``(row, col)`` tuples to mark for **vertical**
            merge into ``(row-1, col)`` (the cell "extends downward" from
            the previous row). If None, auto-merge is attempted as above.
        header_bg: header cell background color.
        header_text_color: header text color.

    Returns:
        meta dict: ``{"rows": int, "cols": int, "merges": int}``.

    Notes:
        This function only sets up the first ``len(header_rows)`` rows of
        the table. The caller is responsible for writing data rows below.
    """
    if not isinstance(header_rows, list) or not header_rows:
        raise ValueError("render_merged_header: header_rows must be non-empty list")
    if not all(isinstance(r, list) for r in header_rows):
        raise ValueError("render_merged_header: header_rows must be list[list]")

    n_rows = len(header_rows)
    n_cols = max(len(r) for r in header_rows)
    norm_rows = [list(r) + [""] * (n_cols - len(r)) for r in header_rows]
    if len(table.rows) < n_rows:
        raise ValueError(
            f"render_merged_header: table has {len(table.rows)} rows, need >= {n_rows}"
        )

    bg_hex = _validate_hex(header_bg, name="header_bg")
    txt_hex = _validate_hex(header_text_color, name="header_text_color")

    cn_heading = tokens.get("font", {}).get("family", {}).get("cn_heading", "SimHei")
    heading_font = tokens.get("font", {}).get("family", {}).get("heading", "Calibri")

    for ri in range(n_rows):
        for ci in range(n_cols):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text = norm_rows[ri][ci]
            r = p.add_run(text)
            _set_run_font(r, heading_font, cn_heading)
            r.font.size = Pt(10)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(txt_hex)
            _set_cell_shading(cell, bg_hex)

    if col_widths_inches is not None:
        if len(col_widths_inches) != n_cols:
            raise ValueError(
                f"render_merged_header: col_widths_inches length {len(col_widths_inches)} != n_cols {n_cols}"
            )
        for ci in range(n_cols):
            for ri in range(n_rows):
                table.rows[ri].cells[ci].width = Inches(float(col_widths_inches[ci]))

    merge_count = 0
    if merges is None:
        # 水平合并：行内连续空字符串继承左侧文字
        for ri in range(n_rows):
            for ci in range(1, n_cols):
                if norm_rows[ri][ci] == "" and norm_rows[ri][ci - 1] != "":
                    try:
                        table.rows[ri].cells[ci].merge(table.rows[ri].cells[ci - 1])
                        merge_count += 1
                    except Exception:
                        pass
        # 垂直合并：列内上方文字与本行相同 → 合并
        for ci in range(n_cols):
            for ri in range(1, n_rows):
                a = norm_rows[ri - 1][ci]
                b = norm_rows[ri][ci]
                if a and a == b:
                    try:
                        table.rows[ri].cells[ci].merge(table.rows[ri - 1].cells[ci])
                        merge_count += 1
                    except Exception:
                        pass
    else:
        for ri, ci in merges:
            if ri < 1 or ri >= n_rows or ci < 0 or ci >= n_cols:
                continue
            try:
                table.rows[ri].cells[ci].merge(table.rows[ri - 1].cells[ci])
                merge_count += 1
            except Exception:
                pass

    return {"rows": n_rows, "cols": n_cols, "merges": merge_count}


# ─── 5) auto_compute_rows：自动算 YoY / 汇总行（数据驱动） ─────────


def auto_compute_rows(
    rows: list,
    *,
    yoy_col_index: int = None,
    summary_mode: str = None,
    summary_label: str = "合计",
    yoy_fmt: str = "{:.1%}",
    yoy_previous_rows: list = None,
    yoy_insert_baseline: bool = False,
) -> dict:
    """Compute derived rows for a table: YoY column, summary row, or both.

    Public wrapper around :mod:`pro_docx_gen.shared.auto_calc`. The return
    can be passed directly to the ``rows=`` field of a ``type: "table"``
    node, or fed back to ``render_conditional_cell`` for per-cell coloring.

    Args:
        rows: data rows (list of list).
        yoy_col_index: column to compute YoY for.
        summary_mode: ``"sum"`` / ``"avg"`` / None.
        summary_label: label for the summary row's first cell.
        yoy_fmt: percent format string.
        yoy_previous_rows: explicit previous-period rows.
        yoy_insert_baseline: include the raw baseline value beside YoY%.

    Returns:
        ``{"rows": new_rows, "meta": combined_meta}``.
    """
    from ..shared.auto_calc import auto_compute_rows as _impl
    new_rows, meta = _impl(
        rows,
        yoy_col_index=yoy_col_index,
        summary_mode=summary_mode,
        summary_label=summary_label,
        yoy_fmt=yoy_fmt,
        yoy_previous_rows=yoy_previous_rows,
        yoy_insert_baseline=yoy_insert_baseline,
    )
    return {"rows": new_rows, "meta": meta}


# ─── 把上面 5 个函数挂到 DocxRenderer 上（同名方法） ────────────────


def _install_v17_methods(cls):
    """Attach the 5 v1.7.0 table-engine helpers as bound methods on the
    DocxRenderer class. This lets existing renderer code do::

        self.render_conditional_cell(cell, value, fmt, self.tokens)

    while the **module-level** functions remain the canonical public API
    (better for unit tests + templates that build their own doc).
    """
    cls.render_conditional_cell = render_conditional_cell
    cls.render_progress_bar = render_progress_bar
    cls.render_star_rating = render_star_rating
    cls.render_merged_header = render_merged_header
    cls.auto_compute_rows = auto_compute_rows


# 实际挂载：让 DocxRenderer 的实例上也有同名方法（self.*）
try:
    _install_v17_methods(DocxRenderer)
except NameError:
    # 模块加载顺序问题：DocxRenderer 在本文件顶部定义，应该可用
    pass


# ═════════════════════════════════════════════════════════════════════
# v1.7.0 模块导出
# ═════════════════════════════════════════════════════════════════════

__all_v17__ = [
    "render_conditional_cell",
    "render_progress_bar",
    "render_star_rating",
    "render_merged_header",
    "auto_compute_rows",
]
