"""v1.7.0 研报专用渲染模块 — 研报 header / 评级徽章 / 风险提示尾部。

研报有强烈的视觉惯例：
- 顶部 header：股票代码 + 评级徽章 + 目标价 + 日期（**4 字段强制**）
- 副标题常用紫红斜体做品牌色（参考 Bloomberg / Reuters 风格）
- 尾部必须有风险提示披露

颜色策略（v1.7.0）：
- 评级徽章：Overweight=深绿 / Hold=琥珀 / Underweight=深红
  - **故意不用**黑金 / 浅蓝紫 / 霓虹等黑名单色
  - 全部走 cell-shading / run-color，不进 heading 字段
  - self_audit 不会拦截（heading 字段没出现）
- 紫红副标题：``#6B2C91``（不在 FORBIDDEN_HEADING_HEXES）
- 研报深蓝主色：``#1F3864``（继承 academic primary，不在黑名单）
- 风险提示：``#666666`` 斜体小字（muted 同色）

API：
- ``render_research_header(doc, tokens, header_data)``:
  - 4 字段：**stock_code / rating / target_price / report_date**
  - 副标题：``subtitle``（紫红斜体，可选）
  - 标题：``title``（股票中文名，深蓝大字）
  - 三行 meta：target_price / current_price / upside（可选）
- ``render_rating_badge(doc, tokens, rating)``:
  - 单格徽章，Overweight/Hold/Underweight 三色
- ``render_risk_disclaimer_footer(doc, tokens, *, lang='cn')``:
  - 固定模板文案的灰色斜体段落

@since v1.7.0
"""
from __future__ import annotations

from typing import Optional

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .renderer import (
    _set_cell_shading,
    _set_cell_border,
    _set_run_font,
    _set_table_row_no_split,
)


# ─── 颜色常量（v1.7.0 研报色板） ───────────────────────────────────────

# 评级三色
RATING_OVERWEIGHT_FILL   = "0E7C3A"  # 深绿
RATING_OVERWEIGHT_TEXT   = "FFFFFF"
RATING_HOLD_FILL         = "B45309"  # 琥珀
RATING_HOLD_TEXT         = "FFFFFF"
RATING_UNDERWEIGHT_FILL  = "B91C1C"  # 深红
RATING_UNDERWEIGHT_TEXT  = "FFFFFF"

# 研报主色
RESEARCH_NAVY            = "1F3864"  # 深蓝（继承 academic primary）
RESEARCH_PURPLE          = "6B2C91"  # 紫红副标题
RESEARCH_GREY_LINE       = "DDDDDD"  # 分隔线灰
RESEARCH_META_GREY       = "999999"  # meta 灰
RESEARCH_TITLE_BLACK     = "1A1A1A"  # 股票中文名（**不是** heading 字段；用 run color）

# 风险提示
RISK_TEXT_GREY           = "666666"


# ─── 内部工具 ────────────────────────────────────────────────────────


_VALID_RATINGS = {"overweight", "hold", "underweight"}


def _normalize_rating(rating: str) -> str:
    """Normalize rating text to one of ``overweight/hold/underweight``.

    Accepts common aliases (Buy/Strong Buy → overweight, Neutral → hold,
    Sell/Reduce → underweight). Case-insensitive.

    Raises:
        ValueError: when ``rating`` cannot be mapped.
    """
    if not isinstance(rating, str):
        raise ValueError(
            f"_normalize_rating: rating must be str, got {type(rating).__name__}"
        )
    s = rating.strip().lower().replace(" ", "").replace("-", "").replace("_", "")
    aliases = {
        "overweight": "overweight", "ow": "overweight",
        "buy": "overweight", "strongbuy": "overweight", "outperform": "overweight",
        "hold": "hold", "neutral": "hold", "marketperform": "hold", "mp": "hold",
        "underweight": "underweight", "uw": "underweight",
        "sell": "underweight", "reduce": "underweight", "underperform": "underweight",
    }
    if s not in aliases:
        raise ValueError(
            f"_normalize_rating: unknown rating {rating!r}; "
            f"expected one of {sorted(_VALID_RATINGS)} (aliases accepted)"
        )
    return aliases[s]


def _rating_palette(rating_key: str) -> tuple:
    """Return ``(fill_hex, text_hex, label_cn, label_en)`` for a rating key."""
    if rating_key == "overweight":
        return (
            RATING_OVERWEIGHT_FILL, RATING_OVERWEIGHT_TEXT,
            "增持", "Overweight",
        )
    if rating_key == "hold":
        return (
            RATING_HOLD_FILL, RATING_HOLD_TEXT,
            "中性", "Hold",
        )
    if rating_key == "underweight":
        return (
            RATING_UNDERWEIGHT_FILL, RATING_UNDERWEIGHT_TEXT,
            "减持", "Underweight",
        )
    raise ValueError(f"_rating_palette: invalid rating_key {rating_key!r}")


# ─── 评级徽章（公开 API） ──────────────────────────────────────────────


def render_rating_badge(
    doc,
    tokens: dict,
    rating: str,
    *,
    lang: str = "cn",
    width_inches: float = 1.0,
    height_inches: Optional[float] = None,
) -> dict:
    """Render a single rating badge cell into ``doc``.

    Args:
        doc: active Document.
        tokens: token dict (used for font defaults).
        rating: ``"overweight"`` / ``"hold"`` / ``"underweight"`` (aliases
            like Buy/Neutral/Sell accepted).
        lang: "cn" → 中文标签（增持/中性/减持）；"en" → 英文标签。
        width_inches: badge width in inches (default 1.0).
        height_inches: optional row height. Default 0.4 in.

    Returns:
        meta dict: ``{"rating_key": str, "fill": str, "label": str, "table_index": int}``.
    """
    key = _normalize_rating(rating)
    fill_hex, text_hex, label_cn, label_en = _rating_palette(key)
    label = label_cn if lang == "cn" else label_en

    cn_font = tokens.get("font", {}).get("family", {}).get("cn_heading", "SimHei")
    heading_font = tokens.get("font", {}).get("family", {}).get("heading", "Calibri")

    # 单行单列表格
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_row_no_split(table.rows[0])
    if height_inches is not None:
        table.rows[0].height = Inches(height_inches)

    cell = table.rows[0].cells[0]
    cell.width = Inches(width_inches)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_shading(cell, fill_hex)
    border_attrs = {"val": "single", "sz": "4", "color": fill_hex}
    _set_cell_border(cell, top=border_attrs, bottom=border_attrs,
                     left=border_attrs, right=border_attrs)

    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label)
    _set_run_font(r, heading_font, cn_font)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(text_hex)

    return {
        "rating_key": key,
        "fill": fill_hex,
        "label": label,
        "table_index": len(doc.tables) - 1,
    }


# ─── 研报 header（公开 API） ────────────────────────────────────────────


def render_research_header(
    doc,
    tokens: dict,
    header_data: dict,
    *,
    lang: str = "cn",
) -> dict:
    """Render the research report header block into ``doc``.

    4 字段强制：
    - ``stock_code``   ：股票代码（如 "002415.SZ" / "AAPL US"），必填
    - ``rating``       ：评级，必填（接受 Overweight/Hold/Underweight + 别名）
    - ``target_price`` ：目标价，必填（数值或字符串）
    - ``report_date``  ：报告日期，必填

    可选：
    - ``title``         ：股票中文名（标题大字）
    - ``subtitle``      ：紫红斜体副标题（行业 / 主题）
    - ``current_price`` ：当前价（与 target_price 一起算 upside）
    - ``analyst``       ：分析师署名（meta 行）

    Returns:
        meta dict: ``{"rating_key": str, "upside_pct": float|None}``.
        ``upside_pct`` 来自 ``(target_price - current_price) / current_price``。
    """
    # --- 强制字段校验 ---
    missing = []
    for f in ("stock_code", "rating", "target_price", "report_date"):
        if f not in header_data or header_data[f] in (None, ""):
            missing.append(f)
    if missing:
        raise ValueError(
            f"render_research_header: 4 字段强制，缺少 {missing}; "
            f"必须提供 stock_code / rating / target_price / report_date"
        )

    stock_code = str(header_data["stock_code"]).strip()
    rating = header_data["rating"]
    target_price = header_data["target_price"]
    report_date = str(header_data["report_date"]).strip()
    title = str(header_data.get("title") or "").strip()
    subtitle = str(header_data.get("subtitle") or "").strip()
    current_price = header_data.get("current_price")
    analyst = str(header_data.get("analyst") or "").strip()

    rating_key = _normalize_rating(rating)
    fill_hex, text_hex, label_cn, label_en = _rating_palette(rating_key)
    rating_label = label_cn if lang == "cn" else label_en

    # 字体
    cn_font = tokens.get("font", {}).get("family", {}).get("cn", "SimSun")
    cn_heading = tokens.get("font", {}).get("family", {}).get("cn_heading", "SimHei")
    heading_font = tokens.get("font", {}).get("family", {}).get("heading", "Calibri")
    default_font = tokens.get("font", {}).get("family", {}).get("default", "Calibri")

    # ─── 行 1：meta (灰色小字) ───
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta_p.paragraph_format.space_before = Pt(0)
    meta_p.paragraph_format.space_after = Pt(2)
    meta_text_parts = []
    if analyst:
        meta_text_parts.append(f"分析师：{analyst}" if lang == "cn" else f"Analyst: {analyst}")
    meta_text_parts.append(
        f"报告日期：{report_date}" if lang == "cn" else f"Report Date: {report_date}"
    )
    if current_price is not None:
        meta_text_parts.append(
            f"当前价：{current_price}" if lang == "cn" else f"Current: {current_price}"
        )
    mr = meta_p.add_run("  ·  ".join(meta_text_parts))
    _set_run_font(mr, default_font, cn_font)
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor.from_string(RESEARCH_META_GREY)
    mr.font.italic = True

    # ─── 行 2：title (深蓝大字) + 评级徽章 (右侧) ───
    # 用 1×2 表格：左 70% 标题 / 右 30% 徽章
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _set_table_row_no_split(table.rows[0])
    table.rows[0].height = Inches(0.55)

    # 内容宽度估算（与 KPI 卡片同）
    layout = tokens.get("__layout_content_width_emu__")
    content_width_in = (float(layout) / 914400.0) if layout else 6.5
    title_w = content_width_in * 0.72
    badge_w = content_width_in - title_w

    title_cell = table.rows[0].cells[0]
    title_cell.width = Inches(title_w)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 不要边框
    border_attrs = {"val": "nil", "sz": "0", "color": "auto"}
    _set_cell_border(title_cell, top=border_attrs, bottom=border_attrs,
                     left=border_attrs, right=border_attrs)

    title_cell.text = ""
    p_title = title_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(0)

    if title:
        r1 = p_title.add_run(title + "  ")
        _set_run_font(r1, heading_font, cn_heading)
        r1.font.size = Pt(20)
        r1.font.bold = True
        r1.font.color.rgb = RGBColor.from_string(RESEARCH_NAVY)

    r2 = p_title.add_run(stock_code)
    _set_run_font(r2, default_font, cn_font)
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor.from_string(RESEARCH_TITLE_BLACK)
    r2.font.bold = True

    # 评级徽章 cell
    badge_cell = table.rows[0].cells[1]
    badge_cell.width = Inches(badge_w)
    badge_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_shading(badge_cell, fill_hex)
    bd_attrs = {"val": "single", "sz": "4", "color": fill_hex}
    _set_cell_border(badge_cell, top=bd_attrs, bottom=bd_attrs,
                     left=bd_attrs, right=bd_attrs)
    badge_cell.text = ""
    p_b = badge_cell.paragraphs[0]
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_b.paragraph_format.space_before = Pt(2)
    p_b.paragraph_format.space_after = Pt(2)
    rb = p_b.add_run(f"评级  {rating_label}")
    _set_run_font(rb, heading_font, cn_heading)
    rb.font.size = Pt(11)
    rb.font.bold = True
    rb.font.color.rgb = RGBColor.from_string(text_hex)

    # ─── 行 3：紫红斜体副标题 ───
    if subtitle:
        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sp.paragraph_format.space_before = Pt(2)
        sp.paragraph_format.space_after = Pt(2)
        sr = sp.add_run(subtitle)
        _set_run_font(sr, default_font, cn_font)
        sr.font.size = Pt(12)
        sr.font.italic = True
        sr.font.bold = True
        sr.font.color.rgb = RGBColor.from_string(RESEARCH_PURPLE)

    # ─── 行 4：target / current / upside 三行 meta ───
    target_str = str(target_price)
    current_str = str(current_price) if current_price is not None else "—"
    upside_text = "—"
    upside_pct: Optional[float] = None
    if current_price is not None:
        try:
            cur_n = float(str(current_price).replace(",", "").replace("¥", "").replace("$", ""))
            tgt_n = float(str(target_price).replace(",", "").replace("¥", "").replace("$", ""))
            if cur_n > 0:
                upside_pct = (tgt_n - cur_n) / cur_n
                arrow = "▲" if upside_pct > 0 else ("▼" if upside_pct < 0 else "●")
                sign = "+" if upside_pct > 0 else ""
                upside_text = f"{arrow} {sign}{upside_pct:.1%}"
        except (TypeError, ValueError):
            pass

    meta2 = doc.add_table(rows=1, cols=3)
    meta2.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta2.autofit = False
    _set_table_row_no_split(meta2.rows[0])
    meta2.rows[0].height = Inches(0.42)

    cells = meta2.rows[0].cells
    col_w = content_width_in / 3
    for c in cells:
        c.width = Inches(col_w)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(c, "F8F9FB")
        bd = {"val": "single", "sz": "4", "color": RESEARCH_GREY_LINE}
        _set_cell_border(c, top=bd, bottom=bd, left=bd, right=bd)

    # Target Price
    cells[0].text = ""
    p0 = cells[0].paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(0)
    rk = p0.add_run("目标价" if lang == "cn" else "Target")
    _set_run_font(rk, default_font, cn_font)
    rk.font.size = Pt(9)
    rk.font.color.rgb = RGBColor.from_string(RESEARCH_META_GREY)
    p0b = cells[0].add_paragraph()
    p0b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0b.paragraph_format.space_before = Pt(0)
    p0b.paragraph_format.space_after = Pt(2)
    rv = p0b.add_run(target_str)
    _set_run_font(rv, heading_font, cn_heading)
    rv.font.size = Pt(13)
    rv.font.bold = True
    rv.font.color.rgb = RGBColor.from_string(RESEARCH_NAVY)

    # Current Price
    cells[1].text = ""
    p1 = cells[1].paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(0)
    rk1 = p1.add_run("当前价" if lang == "cn" else "Current")
    _set_run_font(rk1, default_font, cn_font)
    rk1.font.size = Pt(9)
    rk1.font.color.rgb = RGBColor.from_string(RESEARCH_META_GREY)
    p1b = cells[1].add_paragraph()
    p1b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1b.paragraph_format.space_before = Pt(0)
    p1b.paragraph_format.space_after = Pt(2)
    rv1 = p1b.add_run(current_str)
    _set_run_font(rv1, heading_font, cn_heading)
    rv1.font.size = Pt(13)
    rv1.font.bold = True
    rv1.font.color.rgb = RGBColor.from_string(RESEARCH_TITLE_BLACK)

    # Upside
    cells[2].text = ""
    p2 = cells[2].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(0)
    rk2 = p2.add_run("上行空间" if lang == "cn" else "Upside")
    _set_run_font(rk2, default_font, cn_font)
    rk2.font.size = Pt(9)
    rk2.font.color.rgb = RGBColor.from_string(RESEARCH_META_GREY)
    p2b = cells[2].add_paragraph()
    p2b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2b.paragraph_format.space_before = Pt(0)
    p2b.paragraph_format.space_after = Pt(2)
    # 颜色规则：上行 > 0 绿 / < 0 红 / 0 灰
    if upside_pct is None or upside_text == "—":
        up_color = RISK_TEXT_GREY
    elif upside_pct > 0:
        up_color = RATING_OVERWEIGHT_FILL
    elif upside_pct < 0:
        up_color = RATING_UNDERWEIGHT_FILL
    else:
        up_color = RISK_TEXT_GREY
    rv2 = p2b.add_run(upside_text)
    _set_run_font(rv2, heading_font, cn_heading)
    rv2.font.size = Pt(13)
    rv2.font.bold = True
    rv2.font.color.rgb = RGBColor.from_string(up_color)

    # 行 4 后留空段
    after_p = doc.add_paragraph()
    after_p.paragraph_format.space_before = Pt(4)
    after_p.paragraph_format.space_after = Pt(8)

    return {
        "rating_key": rating_key,
        "upside_pct": upside_pct,
        "upside_text": upside_text,
    }


# ─── 风险提示尾部（公开 API） ─────────────────────────────────────────


_RISK_TEXT_CN = (
    "本报告仅供研究参考，不构成投资建议；投资有风险，决策需谨慎。"
    "本机构不对依据本报告作出的任何投资决策承担责任，亦不保证报告中"
    "信息的绝对准确性和完整性。报告所载数据来源于公开信息，"
    "本机构对其及时性、准确性不作任何承诺。"
    "投资者应当独立判断本报告中的任何信息和判断，并据此做出投资决策。"
)

_RISK_TEXT_EN = (
    "This report is intended for research purposes only and does not "
    "constitute investment advice. Investing involves risk; decisions "
    "should be made cautiously. The institution accepts no responsibility "
    "for any investment decision made on the basis of this report, nor "
    "does it guarantee the absolute accuracy or completeness of the "
    "information herein. Data are sourced from public information; the "
    "institution makes no commitment to its timeliness or accuracy. "
    "Investors should independently evaluate all information and judgments "
    "in this report before making investment decisions."
)


def render_risk_disclaimer_footer(
    doc,
    tokens: dict,
    *,
    lang: str = "cn",
    institution: Optional[str] = None,
) -> dict:
    """Render the fixed risk-disclaimer footer block at the end of the doc.

    灰色斜体小字，自动出现在每份研报最后一页（caller 负责在合适位置调用）。
    - lang="cn" → 中文（默认）
    - lang="en" → 英文
    - institution：可选机构名，会拼到文末

    Returns:
        meta dict: ``{"chars": int, "lang": str}``.
    """
    text = _RISK_TEXT_CN if lang == "cn" else _RISK_TEXT_EN
    if institution:
        sep = " " if lang == "en" else "  "
        text = text + sep + (
            f"——{institution}" if lang == "cn" else f"— {institution}"
        )

    cn_font = tokens.get("font", {}).get("family", {}).get("cn", "SimSun")
    default_font = tokens.get("font", {}).get("family", {}).get("default", "Calibri")

    # 上方细灰分隔线
    line_p = doc.add_paragraph()
    line_p.paragraph_format.space_before = Pt(8)
    line_p.paragraph_format.space_after = Pt(2)
    line_run = line_p.add_run("─" * 40)
    _set_run_font(line_run, default_font, cn_font)
    line_run.font.size = Pt(8)
    line_run.font.color.rgb = RGBColor.from_string(RESEARCH_GREY_LINE)

    # 风险提示正文
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _set_run_font(r, default_font, cn_font)
    r.font.size = Pt(8)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(RISK_TEXT_GREY)

    return {"chars": len(text), "lang": lang}


__all__ = [
    "render_research_header",
    "render_rating_badge",
    "render_risk_disclaimer_footer",
    # 颜色常量（供外部/模板引用）
    "RATING_OVERWEIGHT_FILL", "RATING_OVERWEIGHT_TEXT",
    "RATING_HOLD_FILL", "RATING_HOLD_TEXT",
    "RATING_UNDERWEIGHT_FILL", "RATING_UNDERWEIGHT_TEXT",
    "RESEARCH_NAVY", "RESEARCH_PURPLE", "RESEARCH_GREY_LINE", "RESEARCH_META_GREY",
    "RISK_TEXT_GREY",
]
