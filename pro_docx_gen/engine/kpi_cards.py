"""v1.7.0 KPI 卡片行 — 横向 3-5 卡片（label / value / delta / mini trend）。

场景：研报 / 财务报告首页 / 季度业绩页"关键指标一眼可见"行。

设计：
- 1×N 表格实现横向卡片，卡之间用细灰线分隔。
- 数值用深黑 ``#1A1A1A`` 24pt 加粗；同比用 ▲▼ 符号 + 染色（正绿/负红/0灰/不可计算灰）。
- 趋势 mini 图：可选 ``trend`` 字段（list[float]），用 unicode 块字符在单元格内
  拼一个 sparkline（不依赖 matplotlib，避免插入图片和字体测量问题）。
- 颜色全部走 run color（不进 self_audit 黑名单）；卡片底色用浅灰 ``#F8F9FB`` 显层次。
- 修根因：单元格宽度由可用 content_width 均分；至少 3、最多 5 张卡，超出抛 ValueError。
- 不吞错：所有异常（参数类型 / 颜色字符串）都上抛 ValueError。

API：
- ``render_kpi_card_row(doc, cards, tokens, *, lang='en', card_widths=None)``
  - ``doc``: python-docx ``Document``（由 renderer 传入）
  - ``cards``: list[dict]，每张卡 ``{label, value, delta, trend?, subtext?, suffix?}``
  - ``tokens``: 渲染 token dict（与 ``DocxRenderer.tokens`` 同结构）
  - ``lang``: "en" | "cn"，影响 "QoQ" vs "环比" 等文字（当前只影响分隔符）
  - ``card_widths``: 可选，list[Inches]。None 时按 content_width / N 等分。
- 返回：插入到 doc 的表格对象 + meta dict。

@since v1.7.0
"""
from __future__ import annotations

from typing import List, Optional, Sequence

from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from ..shared.auto_calc import yoy_color_for
from .renderer import (
    _set_cell_shading,
    _set_cell_border,
    _set_run_font,
    _set_table_row_no_split,
)


# v1.7.0 颜色常量（不在 self_audit 黑名单内；这些用于 run/table-cell
# shading，不是 heading 字段）
_KPI_CARD_FILL     = "F8F9FB"   # 卡片底色（极浅冷灰）
_KPI_CARD_BORDER   = "DDDDDD"   # 卡间分隔线
_KPI_CARD_VALUE    = "1A1A1A"   # value 数字（深黑，**不是**heading，但保持高对比）
_KPI_CARD_LABEL    = "666666"   # label 灰（用 muted）
_KPI_CARD_SUBTEXT  = "999999"   # 副文本更淡的灰
_KPI_TREND_POS     = "0E7C3A"   # 正向趋势
_KPI_TREND_NEG     = "B91C1C"   # 负向趋势

# 卡片数限制
_KPI_CARD_MIN = 3
_KPI_CARD_MAX = 5


# ─── 内部工具 ────────────────────────────────────────────────────────


def _sparkline_unicode(values: Sequence[float], *, max_width: int = 7) -> str:
    """Render a tiny sparkline using unicode block characters.

    用 ▁▂▃▄▅▆▇█ 8 个块字符按比例映射到当前 max/min 之间。无 matplotlib
    依赖，避免 chart 资产的预检失败。返回 1 行字符串。
    """
    if not values:
        return ""
    # 取最近 max_width 个点
    series = list(values)[-max_width:]
    if len(series) < 2:
        # 至少 2 个点才能画线
        return ""
    lo = min(series)
    hi = max(series)
    span = hi - lo
    # 8 档 block 字符
    blocks = " ▁▂▃▄▅▆▇█"
    chars: List[str] = []
    if span == 0:
        # 全平
        chars = ["▄"] * len(series)
    else:
        for v in series:
            idx = int(round((v - lo) / span * (len(blocks) - 1)))
            idx = max(0, min(len(blocks) - 1, idx))
            chars.append(blocks[idx])
    return "".join(chars)


def _format_delta(delta, *, lang: str = "en") -> str:
    """Format the delta into ``"▲ 10.5%"`` / ``"▼ -3.2%"`` / ``"—"`` etc.

    delta 可以是 float（已算好的 ratio）或字符串（"10.5%" / "-3.2%" / "—"）。
    """
    if delta is None:
        return "—"
    if isinstance(delta, str):
        s = delta.strip()
        if not s or s in ("—", "-", "n/a", "N/A"):
            return "—"
        # 字符串以 +/- 开头：保留符号 + 加箭头
        if s.startswith("+"):
            return f"▲ {s[1:].strip()}"
        if s.startswith("-"):
            return f"▼ {s}"
        # 普通 pct 字符串：让 sign 决定箭头
        try:
            v = float(s.rstrip("%").replace(",", ""))
            if v > 0:
                return f"▲ +{s.lstrip('+')}"
            if v < 0:
                return f"▼ {s}"
            return f"● {s}"
        except ValueError:
            return s
    if isinstance(delta, bool):
        return "—"
    if isinstance(delta, (int, float)):
        if delta > 0:
            return f"▲ +{delta:.1%}"
        if delta < 0:
            return f"▼ {delta:.1%}"
        return f"● 0.0%"
    return str(delta)


def _delta_color(delta, *, neutral: bool = False) -> str:
    """Return the run-color hex (no ``#``) for a delta value."""
    if neutral:
        return "666666"
    return yoy_color_for(delta)


# ─── 公开 API ────────────────────────────────────────────────────────


def render_kpi_card_row(
    doc,
    cards: Sequence[dict],
    tokens: dict,
    *,
    lang: str = "en",
    card_widths: Optional[Sequence[float]] = None,
) -> dict:
    """Render a horizontal KPI card row into ``doc``.

    Args:
        doc: an active ``docx.Document`` (caller owns the save).
        cards: list of card dicts. Each card requires ``label`` and ``value``;
            optional ``delta`` (float or string), ``subtext``, ``suffix``
            (e.g. "亿元" / "USD mn"), ``trend`` (list[float] for sparkline),
            ``trend_invert`` (True for metrics where down=good, e.g. 成本).
        tokens: token dict matching ``DocxRenderer.tokens``. We only use
            ``tokens["color"]["text"]`` / ``tokens["font"]["family"]["default"]``
            for fallbacks; the canonical KPI colors are module constants.
        lang: "en" | "cn" (currently only affects the QoQ vs 环比 suffix).
        card_widths: optional list of card widths in inches. Must match len(cards).
            Default: content_width / N.

    Returns:
        meta dict: ``{"cards": int, "skipped": int, "table_position": int}``
        where ``table_position`` is the 0-based index of the newly added
        table in ``doc.tables``.

    Raises:
        ValueError: when ``len(cards) < 3`` or ``> 5`` or
            ``card_widths`` length doesn't match.
    """
    if not cards:
        raise ValueError("render_kpi_card_row: cards must be a non-empty sequence")
    n = len(cards)
    if n < _KPI_CARD_MIN:
        raise ValueError(
            f"render_kpi_card_row: requires at least {_KPI_CARD_MIN} cards, got {n}"
        )
    if n > _KPI_CARD_MAX:
        raise ValueError(
            f"render_kpi_card_row: at most {_KPI_CARD_MAX} cards per row, got {n}"
        )
    if card_widths is not None and len(card_widths) != n:
        raise ValueError(
            f"render_kpi_card_row: card_widths length {len(card_widths)} != n_cards {n}"
        )

    # 颜色 / 字体
    cn_font = tokens.get("font", {}).get("family", {}).get("cn", "SimSun")
    default_font = tokens.get("font", {}).get("family", {}).get("default", "Calibri")
    heading_font = tokens.get("font", {}).get("family", {}).get("heading", "Calibri")

    # 内容宽度估算：与 renderer._render_table 一致
    layout = tokens.get("__layout_content_width_emu__")
    if layout is None:
        # 退化估算：A4 默认 1in 边距 → 6.5in
        content_width_in = 6.5
    else:
        content_width_in = float(layout) / 914400.0

    # 计算每张卡宽度
    if card_widths is not None:
        widths = [Inches(float(w)) for w in card_widths]
    else:
        per_card = content_width_in / n
        widths = [Inches(per_card) for _ in range(n)]

    # 表格：1 行 × N 列
    table = doc.add_table(rows=1, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_row_no_split(table.rows[0])

    # 卡间分隔线用细灰边
    border_attrs = {"val": "single", "sz": "4", "color": _KPI_CARD_BORDER}
    inner_border = {"val": "single", "sz": "4", "color": _KPI_CARD_BORDER}
    # 整行底色
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, _KPI_CARD_FILL)

    # 每张卡填内容
    skipped = 0
    for i, card in enumerate(cards):
        if not isinstance(card, dict):
            raise ValueError(
                f"render_kpi_card_row: card #{i} must be a dict, got {type(card).__name__}"
            )
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 边框：左右加细线，制造"卡片分割"感
        _set_cell_border(
            cell,
            top=border_attrs, bottom=border_attrs,
            left=inner_border if i > 0 else {"val": "nil", "sz": "0", "color": "auto"},
            right=inner_border if i < n - 1 else {"val": "nil", "sz": "0", "color": "auto"},
        )

        # 清空默认段
        cell.text = ""
        label = str(card.get("label", "")).strip()
        value = card.get("value", "")
        suffix = str(card.get("suffix", "")).strip()
        delta = card.get("delta")
        subtext = str(card.get("subtext", "")).strip()
        trend = card.get("trend")
        trend_invert = bool(card.get("trend_invert", False))

        if not label and value == "":
            skipped += 1
            continue

        # --- 1) Label (小字灰) ---
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(label or "—")
        _set_run_font(r1, default_font, cn_font)
        r1.font.size = Pt(9)
        r1.font.color.rgb = RGBColor.from_string(_KPI_CARD_LABEL)

        # --- 2) Value (大字深黑) + suffix (小字灰) ---
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        r2a = p2.add_run(str(value))
        _set_run_font(r2a, heading_font, cn_font)
        r2a.font.size = Pt(20)
        r2a.font.bold = True
        r2a.font.color.rgb = RGBColor.from_string(_KPI_CARD_VALUE)
        if suffix:
            r2b = p2.add_run(" " + suffix)
            _set_run_font(r2b, default_font, cn_font)
            r2b.font.size = Pt(10)
            r2b.font.color.rgb = RGBColor.from_string(_KPI_CARD_LABEL)

        # --- 3) Delta (▲▼ + 染色) ---
        delta_text = _format_delta(delta, lang=lang)
        delta_color = _delta_color(delta, neutral=(delta_text == "—"))
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(0)
        r3 = p3.add_run(delta_text)
        _set_run_font(r3, default_font, cn_font)
        r3.font.size = Pt(10)
        r3.font.bold = True
        r3.font.color.rgb = RGBColor.from_string(delta_color)

        # --- 4) Optional sparkline (unicode) ---
        if trend and isinstance(trend, (list, tuple)) and len(trend) >= 2:
            # 颜色：与 delta 同色（trend_invert 时反色）
            line_color = delta_color
            if trend_invert:
                if delta_color == _KPI_TREND_POS:
                    line_color = _KPI_TREND_NEG
                elif delta_color == _KPI_TREND_NEG:
                    line_color = _KPI_TREND_POS
            p4 = cell.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.space_before = Pt(0)
            p4.paragraph_format.space_after = Pt(2)
            spark = _sparkline_unicode(trend, max_width=7)
            r4 = p4.add_run(spark or "—")
            _set_run_font(r4, default_font, cn_font)
            r4.font.size = Pt(12)
            r4.font.color.rgb = RGBColor.from_string(line_color)
        elif subtext:
            p4 = cell.add_paragraph()
            p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p4.paragraph_format.space_before = Pt(0)
            p4.paragraph_format.space_after = Pt(2)
            r4 = p4.add_run(subtext)
            _set_run_font(r4, default_font, cn_font)
            r4.font.size = Pt(9)
            r4.font.color.rgb = RGBColor.from_string(_KPI_CARD_SUBTEXT)

    # 行后空一行（紧凑）
    after_p = doc.add_paragraph()
    after_p.paragraph_format.space_after = Pt(2)

    return {
        "cards": n - skipped,
        "skipped": skipped,
        "table_position": len(doc.tables) - 1,
    }


__all__ = [
    "render_kpi_card_row",
    "_KPI_CARD_MIN",
    "_KPI_CARD_MAX",
]
