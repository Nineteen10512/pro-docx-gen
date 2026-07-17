"""Structural DOCX QA for defects that sampled rendering can miss."""
from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


def expand_review_pages(figure_pages: list[int], page_count: int) -> list[int]:
    """Return every figure page plus its immediate neighbors."""
    pages: set[int] = set()
    for page in figure_pages:
        for candidate in (page - 1, page, page + 1):
            if 1 <= candidate <= page_count:
                pages.add(candidate)
    return sorted(pages)


def _issue(level: str, code: str, message: str, **details) -> dict:
    item = {"level": level, "code": code, "message": message}
    if details:
        item["details"] = details
    return item


def _table_border_conflicts(document) -> list[dict]:
    conflicts: list[dict] = []
    for table_index, table in enumerate(document.tables, start=1):
        table_colors = {
            node.get(qn("w:color"), "").upper()
            for node in table._tbl.tblPr.findall(".//" + qn("w:tblBorders") + "/*")
            if node.get(qn("w:color"))
        }
        if not table_colors:
            continue
        cell_colors = {
            node.get(qn("w:color"), "").upper()
            for node in table._tbl.findall(".//" + qn("w:tcBorders") + "/*")
            if node.get(qn("w:color"))
        }
        unexpected = sorted(cell_colors - table_colors - {"AUTO"})
        if unexpected:
            conflicts.append(
                _issue(
                    "ERROR",
                    "cell_border_color_conflict",
                    f"table {table_index} has direct cell border colors that override the table palette",
                    table_index=table_index,
                    table_colors=sorted(table_colors),
                    cell_colors=unexpected,
                )
            )
    return conflicts


def _font_size_drift(document) -> list[dict]:
    sizes: dict[str, list[float]] = defaultdict(list)
    # The first four paragraphs are the conventional title/rule/subtitle/author
    # cover block and intentionally use different sizes.
    for paragraph in document.paragraphs[4:]:
        role = paragraph.style.name if paragraph.style else "Normal"
        for run in paragraph.runs:
            if run.text.strip() and run.font.size is not None:
                sizes[role].append(round(run.font.size.pt, 2))
    issues: list[dict] = []
    for role, values in sizes.items():
        if len(values) >= 2 and max(values) - min(values) >= 3.0:
            issues.append(
                _issue(
                    "WARNING",
                    "role_font_size_drift",
                    f"paragraph role {role!r} spans {min(values)}–{max(values)} pt",
                    role=role,
                    minimum=min(values),
                    maximum=max(values),
                )
            )
    return issues


def _redundant_page_breaks(document) -> list[dict]:
    issues: list[dict] = []
    paragraphs = document.paragraphs
    for index, paragraph in enumerate(paragraphs):
        if not paragraph.style or paragraph.style.name not in ("Heading 1", "Part Divider"):
            continue
        previous = paragraph._p.getprevious()
        if previous is None or previous.tag != qn("w:p"):
            continue
        if "".join(previous.itertext()).strip():
            continue
        if previous.find(".//" + qn("w:br")) is not None:
            issues.append(
                _issue(
                    "ERROR",
                    "redundant_page_break",
                    "manual page break immediately precedes a page-break heading and can create a blank page",
                    paragraph_index=index,
                    heading=paragraph.text,
                )
            )
    return issues


def _section_break_kind(element) -> str | None:
    if element is None or element.tag != qn("w:p"):
        return None
    p_pr = element.find(qn("w:pPr"))
    sect_pr = p_pr.find(qn("w:sectPr")) if p_pr is not None else None
    if sect_pr is None:
        return None
    page_size = sect_pr.find(qn("w:pgSz"))
    if page_size is not None and page_size.get(qn("w:orient")) == "landscape":
        return "landscape"
    return "portrait"


def _adjacent_section_breaks(document) -> list[dict]:
    children = list(document._element.body.iterchildren())
    for index in range(len(children) - 1):
        if _section_break_kind(children[index]) == "landscape" and _section_break_kind(children[index + 1]) == "portrait":
            return [
                _issue(
                    "ERROR",
                    "adjacent_section_blank_page",
                    "consecutive landscape and portrait section-break paragraphs can emit a blank portrait page",
                    body_index=index,
                )
            ]
    return []


def _toc_integrity(path: Path) -> list[dict]:
    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    refs = set(re.findall(r"PAGEREF\s+(_Toc\d+)", xml))
    bookmarks = set(re.findall(r'w:bookmarkStart[^>]+w:name="(_Toc\d+)"', xml))
    missing = sorted(refs - bookmarks)
    if not missing:
        return []
    return [
        _issue(
            "ERROR",
            "dangling_toc_bookmark",
            f"{len(missing)} PAGEREF fields point to missing TOC bookmarks",
            bookmarks=missing,
        )
    ]


def audit_docx_layout(path: str | Path, minimum_figure_width_inches: float = 5.5) -> dict:
    """Audit structural layout defects before and after render."""
    target = Path(path)
    document = Document(str(target))
    issues: list[dict] = []
    issues.extend(_table_border_conflicts(document))
    issues.extend(_font_size_drift(document))
    issues.extend(_redundant_page_breaks(document))
    issues.extend(_adjacent_section_breaks(document))
    issues.extend(_toc_integrity(target))

    widths = [shape.width / 914400 for shape in document.inline_shapes]
    small = [round(width, 2) for width in widths if width < minimum_figure_width_inches]
    if small:
        issues.append(
            _issue(
                "WARNING",
                "figure_too_small",
                f"{len(small)} inline figures are narrower than {minimum_figure_width_inches} inches",
                displayed_width_inches=small,
            )
        )

    return {
        "path": str(target.resolve()),
        "passed": not any(item["level"] == "ERROR" for item in issues),
        "issues": issues,
        "figure_count": len(widths),
        "requires_full_figure_review": bool(widths),
        "review_policy": (
            "Render all pages when figures exist; review every figure page and immediate neighbors."
            if widths
            else "Sampled rendering is permitted when structural QA passes."
        ),
    }
