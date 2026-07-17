"""Lightweight delivery quality gate for PRO-PPTX / PRO-DOCX outputs.

This script is inspired by the built-in presentation/document skills' delivery
discipline: inspect structure first, render when the local environment supports
it, and return a machine-readable report.

Inspection intentionally uses only the standard library. The optional
``--auto-fix`` path uses python-docx, already a required skill dependency.
LibreOffice rendering is strict by default: warnings and render skips/failures
block delivery unless the caller explicitly opts out for a local diagnostic run.

v1.6.6: added ``check_heading_color_consistency`` so that the quality gate
mirrors the v1.6.6 self-audit discipline. The new check ensures (a) every
heading paragraph uses a non-forbidden colour and (b) all paragraphs at
the same heading level share a single colour. The optional
``--with-self-audit`` flag additionally runs the template-level audit
(from ``pro_docx_gen.self_audit``) so that template regressions are
caught before the renderer even starts.
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from collections import Counter
from pathlib import Path
from xml.etree import ElementTree as ET


PLACEHOLDER_RE = re.compile(r"TODO|TBD|待补|占位|示例文本|\[chart\]|\[image\]", re.I)
AI_TELL_RE = re.compile(
    r"\b(?:lorem ipsum|john doe|jane doe|acme|nexus|smartflow|cloudly|"
    r"synergy|next-gen|cutting-edge|revolutionize|unlock|elevate|"
    r"seamless|world-class|game-changing)\b|"
    r"\u8d4b\u80fd|\u98a0\u8986|\u91cd\u5851|\u65e0\u7f1d|\u4e0b\u4e00\u4ee3",
    re.I,
)
DECORATIVE_DASH_RE = re.compile(r"[\u2013\u2014]")
LONG_TEXT_LIMIT = 240
PPT_SLIDE_TEXT_LIMIT = 1200
DOCX_PARAGRAPH_LIMIT = 900
PRICE_TOKEN_RE = re.compile(r"^\$?\d{2,4}(?:,\d{3})*(?:\.\d{2})?$")
AUTO_FIXABLE_CODES = {
    "table_header_narrow",
    "table_body_word_narrow",
    "table_row_can_split",
    "callout_row_can_split",
    "table_orphan_split_risk",
}
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def _issue(level: str, code: str, message: str) -> dict:
    return {"level": level, "code": code, "message": message}


def _parse_timestamp(value: str) -> float:
    try:
        return float(value)
    except ValueError:
        pass
    text = value.strip()
    if text.endswith("Z"):
        text = text[:-1] + "+00:00"
    try:
        parsed = dt.datetime.fromisoformat(text)
    except ValueError as exc:
        raise argparse.ArgumentTypeError(
            "timestamp must be Unix seconds or ISO-8601, e.g. 1783512000 or 2026-07-09T10:00:00+08:00"
        ) from exc
    if parsed.tzinfo is None:
        parsed = parsed.replace(tzinfo=dt.datetime.now().astimezone().tzinfo)
    return parsed.timestamp()


def select_sample_page_numbers(page_count: int, group_size: int = 3) -> list[int]:
    """Return one-based source pages sampled from consecutive page groups."""
    if page_count < 1:
        raise ValueError("page_count must be at least 1")
    if group_size < 1:
        raise ValueError("group_size must be at least 1")
    return list(range(1, page_count + 1, group_size))


def _read_zip_text(zf: zipfile.ZipFile, name: str) -> str:
    return zf.read(name).decode("utf-8", errors="ignore")


def _xml_text(xml: str) -> list[str]:
    try:
        root = ET.fromstring(xml)
    except ET.ParseError:
        return []
    out: list[str] = []
    for el in root.iter():
        if el.text and el.tag.endswith("}t"):
            out.append(el.text)
    return out


def _element_text(el: ET.Element) -> str:
    return "".join(t.text or "" for t in el.iter() if t.tag.endswith("}t")).strip()


def _twips(value: str | None) -> int | None:
    try:
        return int(value) if value is not None else None
    except (TypeError, ValueError):
        return None


def _cell_width_twips(cell: ET.Element, grid_width: int | None) -> int | None:
    tc_w = cell.find(".//w:tcW", NS)
    return _twips(tc_w.get(f"{{{W_NS}}}w")) if tc_w is not None else grid_width


def _header_min_width_twips(text: str) -> int | None:
    words = re.findall(r"[A-Za-z0-9][A-Za-z0-9/-]*", text)
    longest = max((len(word) for word in words), default=0)
    if longest < 4:
        return None
    return max(900, min(2200, longest * 155))


def _short_cell_min_width_twips(text: str) -> int | None:
    clean = " ".join(text.split())
    if len(clean) > 64 or "://" in clean:
        return None
    if PRICE_TOKEN_RE.match(clean):
        return 1050
    words = re.findall(r"[A-Za-z0-9][A-Za-z0-9/-]*", clean)
    longest = max((len(word) for word in words), default=0)
    if longest < 6:
        return None
    return max(1050, min(1900, longest * 165))


def _row_has_cant_split(row: ET.Element) -> bool:
    return row.find("./w:trPr/w:cantSplit", NS) is not None


def _row_has_keep_next(row: ET.Element) -> bool:
    return row.find(".//w:pPr/w:keepNext", NS) is not None


def _inspect_docx_table_layout(report: dict, xml: str) -> None:
    try:
        root = ET.fromstring(xml)
    except ET.ParseError:
        return
    for table_idx, table in enumerate(root.findall(".//w:tbl", NS), 1):
        grid_widths = [
            _twips(col.get(f"{{{W_NS}}}w"))
            for col in table.findall("./w:tblGrid/w:gridCol", NS)
        ]
        first_row = table.find("./w:tr", NS)
        if first_row is None:
            continue
        cells = first_row.findall("./w:tc", NS)
        table_rows = table.findall("./w:tr", NS)
        if len(table_rows) == 1 and len(cells) == 1:
            row_text = _element_text(table_rows[0])
            if len(row_text) >= 80 and not _row_has_cant_split(table_rows[0]):
                report["warnings"].append(
                    _issue(
                        "warning",
                        "callout_row_can_split",
                        f"One-cell callout/table {table_idx} can split across pages; keep it as one visual block",
                    )
                )
        total_table_text = sum(len(_element_text(row)) for row in table_rows)
        prose_table_risk = 4 <= len(table_rows) <= 6 and total_table_text >= 420
        if prose_table_risk and not all(_row_has_keep_next(row) for row in table_rows[:-1]):
            report["warnings"].append(
                _issue(
                    "warning",
                    "table_orphan_split_risk",
                    f"Table {table_idx} is a compact prose table likely to orphan rows across pages; keep rows together or start it earlier",
                )
            )
        for row_idx, row in enumerate(table_rows, 1):
            row_text = _element_text(row)
            split_risk = len(table_rows) >= 4 and len(row_text) >= 80
            if split_risk and not _row_has_cant_split(row):
                report["warnings"].append(
                    _issue(
                        "warning",
                        "table_row_can_split",
                        f"Table {table_idx} row {row_idx} can split across pages; this can break row readability",
                    )
                )
        for col_idx, cell in enumerate(cells, 1):
            text = _element_text(cell)
            min_width = _header_min_width_twips(text)
            if min_width is None:
                continue
            grid_width = grid_widths[col_idx - 1] if col_idx <= len(grid_widths) else None
            width = _cell_width_twips(cell, grid_width)
            if width is not None and width < min_width:
                report["errors"].append(
                    _issue(
                        "error",
                        "table_header_narrow",
                        (
                            f"Table {table_idx} header column {col_idx} is too narrow for "
                            f"'{text}' ({width} twips < {min_width} twips); likely visible word wrap"
                        ),
                    )
                )
        for row_idx, row in enumerate(table_rows[1:], 2):
            body_cells = row.findall("./w:tc", NS)
            for col_idx, cell in enumerate(body_cells, 1):
                text = _element_text(cell)
                min_width = _short_cell_min_width_twips(text)
                if min_width is None:
                    continue
                grid_width = grid_widths[col_idx - 1] if col_idx <= len(grid_widths) else None
                width = _cell_width_twips(cell, grid_width)
                if width is not None and width < min_width:
                    report["errors"].append(
                        _issue(
                            "error",
                            "table_body_word_narrow",
                            (
                                f"Table {table_idx} row {row_idx} column {col_idx} is too narrow for "
                                f"'{text}' ({width} twips < {min_width} twips); likely visible word break"
                            ),
                        )
                    )


def _natural_slide_key(name: str) -> tuple[int, str]:
    m = re.search(r"slide(\d+)\.xml$", name)
    return (int(m.group(1)) if m else 0, name)


def _inspect_pptx_taste(report: dict, slide_texts: list[str]) -> None:
    """Add lightweight Impeccable/Taste-inspired text warnings for PPTX files."""
    repeated = Counter(t.lower().strip() for t in slide_texts if t.strip())
    if any(count > 1 and len(text) > 20 for text, count in repeated.items()):
        report["warnings"].append(
            _issue("warning", "duplicate_slide_text", "Repeated slide text suggests templated or duplicated content")
        )

    for idx, text in enumerate(slide_texts, 1):
        if AI_TELL_RE.search(text):
            report["warnings"].append(
                _issue("warning", "generic_ai_copy", f"Slide {idx} contains generic AI/marketing copy")
            )
        if DECORATIVE_DASH_RE.search(text):
            report["warnings"].append(
                _issue("warning", "dash_tell", f"Slide {idx} contains em/en dash characters")
            )


def inspect_pptx(path: Path) -> dict:
    report = {"path": str(path), "kind": "pptx", "errors": [], "warnings": [], "metrics": {}}
    try:
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
            if "ppt/presentation.xml" not in names:
                report["errors"].append(_issue("error", "missing_presentation_xml", "Missing ppt/presentation.xml"))
                return report
            slide_names = sorted(
                [n for n in names if re.match(r"ppt/slides/slide\d+\.xml$", n)],
                key=_natural_slide_key,
            )
            report["metrics"]["slide_count"] = len(slide_names)
            if not slide_names:
                report["errors"].append(_issue("error", "empty_deck", "No slide XML files found"))
                return report

            total_text = 0
            slide_texts = []
            for idx, slide_name in enumerate(slide_names, 1):
                text_runs = _xml_text(_read_zip_text(zf, slide_name))
                slide_text = " ".join(t.strip() for t in text_runs if t.strip())
                slide_texts.append(slide_text)
                total_text += len(slide_text)
                if len(slide_text) > PPT_SLIDE_TEXT_LIMIT:
                    report["warnings"].append(
                        _issue("warning", "dense_slide", f"Slide {idx} has {len(slide_text)} text characters")
                    )
                for run in text_runs:
                    clean = run.strip()
                    if len(clean) > LONG_TEXT_LIMIT:
                        report["warnings"].append(
                            _issue("warning", "long_text_run", f"Slide {idx} contains a long text run ({len(clean)} chars)")
                        )
                    if PLACEHOLDER_RE.search(clean):
                        report["warnings"].append(
                            _issue("warning", "placeholder_text", f"Slide {idx} contains placeholder-like text: {clean[:80]}")
                        )
            report["metrics"]["total_text_chars"] = total_text
            _inspect_pptx_taste(report, slide_texts)
    except zipfile.BadZipFile:
        report["errors"].append(_issue("error", "bad_zip", "File is not a valid PPTX zip package"))
    return report


def _run_self_audit_on_templates(args: argparse.Namespace) -> None:
    """Run the template-level self-audit and exit non-zero on any failure.

    Called only when ``--with-self-audit`` is set. The audit results are
    printed so the operator can see which template regressed, and any
    ``fail`` verdict terminates the gate with exit code 1 before the
    per-file DOCX inspection even starts (the v1.6.6 防复发 contract:
    a forbidden template must not be allowed to ship a single file).

    The function writes its verdict to stderr, not stdout, so the JSON
    report (if ``--json-report`` is also set) is not polluted.
    """
    try:
        import importlib

        try:
            self_audit = importlib.import_module("pro_docx_gen.self_audit")
        except ModuleNotFoundError:
            skill_root = Path(__file__).resolve().parents[1]
            if str(skill_root) not in sys.path:
                sys.path.insert(0, str(skill_root))
            self_audit = importlib.import_module("pro_docx_gen.self_audit")
    except Exception as exc:  # pragma: no cover - defensive guard
        print(
            f"ERROR self_audit_unavailable: cannot import pro_docx_gen.self_audit: {exc}",
            file=sys.stderr,
        )
        sys.exit(2)

    reports = self_audit.audit_all_templates()
    fail_count = 0
    for r in reports:
        marker = "PASS" if r.status == "pass" else "FAIL"
        if r.status != "pass":
            fail_count += 1
        print(f"  [self-audit] [{marker}] {r.name}: {len(r.violations)} violation(s)", file=sys.stderr)
        for v in r.violations:
            print(f"    - {v.field}={v.hex_value} L{v.line}: {v.description}", file=sys.stderr)
    if fail_count:
        print(
            f"ERROR self_audit_failed: {fail_count}/{len(reports)} template(s) failed the v1.6.6 colour audit",
            file=sys.stderr,
        )
        sys.exit(1)


def _heading_color_audit(docx_path: Path) -> list[dict]:
    """Return ``Violation``-shaped dicts for the v1.6.6 heading colour gate.

    The implementation reuses ``pro_docx_gen.self_audit.audit_docx_output``
    to avoid duplicating the XML walk / hex regex logic. The function
    is robust to ``self_audit`` not being importable: in that case it
    returns an empty list and the gate is a no-op, which keeps
    ``run_quality_gate.py`` runnable from a clean checkout (e.g. when
    the ``pro_docx_gen`` package is not yet on ``PYTHONPATH``).
    """
    try:
        # Make the import lazy so the gate still works when the package
        # is not installed; e.g. when running from a checked-out zip.
        import importlib

        try:
            self_audit = importlib.import_module("pro_docx_gen.self_audit")
        except ModuleNotFoundError:
            # Fallback: add the skill root to sys.path and retry once.
            skill_root = Path(__file__).resolve().parents[1]
            if str(skill_root) not in sys.path:
                sys.path.insert(0, str(skill_root))
            self_audit = importlib.import_module("pro_docx_gen.self_audit")
    except Exception as exc:  # pragma: no cover - defensive guard
        return [_issue(
            "warning",
            "self_audit_unavailable",
            f"pro_docx_gen.self_audit could not be imported; skipping heading colour check: {exc}",
        )]

    try:
        report = self_audit.audit_docx_output(str(docx_path))
    except Exception as exc:  # pragma: no cover - defensive guard
        return [_issue(
            "warning",
            "self_audit_failed",
            f"self_audit.audit_docx_output raised an exception: {exc}",
        )]

    issues: list[dict] = []
    for v in report.violations:
        # Use the same _issue() factory so the report printer formats
        # the message identically. code + level are derived from the
        # violation's description/field.
        if v.field.startswith("Heading") or v.field == "Title":
            code = "heading_color_inconsistent"
        elif v.field in ("<file>", "<zip>", "<xml>"):
            code = "self_audit_io_error"
        else:
            code = "self_audit_violation"
        issues.append(_issue("error", code, v.description))
    return issues


def check_heading_color_consistency(docx_path: Path) -> list[dict]:
    """Public entry: heading colour consistency check for the quality gate.

    Wraps :func:`_heading_color_audit` and applies a small extra safety net
    so that the gate is **always** informative even if the self_audit
    module changes its public surface: the function returns a list of
    issue-dicts (never raises).
    """
    try:
        return _heading_color_audit(docx_path)
    except Exception as exc:  # pragma: no cover - defensive guard
        return [_issue(
            "error",
            "heading_color_check_crashed",
            f"check_heading_color_consistency crashed: {exc}",
        )]


def inspect_docx(path: Path) -> dict:
    report = {"path": str(path), "kind": "docx", "errors": [], "warnings": [], "metrics": {}}
    try:
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
            if "word/document.xml" not in names:
                report["errors"].append(_issue("error", "missing_document_xml", "Missing word/document.xml"))
                return report
            xml = _read_zip_text(zf, "word/document.xml")
            texts = _xml_text(xml)
            body_text = " ".join(t.strip() for t in texts if t.strip())
            report["metrics"]["text_chars"] = len(body_text)
            report["metrics"]["table_count"] = xml.count("<w:tbl>")
            report["metrics"]["image_count"] = xml.count("<w:drawing>")
            if not body_text:
                report["errors"].append(_issue("error", "empty_document", "No visible document text found"))
            for text in texts:
                clean = text.strip()
                if len(clean) > DOCX_PARAGRAPH_LIMIT:
                    report["warnings"].append(
                        _issue("warning", "long_text_run", f"DOCX contains a long text run ({len(clean)} chars)")
                    )
                if PLACEHOLDER_RE.search(clean):
                    report["warnings"].append(
                        _issue("warning", "placeholder_text", f"DOCX contains placeholder-like text: {clean[:80]}")
                    )
            tbl_count = xml.count("<w:tbl>")
            grid_count = xml.count("<w:tblGrid>")
            if tbl_count and grid_count < tbl_count:
                report["warnings"].append(
                    _issue("warning", "table_geometry", f"{tbl_count - grid_count} table(s) missing explicit tblGrid")
                )
            _inspect_docx_table_layout(report, xml)
    except zipfile.BadZipFile:
        report["errors"].append(_issue("error", "bad_zip", "File is not a valid DOCX zip package"))
    # v1.6.6: heading colour consistency check. Runs after the structural
    # checks so the report is built up in pipeline order, mirroring the
    # agent's mental model: "first the docx is well-formed, then the
    # headings are colour-consistent, then we may render it."
    if not any(err.get("code") in ("bad_zip", "missing_document_xml") for err in report["errors"]):
        for issue in check_heading_color_consistency(path):
            level = issue.get("level", "warning")
            if level == "error":
                report["errors"].append(issue)
            else:
                report["warnings"].append(issue)
    return report


def _has_auto_fixable_issue(report: dict) -> bool:
    issues = report.get("errors", []) + report.get("warnings", [])
    return report.get("kind") == "docx" and any(issue.get("code") in AUTO_FIXABLE_CODES for issue in issues)


def _set_docx_cell_width(cell, width_twips: int) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Twips

    cell.width = Twips(width_twips)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_twips))


def _set_docx_row_no_split(row) -> bool:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is not None:
        return False
    tr_pr.append(OxmlElement("w:cantSplit"))
    return True


def _set_docx_row_keep_next(row) -> int:
    changes = 0
    for cell in row.cells:
        for para in cell.paragraphs:
            if not para.paragraph_format.keep_with_next:
                para.paragraph_format.keep_with_next = True
                changes += 1
    return changes


def auto_fix_docx_layout(path: Path) -> dict:
    from docx import Document

    doc = Document(path)
    changes = 0
    for table in doc.tables:
        if not table.rows:
            continue
        table.autofit = False
        if len(table.rows) == 1 and len(table.columns) == 1:
            row_text = " ".join(cell.text for cell in table.rows[0].cells)
            if len(row_text) >= 80 and _set_docx_row_no_split(table.rows[0]):
                changes += 1
        total_table_text = sum(len(cell.text) for row in table.rows for cell in row.cells)
        if 4 <= len(table.rows) <= 6 and total_table_text >= 420:
            for row in table.rows[:-1]:
                changes += _set_docx_row_keep_next(row)
        if len(table.rows) >= 4:
            for row in table.rows:
                row_text = " ".join(cell.text for cell in row.cells)
                if len(row_text) >= 80 and _set_docx_row_no_split(row):
                    changes += 1
        for col_idx, cell in enumerate(table.rows[0].cells):
            min_width = _header_min_width_twips(cell.text.strip())
            if min_width is None:
                continue
            current = getattr(cell.width, "twips", None) or 0
            if current >= min_width:
                continue
            for row in table.rows:
                if col_idx < len(row.cells):
                    _set_docx_cell_width(row.cells[col_idx], min_width)
            changes += 1
        required_by_col: dict[int, int] = {}
        for row in table.rows[1:]:
            for col_idx, cell in enumerate(row.cells):
                min_width = _short_cell_min_width_twips(cell.text.strip())
                if min_width is not None:
                    required_by_col[col_idx] = max(required_by_col.get(col_idx, 0), min_width)
        for col_idx, min_width in required_by_col.items():
            current = getattr(table.rows[0].cells[col_idx].width, "twips", None) or 0
            if current >= min_width:
                continue
            for row in table.rows:
                if col_idx < len(row.cells):
                    _set_docx_cell_width(row.cells[col_idx], min_width)
            changes += 1
    if changes:
        doc.save(path)
    return {"applied": bool(changes), "changes": changes}


def _office_candidates() -> list[str]:
    names = ("soffice", "libreoffice")
    paths = []
    for name in names:
        found = shutil.which(name)
        if found and found not in paths:
            paths.append(found)
    if os.name == "nt":
        for candidate in (
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ):
            if Path(candidate).exists() and candidate not in paths:
                paths.append(candidate)
    return paths


def _render_with_wps_com(path: Path, out_dir: Path) -> dict:
    powershell = shutil.which("powershell") or shutil.which("pwsh")
    if os.name != "nt" or not powershell:
        return {"status": "skipped", "reason": "WPS COM unavailable"}

    pdf_path = out_dir / f"{path.stem}.pdf"
    script = out_dir / "_wps_export_pdf.ps1"
    script.write_text(
        """
param([string]$DocPath, [string]$PdfPath)
$ErrorActionPreference = 'Stop'
$wps = New-Object -ComObject KWPS.Application
$wps.Visible = $false
$wps.DisplayAlerts = 0
try {
    $doc = $wps.Documents.Open($DocPath, $false, $true)
    try {
        $doc.ExportAsFixedFormat($PdfPath, 17)
    } finally {
        $doc.Close($false)
    }
} finally {
    $wps.Quit()
}
""".strip(),
        encoding="utf-8",
    )
    try:
        proc = subprocess.run(
            [powershell, "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(script), str(path), str(pdf_path)],
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            timeout=60,
        )
    except subprocess.TimeoutExpired:
        return {"status": "failed", "engine": "wps_com", "reason": "WPS COM export timed out"}
    if proc.returncode == 0 and pdf_path.exists() and pdf_path.stat().st_size > 1000:
        return {"status": "passed", "pdf": str(pdf_path), "engine": "wps_com"}
    return {
        "status": "failed",
        "engine": "wps_com",
        "returncode": proc.returncode,
        "stdout": (proc.stdout or "")[-2000:],
        "stderr": (proc.stderr or "")[-2000:],
    }


def _render_with_word_com(path: Path, out_dir: Path) -> dict:
    powershell = shutil.which("powershell") or shutil.which("pwsh")
    if os.name != "nt" or not powershell:
        return {"status": "skipped", "reason": "Word COM unavailable"}

    pdf_path = out_dir / f"{path.stem}.pdf"
    script = out_dir / "_word_export_pdf.ps1"
    script.write_text(
        """
param([string]$DocPath, [string]$PdfPath)
$ErrorActionPreference = 'Stop'
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {
    $doc = $word.Documents.Open($DocPath, $false, $true)
    try {
        $doc.ExportAsFixedFormat($PdfPath, 17)
    } finally {
        $doc.Close($false)
    }
} finally {
    $word.Quit()
}
""".strip(),
        encoding="utf-8",
    )
    proc = subprocess.run(
        [powershell, "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(script), str(path), str(pdf_path)],
        text=True,
        encoding="utf-8",
        errors="replace",
        capture_output=True,
    )
    if proc.returncode == 0 and pdf_path.exists() and pdf_path.stat().st_size > 1000:
        return {"status": "passed", "pdf": str(pdf_path), "engine": "word_com"}
    return {
        "status": "failed",
        "engine": "word_com",
        "returncode": proc.returncode,
        "stdout": (proc.stdout or "")[-2000:],
        "stderr": (proc.stderr or "")[-2000:],
    }


def _selective_export_script(application_progid: str) -> str:
    """Build a Word-compatible COM script that exports sampled pages only."""
    return f"""
param([string]$DocPath, [string]$OutDir, [int]$GroupSize)
$ErrorActionPreference = 'Stop'
if ($GroupSize -lt 1) {{ throw 'GroupSize must be at least 1' }}
$app = New-Object -ComObject {application_progid}
$app.Visible = $false
try {{ $app.DisplayAlerts = 0 }} catch {{}}
try {{
    $doc = $app.Documents.Open($DocPath, $false, $true)
    try {{
        $pageCount = [int]$doc.ComputeStatistics(2)
        if ($pageCount -lt 1) {{ throw 'Document page count is zero' }}
        $pages = @()
        $pdfs = @()
        for ($page = 1; $page -le $pageCount; $page += $GroupSize) {{
            $pdfPath = Join-Path $OutDir (([IO.Path]::GetFileNameWithoutExtension($DocPath)) + ('-page-{{0:D4}}.pdf' -f $page))
            $doc.ExportAsFixedFormat($pdfPath, 17, $false, 0, 3, $page, $page, 0, $true, $true, 0, $true, $true, $false)
            $pages += $page
            $pdfs += $pdfPath
        }}
        [PSCustomObject]@{{page_count=$pageCount; sampled_page_numbers=$pages; sampled_pdfs=$pdfs}} | ConvertTo-Json -Compress
    }} finally {{
        $doc.Close($false)
    }}
}} finally {{
    try {{ $app.Quit() }} catch {{}}
}}
""".strip()


def _targeted_export_script(application_progid: str) -> str:
    """Build a Word-compatible COM script that exports explicit review pages only."""
    return f"""
param([string]$DocPath, [string]$OutDir, [string]$PagesCsv)
$ErrorActionPreference = 'Stop'
$pages = @($PagesCsv.Split(',') | ForEach-Object {{ [int]$_.Trim() }} | Sort-Object -Unique)
if ($pages.Count -lt 1) {{ throw 'At least one review page is required' }}
$app = New-Object -ComObject {application_progid}
$app.Visible = $false
try {{ $app.DisplayAlerts = 0 }} catch {{}}
try {{
    $doc = $app.Documents.Open($DocPath, $false, $true)
    try {{
        $pageCount = [int]$doc.ComputeStatistics(2)
        if ($pageCount -lt 1) {{ throw 'Document page count is zero' }}
        $pdfs = @()
        foreach ($page in $pages) {{
            if ($page -lt 1 -or $page -gt $pageCount) {{ throw "Review page $page is outside 1..$pageCount" }}
            $pdfPath = Join-Path $OutDir ('target-page-{{0:D4}}.pdf' -f $page)
            $doc.ExportAsFixedFormat($pdfPath, 17, $false, 0, 3, $page, $page, 0, $true, $true, 0, $true, $true, $false)
            $pdfs += $pdfPath
        }}
        [PSCustomObject]@{{page_count=$pageCount; sampled_page_numbers=$pages; sampled_pdfs=$pdfs}} | ConvertTo-Json -Compress
    }} finally {{
        $doc.Close($false)
    }}
}} finally {{
    try {{ $app.Quit() }} catch {{}}
}}
""".strip()


def _render_sampled_with_com(
    path: Path,
    out_dir: Path,
    group_size: int,
    application_progid: str,
    engine: str,
) -> dict:
    powershell = shutil.which("powershell") or shutil.which("pwsh")
    if os.name != "nt" or not powershell:
        return {"status": "skipped", "engine": engine, "reason": "COM automation unavailable"}
    if group_size < 1:
        return {"status": "failed", "engine": engine, "reason": "sample group size must be at least 1"}

    out_dir.mkdir(parents=True, exist_ok=True)
    script = out_dir / f"_{engine}_export_sampled_pages.ps1"
    script.write_text(_selective_export_script(application_progid), encoding="utf-8")
    try:
        proc = subprocess.run(
            [
                powershell,
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(script),
                str(path.resolve()),
                str(out_dir.resolve()),
                str(group_size),
            ],
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            timeout=120,
        )
    except subprocess.TimeoutExpired:
        return {"status": "failed", "engine": engine, "reason": "selective COM export timed out"}

    payload = None
    if proc.returncode == 0:
        for line in reversed((proc.stdout or "").splitlines()):
            try:
                payload = json.loads(line)
                break
            except json.JSONDecodeError:
                continue
    if payload:
        page_count = int(payload.get("page_count") or 0)
        page_numbers = [int(page) for page in payload.get("sampled_page_numbers") or []]
        sampled_pdfs = [str(Path(pdf)) for pdf in payload.get("sampled_pdfs") or []]
        expected_pages = select_sample_page_numbers(page_count, group_size) if page_count else []
        valid_pdfs = [pdf for pdf in sampled_pdfs if Path(pdf).exists() and Path(pdf).stat().st_size > 1000]
        if page_numbers == expected_pages and len(valid_pdfs) == len(expected_pages):
            return {
                "status": "passed",
                "engine": engine,
                "page_count": page_count,
                "sample_group_size": group_size,
                "sampled_page_numbers": page_numbers,
                "sampled_pdfs": valid_pdfs,
            }
    return {
        "status": "failed",
        "engine": engine,
        "reason": "selective COM page export failed",
        "returncode": proc.returncode,
        "stdout": (proc.stdout or "")[-2000:],
        "stderr": (proc.stderr or "")[-2000:],
    }


def _render_sampled_with_wps_com(path: Path, out_dir: Path, group_size: int = 3) -> dict:
    return _render_sampled_with_com(path, out_dir, group_size, "KWPS.Application", "wps_com_selective")


def _render_sampled_with_word_com(path: Path, out_dir: Path, group_size: int = 3) -> dict:
    return _render_sampled_with_com(path, out_dir, group_size, "Word.Application", "word_com_selective")


def _render_targeted_with_com(
    path: Path,
    out_dir: Path,
    page_numbers: list[int],
    application_progid: str,
    engine: str,
) -> dict:
    powershell = shutil.which("powershell") or shutil.which("pwsh")
    if os.name != "nt" or not powershell:
        return {"status": "skipped", "engine": engine, "reason": "COM automation unavailable"}
    requested = sorted({int(page) for page in page_numbers})
    if not requested or any(page < 1 for page in requested):
        return {"status": "failed", "engine": engine, "reason": "review pages must be positive integers"}

    out_dir.mkdir(parents=True, exist_ok=True)
    script = out_dir / f"_{engine}_export_targeted_pages.ps1"
    script.write_text(_targeted_export_script(application_progid), encoding="utf-8")
    try:
        proc = subprocess.run(
            [
                powershell,
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                str(script),
                str(path.resolve()),
                str(out_dir.resolve()),
                ",".join(str(page) for page in requested),
            ],
            text=True,
            encoding="utf-8",
            errors="replace",
            capture_output=True,
            timeout=120,
        )
    except subprocess.TimeoutExpired:
        return {"status": "failed", "engine": engine, "reason": "targeted COM export timed out"}

    payload = None
    if proc.returncode == 0:
        for line in reversed((proc.stdout or "").splitlines()):
            try:
                payload = json.loads(line)
                break
            except json.JSONDecodeError:
                continue
    if payload:
        page_count = int(payload.get("page_count") or 0)
        rendered_pages = [int(page) for page in payload.get("sampled_page_numbers") or []]
        sampled_pdfs = [str(Path(pdf)) for pdf in payload.get("sampled_pdfs") or []]
        valid_pdfs = [pdf for pdf in sampled_pdfs if Path(pdf).exists() and Path(pdf).stat().st_size > 1000]
        if rendered_pages == requested and len(valid_pdfs) == len(requested):
            return {
                "status": "passed",
                "engine": engine,
                "page_count": page_count,
                "sampled_page_numbers": rendered_pages,
                "sampled_pdfs": valid_pdfs,
            }
    return {
        "status": "failed",
        "engine": engine,
        "reason": "targeted COM page export failed",
        "returncode": proc.returncode,
        "stdout": (proc.stdout or "")[-2000:],
        "stderr": (proc.stderr or "")[-2000:],
    }


def _render_targeted_with_wps_com(path: Path, out_dir: Path, page_numbers: list[int]) -> dict:
    return _render_targeted_with_com(path, out_dir, page_numbers, "KWPS.Application", "wps_com_targeted")


def _render_targeted_with_word_com(path: Path, out_dir: Path, page_numbers: list[int]) -> dict:
    return _render_targeted_with_com(path, out_dir, page_numbers, "Word.Application", "word_com_targeted")


def try_render_specific_pages(path: Path, output_dir: Path, page_numbers: list[int]) -> dict:
    """Render explicit source pages; callers may pass every page for figure-heavy QA."""
    attempts: list[dict] = []
    for backend in (_render_targeted_with_wps_com, _render_targeted_with_word_com):
        result = backend(path, output_dir, page_numbers)
        attempts.append(dict(result))
        if result.get("status") == "passed":
            result["attempts"] = attempts
            return result
    return {
        "status": "failed",
        "reason": "no targeted page render backend succeeded",
        "attempts": attempts,
    }


def try_render_sampled(path: Path, output_dir: Path, group_size: int = 3) -> dict:
    """Render only sampled source pages; never fall back to full-document export."""
    attempts: list[dict] = []
    for backend in (_render_sampled_with_wps_com, _render_sampled_with_word_com):
        result = backend(path, output_dir, group_size)
        attempts.append(dict(result))
        if result.get("status") == "passed":
            result["attempts"] = attempts
            return result
    return {
        "status": "failed",
        "reason": "no selective page render backend succeeded",
        "attempts": attempts,
    }


def _render_with_office_cli(path: Path, out_dir: Path, office: str) -> dict:
    cmd = [
        office,
        "--headless",
        "--convert-to",
        "pdf",
        "--outdir",
        str(out_dir),
        str(path),
    ]
    proc = subprocess.run(cmd, text=True, capture_output=True)
    pdf_path = out_dir / f"{path.stem}.pdf"
    if proc.returncode == 0 and pdf_path.exists() and pdf_path.stat().st_size > 1000:
        return {"status": "passed", "pdf": str(pdf_path), "engine": Path(office).name}
    return {
        "status": "failed",
        "engine": Path(office).name,
        "returncode": proc.returncode,
        "stdout": proc.stdout[-2000:],
        "stderr": proc.stderr[-2000:],
    }


def try_render(path: Path, output_dir: Path | None) -> dict:
    out_dir = output_dir or Path.cwd() / "_work" / f"pro_quality_render_{os.getpid()}"
    out_dir.mkdir(parents=True, exist_ok=True)
    attempts: list[dict] = []

    for office in _office_candidates():
        result = _render_with_office_cli(path, out_dir, office)
        attempts.append(dict(result))
        if result.get("status") == "passed":
            result["attempts"] = attempts
            return result

    for backend in (_render_with_wps_com, _render_with_word_com):
        result = backend(path, out_dir)
        attempts.append(dict(result))
        if result.get("status") == "passed":
            result["attempts"] = attempts
            return result
    return {"status": "failed", "reason": "no PDF render backend succeeded", "attempts": attempts}


def _render_pdf_previews_with_pdftoppm(pdf_path: Path, out_dir: Path, max_pages: int) -> dict:
    pdftoppm = _find_pdftoppm()
    if not pdftoppm:
        return {"status": "skipped", "reason": "pdftoppm not found"}
    preview_dir = out_dir / "previews"
    preview_dir.mkdir(parents=True, exist_ok=True)
    prefix = preview_dir / pdf_path.stem
    cmd = [
        pdftoppm,
        "-png",
        "-r",
        "144",
        "-f",
        "1",
        "-l",
        str(max_pages),
        str(pdf_path),
        str(prefix),
    ]
    proc = subprocess.run(cmd, text=True, encoding="utf-8", errors="replace", capture_output=True)
    pngs = sorted(preview_dir.glob(f"{pdf_path.stem}-*.png"))
    valid_pngs = [p for p in pngs if p.stat().st_size > 1000]
    if proc.returncode == 0 and valid_pngs:
        return {"status": "passed", "pngs": [str(p) for p in valid_pngs], "engine": "pdftoppm"}
    return {
        "status": "failed",
        "reason": "PNG preview generation failed",
        "returncode": proc.returncode,
        "stdout": (proc.stdout or "")[-2000:],
        "stderr": (proc.stderr or "")[-2000:],
    }


def _find_pdftoppm() -> str | None:
    found = shutil.which("pdftoppm")
    candidates = [Path(found)] if found else []
    for parent in Path(sys.executable).resolve().parents:
        candidates.append(parent / "native" / "poppler" / "Library" / "bin" / "pdftoppm.exe")
        candidates.append(parent / "dependencies" / "native" / "poppler" / "Library" / "bin" / "pdftoppm.exe")
    for candidate in candidates:
        if candidate and candidate.exists() and candidate.suffix.lower() == ".exe":
            return str(candidate)
    return found


def _find_poppler_bin() -> str | None:
    pdftoppm = _find_pdftoppm()
    if pdftoppm:
        return str(Path(pdftoppm).resolve().parent)
    return None


def _combine_preview_pngs(pngs: list[str], out_dir: Path, stem: str) -> str | None:
    if not pngs:
        return None
    try:
        from PIL import Image, ImageOps
    except ImportError:
        return None

    images = []
    try:
        for png in pngs:
            image = Image.open(png).convert("RGB")
            image.thumbnail((900, 1200))
            images.append(ImageOps.expand(image, border=16, fill="white"))
        if not images:
            return None
        gap = 24
        width = max(img.width for img in images)
        height = sum(img.height for img in images) + gap * (len(images) - 1)
        sheet = Image.new("RGB", (width, height), "white")
        y = 0
        for img in images:
            x = (width - img.width) // 2
            sheet.paste(img, (x, y))
            y += img.height + gap
        out_path = out_dir / f"{stem}.submission_preview.png"
        sheet.save(out_path, "PNG", optimize=True)
        return str(out_path) if out_path.stat().st_size > 1000 else None
    finally:
        for image in images:
            image.close()


def _render_pdf_previews_with_pymupdf(pdf_path: Path, out_dir: Path, max_pages: int) -> dict:
    try:
        import fitz  # type: ignore
    except ImportError:
        return {"status": "skipped", "reason": "PyMuPDF/fitz not installed"}
    preview_dir = out_dir / "previews"
    preview_dir.mkdir(parents=True, exist_ok=True)
    pngs = []
    try:
        doc = fitz.open(str(pdf_path))
        try:
            for index in range(min(max_pages, len(doc))):
                page = doc.load_page(index)
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                out = preview_dir / f"{pdf_path.stem}-{index + 1}.png"
                pix.save(str(out))
                if out.stat().st_size > 1000:
                    pngs.append(str(out))
        finally:
            doc.close()
    except Exception as exc:
        return {"status": "failed", "reason": f"PyMuPDF preview failed: {exc}"}
    if pngs:
        return {"status": "passed", "pngs": pngs, "engine": "pymupdf"}
    return {"status": "failed", "reason": "PyMuPDF produced no PNG previews"}


def _render_pdf_previews_with_pdf2image(pdf_path: Path, out_dir: Path, max_pages: int) -> dict:
    try:
        from pdf2image import convert_from_path  # type: ignore
    except ImportError:
        return {"status": "skipped", "reason": "pdf2image not installed"}
    preview_dir = out_dir / "previews"
    preview_dir.mkdir(parents=True, exist_ok=True)
    try:
        images = convert_from_path(
            str(pdf_path),
            dpi=144,
            first_page=1,
            last_page=max_pages,
            poppler_path=_find_poppler_bin(),
        )
        pngs = []
        for index, image in enumerate(images, 1):
            out = preview_dir / f"{pdf_path.stem}-{index}.png"
            image.save(out, "PNG")
            if out.stat().st_size > 1000:
                pngs.append(str(out))
    except Exception as exc:
        return {"status": "failed", "reason": f"pdf2image preview failed: {exc}"}
    if pngs:
        return {"status": "passed", "pngs": pngs, "engine": "pdf2image"}
    return {"status": "failed", "reason": "pdf2image produced no PNG previews"}


def _render_pdf_previews(pdf_path: Path, out_dir: Path, max_pages: int = 200) -> dict:
    attempts = []
    for backend in (
        _render_pdf_previews_with_pdftoppm,
        _render_pdf_previews_with_pymupdf,
        _render_pdf_previews_with_pdf2image,
    ):
        result = backend(pdf_path, out_dir, max_pages)
        attempts.append(dict(result))
        if result.get("status") == "passed":
            page_pngs = list(result.get("pngs") or [])
            result["attempts"] = attempts
            result["page_count_previewed"] = len(page_pngs)
            result["page_pngs"] = page_pngs
            result["pngs"] = page_pngs
            return result
    return {"status": "failed", "reason": "no PNG preview backend succeeded", "attempts": attempts}


def _render_sampled_pdf_previews(render_result: dict, out_dir: Path) -> dict:
    """Rasterize one-page sampled PDFs while retaining their source page numbers."""
    page_numbers = [int(page) for page in render_result.get("sampled_page_numbers") or []]
    sampled_pdfs = [str(path) for path in render_result.get("sampled_pdfs") or []]
    if not page_numbers or len(page_numbers) != len(sampled_pdfs):
        return {"status": "failed", "reason": "sampled page metadata is incomplete"}

    page_pngs: list[str] = []
    attempts: list[dict] = []
    alias_dir = out_dir / "sampled_previews"
    alias_dir.mkdir(parents=True, exist_ok=True)
    for page_number, pdf in zip(page_numbers, sampled_pdfs):
        source_pdf = Path(pdf)
        short_pdf = alias_dir / f"p{page_number:04d}.pdf"
        shutil.copyfile(source_pdf, short_pdf)
        result = _render_pdf_previews(short_pdf, alias_dir, max_pages=1)
        attempts.append({"source_page": page_number, **dict(result)})
        generated = list(result.get("page_pngs") or result.get("pngs") or [])
        if result.get("status") != "passed" or len(generated) != 1:
            return {
                "status": "failed",
                "reason": f"PNG preview failed for sampled source page {page_number}",
                "attempts": attempts,
            }
        page_pngs.append(str(generated[0]))
    return {
        "status": "passed",
        "engine": "sampled_single_page_pdfs",
        "page_numbers": page_numbers,
        "page_pngs": page_pngs,
        "pngs": page_pngs,
        "page_count_previewed": len(page_pngs),
        "attempts": attempts,
    }


def inspect_file(
    path: Path,
    render: bool,
    output_dir: Path | None,
    allow_render_skip: bool = False,
    preview_png: bool = True,
    created_after: float | None = None,
    sample_group_size: int = 3,
) -> dict:
    suffix = path.suffix.lower()
    if suffix == ".pptx":
        report = inspect_pptx(path)
    elif suffix == ".docx":
        report = inspect_docx(path)
    else:
        report = {"path": str(path), "kind": "unknown", "errors": [], "warnings": [], "metrics": {}}
        report["errors"].append(_issue("error", "unsupported_file", f"Unsupported file type: {suffix}"))
    if created_after is not None:
        try:
            mtime = path.stat().st_mtime
        except FileNotFoundError:
            mtime = 0
        report["metrics"]["mtime"] = mtime
        report["metrics"]["created_after"] = created_after
        if mtime <= created_after:
            report["errors"].append(
                _issue(
                    "error",
                    "stale_output",
                    (
                        "Output file was not modified after the recorded generation start time; "
                        "generation likely failed or quality gate is inspecting an old file"
                    ),
                )
            )
    if render and not report["errors"]:
        render_dir = output_dir / "renders" if output_dir else Path.cwd() / "_work" / f"pro_quality_render_{os.getpid()}"
        if suffix == ".docx":
            report["render"] = try_render_sampled(path, render_dir, sample_group_size)
        else:
            report["render"] = try_render(path, render_dir)
        if report["render"].get("status") == "failed":
            report["errors"].append(_issue("error", "render_failed", "Selective page rendering failed"))
        if report["render"].get("status") == "skipped" and not allow_render_skip:
            report["errors"].append(
                _issue("error", "render_skipped", "No DOCX-to-PDF backend was available; render verification was not run")
            )
        if preview_png and report["render"].get("status") == "passed":
            if suffix == ".docx":
                report["preview_png"] = _render_sampled_pdf_previews(report["render"], render_dir)
            else:
                pdf = Path(report["render"]["pdf"])
                report["preview_png"] = _render_pdf_previews(pdf, render_dir)
            if report["preview_png"].get("status") != "passed":
                report["errors"].append(
                    _issue(
                        "error",
                        "preview_png_failed",
                        report["preview_png"].get("reason", "PNG preview verification failed"),
                    )
                )
    return report


def main() -> int:
    parser = argparse.ArgumentParser(description="Run PRO-PPTX/PRO-DOCX delivery quality gates.")
    parser.add_argument("files", nargs="+", type=Path, help="PPTX/DOCX files to inspect")
    parser.add_argument("--json-report", type=Path, help="Write full JSON report to this path")
    parser.add_argument("--output-dir", type=Path, help="Directory for optional render outputs")
    parser.add_argument("--no-render", action="store_true", help="Skip optional LibreOffice PDF rendering")
    parser.add_argument("--no-preview-png", action="store_true", help="Diagnostic only: skip PNG preview generation")
    parser.add_argument("--auto-fix", action="store_true", help="Apply safe DOCX layout fixes, then re-run gates")
    parser.add_argument(
        "--sample-group-size",
        type=int,
        default=3,
        help="Render one DOCX source page from each consecutive page group (default: 3)",
    )
    parser.add_argument(
        "--created-after",
        type=_parse_timestamp,
        help="Fail if output file mtime is not newer than this Unix/ISO timestamp; prevents old-file false passes after generation failure",
    )
    parser.add_argument("--allow-warnings", action="store_true", help="Do not fail on warnings; diagnostic use only")
    parser.add_argument(
        "--allow-render-skip",
        action="store_true",
        help="Do not fail when LibreOffice/soffice is unavailable; diagnostic use only",
    )
    parser.add_argument(
        "--with-self-audit",
        action="store_true",
        help=(
            "Also run pro_docx_gen.self_audit against all 16 templates (template-level "
            "audit) and re-run the per-DOCX audit. Use this in CI to enforce that no "
            "template or output re-introduces a forbidden heading colour (e.g. #D4AF37)."
        ),
    )
    args = parser.parse_args()
    if args.sample_group_size < 1:
        parser.error("--sample-group-size must be at least 1")

    # v1.6.6: optional template-level self-audit. This catches regressions
    # at the *source* of the issue: a template that injects a forbidden hex
    # should not be allowed to ship, even if no DOCX has been generated yet.
    if args.with_self_audit:
        _run_self_audit_on_templates(args)

    reports = []
    for path in args.files:
        report = inspect_file(
            path,
            render=not args.no_render,
            output_dir=args.output_dir,
            allow_render_skip=args.allow_render_skip,
            preview_png=not args.no_preview_png,
            created_after=args.created_after,
            sample_group_size=args.sample_group_size,
        )
        if args.auto_fix and _has_auto_fixable_issue(report):
            before = [
                issue
                for issue in (report.get("errors", []) + report.get("warnings", []))
                if issue.get("code") in AUTO_FIXABLE_CODES
            ]
            fix_result = auto_fix_docx_layout(path)
            report = inspect_file(
                path,
                render=not args.no_render,
                output_dir=args.output_dir,
                allow_render_skip=args.allow_render_skip,
                preview_png=not args.no_preview_png,
                created_after=args.created_after,
                sample_group_size=args.sample_group_size,
            )
            report["auto_fix"] = {"before": before, **fix_result}
        reports.append(report)
    summary = {
        "files": reports,
        "error_count": sum(len(r["errors"]) for r in reports),
        "warning_count": sum(len(r["warnings"]) for r in reports),
    }
    if args.json_report:
        args.json_report.parent.mkdir(parents=True, exist_ok=True)
        args.json_report.write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")

    for report in reports:
        status = "FAIL" if report["errors"] else "WARN" if report["warnings"] else "PASS"
        print(f"{status} {report['kind']} {report['path']}")
        for err in report["errors"]:
            print(f"  ERROR {err['code']}: {err['message']}")
        for warn in report["warnings"]:
            print(f"  WARN {warn['code']}: {warn['message']}")
        if "render" in report:
            render = report["render"]
            detail = render.get("reason") or render.get("pdf") or render.get("sampled_page_numbers") or ""
            print(f"  render: {render.get('status')} {detail}")
        if "preview_png" in report:
            preview = report["preview_png"]
            pngs = preview.get("pngs") or []
            print(f"  preview_png: {preview.get('status')} count={len(pngs)}")
        if "auto_fix" in report:
            fix = report["auto_fix"]
            print(f"  auto_fix: {'applied' if fix.get('applied') else 'none'} changes={fix.get('changes', 0)}")

    if summary["error_count"]:
        return 1
    if summary["warning_count"] and not args.allow_warnings:
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
