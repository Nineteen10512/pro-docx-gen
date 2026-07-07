"""Translation helpers for preserving DOCX structure while replacing text.

This module does not perform machine translation itself. It exposes a stable
extract/apply workflow so an LLM can translate collected segments and then
write them back into a copy of the original document with minimal visual
change.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


TRANSLATION_SELF_AUDIT_RULES = [
    "Translate only the meaning already present in the source file.",
    "Do not add any fact, example, explanation, qualifier, summary, inferred intent, or background detail that is not explicitly present in the source file.",
    "If the source text is ambiguous, preserve the ambiguity or translate conservatively instead of filling gaps.",
    "Preserve numbers, dates, units, names, citations, formulas, labels, figure numbers, and placeholders unless the user explicitly asked to localize them.",
    "Return one translation for each segment id and do not invent or drop ids.",
]


def build_translation_prompt(
    segments,
    *,
    source_lang: str | None = None,
    target_lang: str | None = None,
    bilingual: bool = False,
) -> str:
    """Build a default LLM prompt for DOCX segment translation."""

    payload = json.dumps(list(segments), ensure_ascii=False, indent=2)
    source_label = source_lang or "source language"
    target_label = target_lang or "target language"
    mode = "bilingual" if bilingual else "replace"
    rules = "\n".join(f"{idx}. {rule}" for idx, rule in enumerate(TRANSLATION_SELF_AUDIT_RULES, start=1))
    return (
        "You are a document translation engine for DOCX segments.\n"
        f"Task: translate the provided segments from {source_label} to {target_label}.\n"
        f"Mode: {mode}.\n"
        "Hard rules:\n"
        f"{rules}\n\n"
        "Output format:\n"
        '- Return JSON only.\n'
        '- Return an array of objects in the form {"id": "...", "translation": "..."}.\n'
        "- Keep ids unchanged.\n"
        "- Do not include commentary, markdown, or extra keys.\n\n"
        f"Segments:\n{payload}\n"
    )


def _iter_body_paragraphs(doc: Document):
    for index, paragraph in enumerate(doc.paragraphs):
        yield {
            "id": f"body:p:{index}",
            "scope": "body",
            "kind": "paragraph",
            "text": paragraph.text,
        }, paragraph


def _iter_table_paragraphs(doc: Document):
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                for para_index, paragraph in enumerate(cell.paragraphs):
                    yield {
                        "id": f"table:{table_index}:r:{row_index}:c:{cell_index}:p:{para_index}",
                        "scope": "table",
                        "kind": "table_paragraph",
                        "table_index": table_index,
                        "row_index": row_index,
                        "cell_index": cell_index,
                        "text": paragraph.text,
                    }, paragraph


def _iter_section_part_paragraphs(doc: Document, include_headers: bool, include_footers: bool):
    for section_index, section in enumerate(doc.sections):
        if include_headers:
            for para_index, paragraph in enumerate(section.header.paragraphs):
                yield {
                    "id": f"section:{section_index}:header:p:{para_index}",
                    "scope": "header",
                    "kind": "paragraph",
                    "text": paragraph.text,
                }, paragraph
        if include_footers:
            for para_index, paragraph in enumerate(section.footer.paragraphs):
                yield {
                    "id": f"section:{section_index}:footer:p:{para_index}",
                    "scope": "footer",
                    "kind": "paragraph",
                    "text": paragraph.text,
                }, paragraph


def _iter_segment_records(
    doc: Document,
    *,
    include_tables: bool = True,
    include_headers: bool = True,
    include_footers: bool = True,
):
    yield from _iter_body_paragraphs(doc)
    if include_tables:
        yield from _iter_table_paragraphs(doc)
    yield from _iter_section_part_paragraphs(
        doc,
        include_headers=include_headers,
        include_footers=include_footers,
    )


def _normalize_output_path(input_path: str, output_path: str | None, target_lang: str | None) -> str:
    if output_path:
        return output_path
    src = Path(input_path)
    suffix = f".{target_lang}" if target_lang else ".translated"
    return str(src.with_name(f"{src.stem}{suffix}{src.suffix}"))


def _replace_paragraph_text_preserve_runs(paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return

    editable_runs = [run for run in runs if not _run_has_protected_content(run)]
    if not editable_runs:
        paragraph.add_run(text)
        return

    anchor = None
    for run in editable_runs:
        if _run_has_text_content(run):
            anchor = run
            break
    if anchor is None:
        anchor = editable_runs[0]

    for run in editable_runs:
        _clear_text_nodes(run._r)

    _append_text_to_run(anchor._r, text)


_PROTECTED_DESCENDANT_TAGS = {
    qn("w:drawing"),
    qn("w:pict"),
    qn("w:object"),
    qn("w:fldChar"),
    qn("w:instrText"),
    qn("w:footnoteReference"),
    qn("w:endnoteReference"),
    qn("w:commentReference"),
}

_TEXT_NODE_TAGS = {
    qn("w:t"),
    qn("w:tab"),
    qn("w:br"),
    qn("w:cr"),
    qn("w:noBreakHyphen"),
    qn("w:softHyphen"),
}


def _run_has_protected_content(run) -> bool:
    for node in run._r.iter():
        if node is run._r:
            continue
        if node.tag in _PROTECTED_DESCENDANT_TAGS:
            return True
    return False


def _run_has_text_content(run) -> bool:
    for node in run._r.iter():
        if node is run._r:
            continue
        if node.tag == qn("w:t") and node.text:
            return True
    return False


def _clear_text_nodes(run_el) -> None:
    for child in list(run_el):
        if child.tag in _TEXT_NODE_TAGS:
            run_el.remove(child)


def _append_text_to_run(run_el, text: str) -> None:
    if not text:
        return

    buffer: list[str] = []

    def flush_buffer():
        if not buffer:
            return
        segment = "".join(buffer)
        t = OxmlElement("w:t")
        if (
            segment.startswith(" ")
            or segment.endswith(" ")
            or "  " in segment
        ):
            t.set(qn("xml:space"), "preserve")
        t.text = segment
        run_el.append(t)
        buffer.clear()

    for ch in text:
        if ch == "\n":
            flush_buffer()
            run_el.append(OxmlElement("w:br"))
        elif ch == "\t":
            flush_buffer()
            run_el.append(OxmlElement("w:tab"))
        elif ch == "\r":
            continue
        else:
            buffer.append(ch)

    flush_buffer()


def _normalize_translation_map(translations) -> dict[str, str]:
    if isinstance(translations, dict):
        return {str(key): str(value) for key, value in translations.items()}

    mapping: dict[str, str] = {}
    for item in translations:
        if not isinstance(item, dict):
            continue
        seg_id = item.get("id")
        if not seg_id:
            continue
        value = item.get("translation")
        if value is None:
            value = item.get("text")
        if value is None:
            continue
        mapping[str(seg_id)] = str(value)
    return mapping


def _estimate_text_growth(source_text: str, translated_text: str) -> float:
    source_len = max(len((source_text or "").strip()), 1)
    target_len = len((translated_text or "").strip())
    return target_len / source_len


def _table_format_action(growth_ratio: float, column_count: int) -> str:
    if growth_ratio >= 2.2 or (growth_ratio >= 1.7 and column_count >= 4):
        return "high"
    if growth_ratio >= 1.35:
        return "medium"
    return "low"


def _iter_tables(doc: Document):
    for table_index, table in enumerate(doc.tables):
        yield table_index, table


def _scale_table_fonts(table, scale: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is None:
                        continue
                    run.font.size = Pt(max(8, round(run.font.size.pt * scale, 1)))


def _format_docx_table(table, severity: str) -> None:
    table.autofit = True
    if severity == "high":
        _scale_table_fonts(table, 0.9)
    elif severity == "medium":
        _scale_table_fonts(table, 0.95)


def assess_translation_risk(docx_path: str, translations) -> dict:
    """Estimate whether translated table content is likely to disturb layout."""

    mapping = _normalize_translation_map(translations)
    doc = Document(docx_path)
    table_reports: dict[int, dict] = {}

    for record, _paragraph in _iter_segment_records(doc):
        if record.get("scope") != "table":
            continue
        translated = mapping.get(record["id"])
        if translated is None:
            continue
        table_index = record["table_index"]
        table_report = table_reports.setdefault(
            table_index,
            {
                "table_index": table_index,
                "max_growth_ratio": 1.0,
                "column_count": 0,
                "risk": "low",
                "cells": [],
            },
        )
        growth_ratio = _estimate_text_growth(record["text"], translated)
        table = doc.tables[table_index]
        column_count = len(table.columns) if table.columns else 0
        severity = _table_format_action(growth_ratio, column_count)
        table_report["column_count"] = column_count
        table_report["max_growth_ratio"] = max(table_report["max_growth_ratio"], growth_ratio)
        if severity == "high" or (severity == "medium" and table_report["risk"] == "low"):
            table_report["risk"] = severity
        if severity != "low":
            table_report["cells"].append(
                {
                    "id": record["id"],
                    "row_index": record["row_index"],
                    "cell_index": record["cell_index"],
                    "growth_ratio": round(growth_ratio, 3),
                    "risk": severity,
                }
            )

    tables = list(table_reports.values())
    return {
        "has_risk": any(item["risk"] != "low" for item in tables),
        "tables": tables,
    }


def collect_translation_segments(
    docx_path: str,
    *,
    include_tables: bool = True,
    include_headers: bool = True,
    include_footers: bool = True,
    skip_empty: bool = True,
) -> list[dict]:
    """Collect stable text segments from a DOCX file for external translation."""

    doc = Document(docx_path)
    segments: list[dict] = []
    for record, _paragraph in _iter_segment_records(
        doc,
        include_tables=include_tables,
        include_headers=include_headers,
        include_footers=include_footers,
    ):
        text = record["text"]
        if skip_empty and not text.strip():
            continue
        segments.append(record)
    return segments


def apply_translation_map(
    docx_path: str,
    translations,
    *,
    output_path: str | None = None,
    target_lang: str | None = None,
    include_tables: bool = True,
    include_headers: bool = True,
    include_footers: bool = True,
    auto_format_tables: bool = True,
) -> str:
    """Write translated text back into a copy of the original DOCX file."""

    mapping = _normalize_translation_map(translations)
    doc = Document(docx_path)
    risk_report = assess_translation_risk(docx_path, mapping)
    severity_by_table = {item["table_index"]: item["risk"] for item in risk_report["tables"]}

    for record, paragraph in _iter_segment_records(
        doc,
        include_tables=include_tables,
        include_headers=include_headers,
        include_footers=include_footers,
    ):
        replacement = mapping.get(record["id"])
        if replacement is None:
            continue
        _replace_paragraph_text_preserve_runs(paragraph, replacement)

    if auto_format_tables:
        for table_index, table in _iter_tables(doc):
            severity = severity_by_table.get(table_index)
            if severity:
                _format_docx_table(table, severity)

    resolved_output = _normalize_output_path(docx_path, output_path, target_lang)
    doc.save(resolved_output)
    return resolved_output
